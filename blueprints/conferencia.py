"""
blueprints/conferencia.py — Módulo Conferência de Contrato Social v5
Relatório comparativo e documental — SEM análise jurídica por IA.
A conferência principal NÃO chama Claude/Anthropic.
"""
import json
import os
import re
import tempfile
import threading
import time
import unicodedata
import uuid
from datetime import date, datetime, timedelta

from flask import (Blueprint, jsonify, redirect, render_template,
                   request, session, url_for)
from werkzeug.utils import secure_filename

from blueprints.auth import login_obrigatorio

conferencia_bp = Blueprint('conferencia', __name__, url_prefix='/conferencia')

# ─── Configuração ──────────────────────────────────────────────────────────────
USE_AI_FOR_MAIN_CONFERENCE = True  # Conferência principal com IA (claude-sonnet-4-6)

# ─── Sistema de jobs assíncronos ───────────────────────────────────────────────
# O frontend recebe job_id imediatamente e faz polling a cada 2s.
# O relatório é gerado em thread separada (Python puro, sem Claude).

_jobs: dict = {}
_jobs_lock = threading.Lock()


def _cleanup_old_jobs():
    """Remove jobs com mais de 2 horas para liberar memória."""
    cutoff = datetime.now() - timedelta(hours=2)
    with _jobs_lock:
        antigos = [jid for jid, j in _jobs.items() if j['criado'] < cutoff]
        for jid in antigos:
            del _jobs[jid]


ALLOWED = {'pdf', 'docx', 'doc'}

MESES_PT = {
    'janeiro': 1, 'fevereiro': 2, 'março': 3, 'marco': 3,
    'abril': 4, 'maio': 5, 'junho': 6, 'julho': 7,
    'agosto': 8, 'setembro': 9, 'outubro': 10,
    'novembro': 11, 'dezembro': 12,
}

TIPOS_DOC = ['viabilidade', 'fcpj', 'dbe', 'fcn', 'uso_do_solo', 'numero_predial']

LABELS_DOC = {
    'viabilidade':    'Consulta Prévia / Viabilidade',
    'fcpj':           'FCPJ (Receita Federal)',
    'dbe':            'DBE (Documento Básico de Entrada)',
    'fcn':            'FCN (Ficha de Cadastro Nacional)',
    'uso_do_solo':    'Uso do Solo',
    'numero_predial': 'Número Predial',
}

TIPO_PROCESSO_LABEL = {
    'abertura':                    'Abertura de Empresa',
    'alteracao_contratual':        'Alteração Contratual',
    'baixa':                       'Baixa / Encerramento',
    'transformacao':               'Transformação Societária',
    'alteracao_natureza_juridica': 'Alteração de Natureza Jurídica',
}

AVISO_JURIDICO_FINAL = (
    '<div class="item info" style="margin-top:24px; border-left-color:#2563eb; background:#eff6ff;">'
    '<strong>Aviso — Análise Jurídica:</strong> Esta conferência possui natureza exclusivamente '
    'documental e comparativa. Ela não substitui a análise jurídica das cláusulas contratuais '
    'pela liderança do setor ou pelo profissional responsável.'
    '</div>'
)

# Termos que indicam trecho de qualificação civil — invalidam nome de sócio extraído
TERMOS_QUALIFICACAO_CIVIL = [
    'expedida', 'expedido', 'SSP', 'SSPII', 'SSPC', 'SSPDF', 'SESPII',
    'Carteira de Identidade', 'carteira de identidade',
    'portador', 'portadora',
    'residente', 'domiciliado', 'domiciliada',
    'brasileiro', 'brasileira', 'estrangeiro', 'estrangeira',
    'casado', 'casada', 'solteiro', 'solteira', 'viúvo', 'viúva',
    'divorciado', 'divorciada', 'separado', 'separada',
    'profissão', 'profissao', 'natural de',
    'filiado', 'filiada',
]

# ─── Configuração IA ──────────────────────────────────────────────────────────
MODELO_IA = 'claude-sonnet-4-6'
TEMPERATURA_IA = 0
MAX_TOKENS_IA = 8000

TERMOS_JURIDICOS_PROIBIDOS = [
    'risco jurídico', 'risco juridico',
    'juridicamente impreciso', 'juridicamente incorreto',
    'segurança jurídica', 'seguranca juridica',
    'recomenda-se alterar', 'recomenda-se revisar',
    'validade jurídica', 'validade juridica',
    'juridicamente inválido', 'nulo de pleno direito',
    'vício contratual', 'vicio contratual',
    'irregularidade jurídica', 'irregularidade juridica',
    'pode ser questionado judicialmente',
]

STATUS_VALIDOS_IA = {
    'Conforme', 'Divergente', 'Ausente', 'Ignorado pelo usuário',
    'Não localizado', 'Atenção para conferência manual',
    'Possível falha de extração', 'Não aplicável',
    'Alterada conforme evento', 'Nova cláusula',
}

SYSTEM_PROMPT_CONFERENCIA = """\
Você é um robô conferente documental e comparativo para escritório contábil e societário.
Sua única função é: extrair informações objetivas dos documentos recebidos, compará-las entre si \
e reportar divergências, ausências e confirmações como fatos concretos.

VEDAÇÕES ABSOLUTAS — NUNCA faça nenhuma das seguintes ações:
1. Não faça análise jurídica de nenhuma cláusula contratual.
2. Não emita parecer jurídico de nenhuma espécie.
3. Não sugira reescrita, reformulação ou alteração de cláusulas.
4. Não avalie se uma cláusula é válida, inválida, adequada ou inadequada do ponto de vista jurídico.
5. Não use os termos: "risco jurídico", "juridicamente impreciso", \
"recomenda-se alterar por segurança jurídica" ou expressões equivalentes.
6. Não crie obrigatoriedade automática de documentos que o usuário não marcou como obrigatórios.
7. Não classifique baixa de filial como baixa total da empresa.
8. Não invente dados ausentes — use "Não localizado" quando não encontrar.
9. Não avalie se o objeto social é adequado, específico, genérico ou abrangente.
10. Não critique cláusulas por conteúdo jurídico.
11. Não emita opinião sobre decisões societárias (ingresso, retirada, cessão de quotas).
12. Não use os termos "juridicamente incorreto", "nulidade", "vício", "irregularidade jurídica".
13. Não diga que algo "pode ser questionado judicialmente" ou expressão similar.
14. Não mencione "segurança jurídica" como critério de avaliação.
15. Na dúvida sobre qualquer extração ou comparação, use status "Atenção para conferência manual".

REGRAS POSITIVAS:
- Extraia e compare apenas fatos objetivos presentes nos documentos.
- Se uma informação não estiver presente, retorne "Não localizado".
- Se houver dúvida na extração, retorne "Possível falha de extração".
- Ignore diferenças irrelevantes de formatação (espaços extras, quebras de linha, caixa alta/baixa).
- Ao comparar consolidações, aponte apenas se o conteúdo mudou, não mudou, desapareceu ou surgiu.

REGRAS PARA NOMES DE SÓCIOS:
NÃO aceite como nome de sócio nenhum destes trechos:
"expedida pela SSP", "expedida pela SSPII", "portador da Carteira de Identidade",
"portadora da Carteira de Identidade", "Carteira de Identidade", "RG", "CPF",
"residente", "domiciliado", "brasileiro", "casado", "solteiro", ou qualquer
fragmento de qualificação civil.
Se localizar CPF/CNPJ mas não conseguir identificar o nome com segurança:
retorne movimento "Não identificado", duvida_extracao: true.

STATUS PERMITIDOS — use APENAS estes dez status (exatamente como escritos abaixo):
- Conforme
- Divergente
- Ausente
- Ignorado pelo usuário
- Não localizado
- Atenção para conferência manual
- Possível falha de extração
- Não aplicável
- Alterada conforme evento
- Nova cláusula

RETORNO OBRIGATÓRIO:
Responda APENAS com um objeto JSON válido seguindo exatamente o schema fornecido.
Não inclua nenhum texto antes ou depois do JSON.
Não use markdown code blocks (não use ```json).
"""


# ─── Extração de texto ─────────────────────────────────────────────────────────

def _extrair_texto(caminho, ext):
    if ext == 'pdf':
        import pdfplumber
        partes = []
        with pdfplumber.open(caminho) as pdf:
            for pagina in pdf.pages:
                t = pagina.extract_text()
                if t:
                    partes.append(t)
        return '\n'.join(partes)
    elif ext in ('docx', 'doc'):
        from docx import Document
        try:
            doc = Document(caminho)
            partes = []
            for p in doc.paragraphs:
                if p.text.strip():
                    partes.append(p.text)
            for tabela in doc.tables:
                for linha in tabela.rows:
                    celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
                    if celulas:
                        partes.append(' | '.join(celulas))
            return '\n'.join(partes)
        except Exception:
            if ext == 'doc':
                import subprocess
                for cmd in (['antiword', caminho], ['catdoc', caminho]):
                    try:
                        r = subprocess.run(cmd, capture_output=True, text=True, timeout=15)
                        if r.returncode == 0 and r.stdout.strip():
                            return r.stdout
                    except (FileNotFoundError, subprocess.TimeoutExpired):
                        continue
            raise
    return ''


def _salvar_e_extrair(arquivo):
    if not arquivo or arquivo.filename == '':
        return None
    nome = secure_filename(arquivo.filename)
    ext = nome.rsplit('.', 1)[1].lower() if '.' in nome else ''
    if ext not in ALLOWED:
        return None
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}') as tmp:
        arquivo.save(tmp.name)
        try:
            texto = _extrair_texto(tmp.name, ext)
        finally:
            os.unlink(tmp.name)
    return texto.strip() if texto.strip() else None


def _extrair_data_viabilidade(texto):
    t = texto.lower()
    for m in re.finditer(r'(\d{1,2})\s+de\s+([a-zç]+)\s+de\s+(\d{4})', t):
        d, mes_nome, a = m.groups()
        mes = MESES_PT.get(mes_nome.strip())
        if mes:
            try:
                return date(int(a), mes, int(d))
            except ValueError:
                continue
    for m in re.finditer(r'(\d{1,2})[/\-](\d{1,2})[/\-](\d{4})', t):
        d, mes, a = m.groups()
        try:
            return date(int(a), int(mes), int(d))
        except ValueError:
            continue
    return None


# ─── Parsing heurístico — classificação do processo ───────────────────────────

def _extrair_heuristicas(texto):
    """Extrai dados estruturados via regex/heurísticas.
    Usado exclusivamente para classificar o tipo de processo (cabeçalho do relatório).
    NÃO abre blocos de estabelecimento automaticamente.
    """
    t = texto[:30000]
    resultado = {
        'cnpjs': [],
        'nires': [],
        'tem_filial': False,
        'filiais_detectadas': [],
        'tipo_inferido': None,
        'numero_alteracao': None,
        'razao_social': None,
        'alteracoes_detectadas': [],
    }

    # CNPJs
    cnpjs_raw = re.findall(r'\b\d{2}[.\s]?\d{3}[.\s]?\d{3}[/\s]?\d{4}[-\s]?\d{2}\b', t)
    seen_cnpj = set()
    cnpjs = []
    for c in cnpjs_raw:
        c_clean = re.sub(r'[^\d]', '', c)
        if len(c_clean) == 14 and c_clean not in seen_cnpj:
            seen_cnpj.add(c_clean)
            cnpjs.append(f'{c_clean[:2]}.{c_clean[2:5]}.{c_clean[5:8]}/{c_clean[8:12]}-{c_clean[12:]}')
    resultado['cnpjs'] = cnpjs

    # NIREs
    nires_raw = re.findall(r'NIRE[:\s#nº°]*([0-9][\d.\-/]{5,18})', t, re.IGNORECASE)
    resultado['nires'] = list(dict.fromkeys([n.strip() for n in nires_raw]))

    # Menção a filial
    resultado['tem_filial'] = bool(re.search(r'\bfilial\b', t, re.IGNORECASE))

    # Tipo de processo — ordem da verificação é relevante
    if re.search(
        r'(contrato\s+de\s+constitu[ií][çc][aã]o|ato\s+constitutivo|'
        r'constitui[çc][aã]o\s+de\s+sociedade)',
        t, re.IGNORECASE,
    ):
        resultado['tipo_inferido'] = 'abertura'

    elif re.search(
        r'(?:instrumento\s+particular|instrumento\s+p[uú]blico)[^\n]{0,80}altera[çc][aã]o',
        t, re.IGNORECASE,
    ):
        # Alteração contratual explícita tem precedência sobre qualquer "baixa" mencionada
        # (ex: alteração que inclui baixa de filial deve ser classificada como alteração)
        resultado['tipo_inferido'] = 'alteracao_contratual'

    elif re.search(r'(?:altera[çc][aã]o|altera-se|passa\s+a\s+ter)', t, re.IGNORECASE):
        resultado['tipo_inferido'] = 'alteracao_contratual'

    elif re.search(
        r'\b(?:distrato|dissolu[çc][aã]o|liquida[çc][aã]o|extin[çc][aã]o)\b'
        r'[^\n]{0,80}\b(?:sociedad|empresa|s\.?\s*a\.?|ltda|eireli|s\.?\s*s\.?)\b',
        t, re.IGNORECASE,
    ) or re.search(
        r'\bencerramento\b[^\n]{0,80}\b(?:sociedad|empresa|cnpj|registro\s+empres)\b',
        t, re.IGNORECASE,
    ):
        # Baixa/encerramento TOTAL da empresa — distrato, dissolução, extinção
        # Não inclui "baixa de filial" dentro de uma alteração contratual
        resultado['tipo_inferido'] = 'baixa'

    elif re.search(
        r'transforma[çc][aã]o\s+(?:social|da\s+sociedade|de\s+tipo)',
        t, re.IGNORECASE,
    ):
        resultado['tipo_inferido'] = 'transformacao'

    # Número da alteração — usar função unificada para consistência com o relatório
    num_alt, conf_alt = _extrair_numero_alteracao_unificado(t)
    resultado['numero_alteracao'] = num_alt
    resultado['numero_alteracao_confianca'] = conf_alt

    # Razão social
    for pattern in [
        r'(?:denominada|denominado|raz[aã]o\s+social\s*[:\-]?\s*)([A-ZÁÉÍÓÚÂÊÔÃÕÇÜÀ][^\n,;]{3,80})',
        r'empresa\s+([A-ZÁÉÍÓÚÂÊÔÃÕÇÜÀ][^\n,;]{3,60})\s+(?:LTDA|S\.?A\.?|EIRELI|S\.?S\.?)',
    ]:
        m = re.search(pattern, t, re.IGNORECASE)
        if m:
            rs = m.group(1).strip().rstrip('.,;')
            if len(rs) > 3:
                resultado['razao_social'] = rs
                break

    # Tipos de alteração detectados (informativo apenas)
    ctx_alt = r'(?:altera-se|passa\s+a\s+(?:ter|ser|constar|vigorar)|fica\s+(?:alterado|alterada)|nova\s+reda[çc][aã]o)'
    kw_map = {
        'quadro_societario': r'(?:ingresso|admissão|retirada|retira|cede\s+e\s+transfer|cessão|cedente|cessionário|s[oó]cio)',
        'atividade':         r'(?:objeto\s+social|atividade(?:s)?\s+econ[oô]mica|cnae)',
        'endereco':          r'(?:endere[çc]o|sede|logradouro|cep|bairro)',
        'capital':           r'capital\s+social',
        'nome_empresarial':  r'(?:raz[aã]o\s+social|nome\s+empresarial|denomina[çc][aã]o)',
    }
    alt_detectadas = []
    for tipo_alt, kw in kw_map.items():
        if re.search(ctx_alt + r'[^\n]{0,300}' + kw, t, re.IGNORECASE | re.DOTALL):
            alt_detectadas.append(tipo_alt)
        elif re.search(kw + r'[^\n]{0,300}' + ctx_alt, t, re.IGNORECASE | re.DOTALL):
            alt_detectadas.append(tipo_alt)
    resultado['alteracoes_detectadas'] = alt_detectadas if alt_detectadas else ['outros']

    return resultado


# ─── Utilitários de comparação ─────────────────────────────────────────────────

def _normalizar(texto):
    """Normaliza texto para comparação: minúsculas, sem acentos, sem espaços duplos."""
    if not texto:
        return ''
    t = texto.lower().strip()
    t = unicodedata.normalize('NFD', t)
    t = ''.join(c for c in t if unicodedata.category(c) != 'Mn')
    t = re.sub(r'\s+', ' ', t)
    return t


def _normalizar_cnpj(cnpj):
    return re.sub(r'[^\d]', '', cnpj) if cnpj else ''


def _cnpjs_iguais(a, b):
    na, nb = _normalizar_cnpj(a), _normalizar_cnpj(b)
    return bool(na) and na == nb


def _extrair_cnpjs(texto):
    cnpjs_raw = re.findall(r'\b\d{2}[.\s]?\d{3}[.\s]?\d{3}[/\s]?\d{4}[-\s]?\d{2}\b', texto)
    seen, result = set(), []
    for c in cnpjs_raw:
        c_clean = re.sub(r'[^\d]', '', c)
        if len(c_clean) == 14 and c_clean not in seen:
            seen.add(c_clean)
            result.append(f'{c_clean[:2]}.{c_clean[2:5]}.{c_clean[5:8]}/{c_clean[8:12]}-{c_clean[12:]}')
    return result


def _extrair_nires(texto):
    nires = re.findall(r'NIRE[:\s#nº°]*([0-9][\d.\-/]{5,18})', texto, re.IGNORECASE)
    return list(dict.fromkeys([re.sub(r'[^\d]', '', n) for n in nires]))


def _extrair_numero_alteracao(texto):
    """Compatibilidade retroativa — usa a função unificada."""
    num, _ = _extrair_numero_alteracao_unificado(texto)
    return num


def _extrair_numero_alteracao_unificado(texto: str):
    """Extrai número da alteração contratual com grau de confiança.

    Retorna (numero: int | None, confianca: str).
    confianca: 'Alta' | 'Média' | 'Baixa' | 'Não localizado'
    """
    patterns_alta = [
        r'(\d+)[ªº°]\s*(?:altera[çc][aã]o\s+(?:ao\s+|do\s+)?contrato\s+social)',
        r'(\d+)[ªº°]\s*(?:altera[çc][aã]o\s+contratual)',
        r'instrumento\s+particular[^\n]{0,80}?(\d+)[ªº°]\s*altera[çc][aã]o',
    ]
    patterns_media = [
        r'(\d+)[ªº°]\s+altera[çc][aã]o\b',
        r'(\d+)[ªº°]\s*altera[çc][aã]o',
        r'(\d+)[ªº°]?\s*(?:altera[çc][aã]o\s+(?:ao\s+|do\s+)?contrato)\b',
    ]
    patterns_baixa = [
        r'\b(\d+)[ªº°]\b[^\n]{0,40}contrato',
    ]
    for p in patterns_alta:
        m = re.search(p, texto, re.IGNORECASE)
        if m:
            try:
                return int(m.group(1)), 'Alta'
            except (ValueError, IndexError):
                continue
    for p in patterns_media:
        m = re.search(p, texto, re.IGNORECASE)
        if m:
            try:
                return int(m.group(1)), 'Média'
            except (ValueError, IndexError):
                continue
    for p in patterns_baixa:
        m = re.search(p, texto, re.IGNORECASE)
        if m:
            try:
                return int(m.group(1)), 'Baixa'
            except (ValueError, IndexError):
                continue
    return None, 'Não localizado'


def _extrair_cnaes(texto):
    return sorted(set(re.findall(r'\b\d{2}\.\d{2}-\d/\d{2}\b', texto)))


def _validar_nome_socio(nome: str) -> bool:
    """Retorna False se o nome contiver termos de qualificação civil."""
    if not nome or len(nome.strip()) < 5:
        return False
    nome_lower = nome.lower()
    for termo in TERMOS_QUALIFICACAO_CIVIL:
        if termo.lower() in nome_lower:
            return False
    # Nome deve ter pelo menos 2 palavras
    palavras = [p for p in nome.strip().split() if len(p) > 1]
    if len(palavras) < 2:
        return False
    # Não pode começar com preposição ou artigo
    if re.match(r'^(de|da|do|das|dos|a|o|e|em|por|para|com|ao|à)\b',
                nome.strip(), re.IGNORECASE):
        return False
    return True


def _extrair_nome_do_contexto(contexto: str):
    """Extrai o nome próprio mais próximo ao CPF no contexto anterior.

    Retorna (nome: str, confianca: str).
    """
    # Padrão 1 (Alta): Nome seguido de vírgula + qualificação civil
    m = re.search(
        r'([A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-ZÁÉÍÓÚÂÊÔÃÕÇa-záéíóúâêôãõç]{1,}(?:\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇa-záéíóúâêôãõç]{2,}){1,6})'
        r'\s*,\s*(?:brasileiro|brasileira|casado|casada|solteiro|solteira|portador|portadora|'
        r'divorciado|divorciada|viúvo|viúva|estrangeiro|estrangeira)',
        contexto, re.IGNORECASE,
    )
    if m:
        nome = m.group(1).strip()
        if _validar_nome_socio(nome):
            return nome, 'Alta'

    # Padrão 2 (Alta): Nome todo em maiúsculas (comum em contratos)
    m2 = re.search(
        r'([A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,}(?:\s+(?:DE|DA|DO|DAS|DOS|E)\s+)?'
        r'(?:[A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,}\s*){1,6})'
        r'\s*,?\s*(?:brasileiro|brasileira|portador|portadora|casado|casada)',
        contexto,
    )
    if m2:
        nome = m2.group(1).strip().rstrip(',')
        if _validar_nome_socio(nome):
            return nome, 'Alta'

    # Padrão 3 (Média): Nome antes de "inscrito" ou "inscrita"
    m3 = re.search(
        r'([A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-ZÁÉÍÓÚÂÊÔÃÕÇa-záéíóúâêôãõç]{1,}(?:\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇa-záéíóúâêôãõç]{2,}){1,6})'
        r'\s*,?\s*inscri[to|ta]',
        contexto, re.IGNORECASE,
    )
    if m3:
        nome = m3.group(1).strip()
        if _validar_nome_socio(nome):
            return nome, 'Média'

    # Padrão 4 (Baixa): Última sequência de palavras com inicial maiúscula
    candidatos = re.findall(
        r'([A-ZÁÉÍÓÚÂÊÔÃÕÇ][a-záéíóúâêôãõç]+(?:\s+(?:[A-ZÁÉÍÓÚÂÊÔÃÕÇ][a-záéíóúâêôãõç]+|DE|DA|DO|DAS|DOS|E)){1,6})',
        contexto,
    )
    if candidatos:
        nome = candidatos[-1].strip()
        if _validar_nome_socio(nome):
            return nome, 'Baixa'

    return '', 'Não localizado'


def _deduplicar_socios(lista: list) -> list:
    """Remove duplicatas por CPF, mantendo o de maior confiança."""
    visto = {}
    ordem_conf = {'Alta': 3, 'Média': 2, 'Baixa': 1, 'Não localizado': 0}
    for s in lista:
        cpf = s['cpf']
        if cpf not in visto:
            visto[cpf] = s
        else:
            if ordem_conf.get(s['confianca'], 0) > ordem_conf.get(visto[cpf]['confianca'], 0):
                visto[cpf] = s
    return list(visto.values())


def _extrair_socios(texto):
    """Extrai sócios com CPF do texto — versão v2 com validação de qualidade."""
    return _extrair_socios_v2(texto)


def _extrair_socios_v2(texto: str) -> list:
    """Extrai sócios com CPF, validando se o nome extraído é realmente um nome próprio."""
    resultado = []
    cpf_pattern = re.compile(r'(\d{3}\.?\d{3}\.?\d{3}[-/]?\d{2})')

    for m_cpf in cpf_pattern.finditer(texto):
        cpf_raw = m_cpf.group(1)
        cpf_norm = re.sub(r'[^\d]', '', cpf_raw)
        if len(cpf_norm) != 11:
            continue

        # Contexto antes do CPF (até 400 chars)
        inicio = max(0, m_cpf.start() - 400)
        contexto = texto[inicio:m_cpf.start()]

        nome, confianca = _extrair_nome_do_contexto(contexto)
        valido = _validar_nome_socio(nome) if nome else False

        resultado.append({
            'cpf': cpf_norm,
            'nome': nome if nome else f'[CPF {cpf_norm} — nome não identificado]',
            'confianca': confianca,
            'valido': valido,
        })

    return _deduplicar_socios(resultado)


def _extrair_socios_legado(texto):
    """Método legado — mantido apenas para referência."""
    return []


def _extrair_endereco_contexto(texto, cnpj_ref=''):
    """Extrai endereço da sede ou de um estabelecimento específico."""
    m = re.search(
        r'(?:localizada|sediada|sede|endere[çc]o)[:\s]+(?:na|em|à)?\s*'
        r'([^\n,;]{10,150})',
        texto, re.IGNORECASE,
    )
    if m:
        return m.group(1).strip()
    m2 = re.search(
        r'(?:Rua|Av(?:enida)?|R\.|Travessa|Alameda|Praça)[,\s]+[^\n]{5,120}',
        texto, re.IGNORECASE,
    )
    return m2.group(0).strip() if m2 else None


# ─── Componentes HTML do relatório ────────────────────────────────────────────

def _esc(texto):
    if not texto:
        return ''
    return (str(texto)
        .replace('&', '&amp;')
        .replace('<', '&lt;')
        .replace('>', '&gt;')
        .replace('"', '&quot;'))


def _ok(msg):
    return f'<div class="item ok">{msg}</div>'


def _erro(msg):
    return f'<div class="item erro">{msg}</div>'


def _atencao(msg):
    return f'<div class="item atencao">{msg}</div>'


def _info(msg):
    return f'<div class="item info">{msg}</div>'


def _tag_doc(nome):
    return f'<span class="doc-tag">{_esc(nome)}</span>'


def _secao(titulo, itens):
    conteudo = '\n'.join(itens)
    return (
        f'<div class="secao">'
        f'<h2 class="sec-titulo">{_esc(titulo)}</h2>'
        f'{conteudo}'
        f'</div>'
    )


def _sub(titulo):
    return f'<h3 class="sec-sub">{_esc(titulo)}</h3>'


# ─── Funções de Conferência com IA ───────────────────────────────────────────

def _item_html(item: dict) -> str:
    """Renderiza um item do JSON da IA em HTML usando as classes existentes."""
    status = item.get('status', 'Não localizado')
    campo = item.get('campo', '')
    obs = item.get('observacao', '')
    conf = item.get('confianca', '')
    doc = item.get('documento', '')
    estab = item.get('estabelecimento', '')

    partes = []
    if doc:
        partes.append(_tag_doc(doc.replace('_', ' ').title()))
    if campo:
        partes.append(f'<strong>{_esc(campo)}</strong>')
    if estab:
        partes.append(f'({_esc(estab)})')

    linha1 = ' — '.join(partes) if partes else ''
    linha2 = _esc(obs) if obs else ''
    if conf and conf not in ('Alta',):
        linha2 += f' <span style="font-size:10px;color:#92400e;">[Confiança: {_esc(conf)}]</span>'

    conteudo = f'{linha1}: {linha2}' if linha1 and linha2 else (linha1 or linha2)
    st = status
    if st in ('Conforme', 'Alterada conforme evento', 'Nova cláusula'):
        return _ok(f'<strong>{_esc(st)}:</strong> {conteudo}')
    elif st in ('Divergente', 'Ausente'):
        return _erro(f'<strong>{_esc(st)}:</strong> {conteudo}')
    elif st in ('Atenção para conferência manual', 'Não localizado', 'Possível falha de extração'):
        return _atencao(f'<strong>{_esc(st)}:</strong> {conteudo}')
    else:
        return _info(f'<strong>{_esc(st)}:</strong> {conteudo}')


def _montar_prompt_usuario_ia(
    texto_minuta, texto_ultima, ultima_ignorada,
    docs_por_estab, estabelecimentos,
    tipo_processo, numero_alteracao, razao_social, cnpj_sociedade,
) -> str:
    """Monta a mensagem de usuário para a IA com todos os documentos e estrutura."""
    linhas = []

    linhas.append('=== CONTEXTO DO PROCESSO ===')
    linhas.append(f'Tipo de processo (inferido): {tipo_processo or "não identificado"}')
    if numero_alteracao:
        linhas.append(f'Número da alteração (inferido): {numero_alteracao}')
    if razao_social:
        linhas.append(f'Razão social (inferida): {razao_social}')
    if cnpj_sociedade:
        linhas.append(f'CNPJ da sociedade (inferido): {cnpj_sociedade}')
    linhas.append('')

    linhas.append('=== ESTRUTURA DE ESTABELECIMENTOS (criada manualmente pelo usuário) ===')
    for i, estab in enumerate(estabelecimentos):
        tipo_e = estab.get('tipo', 'matriz')
        cnpj_e = estab.get('cnpj', '') or ''
        nire_e = estab.get('nire', '') or ''
        descricao_e = estab.get('descricao', '') or ''
        linhas.append(f'Estabelecimento {i}: {tipo_e.upper()}')
        if cnpj_e:
            linhas.append(f'  CNPJ declarado: {cnpj_e}')
        if nire_e:
            linhas.append(f'  NIRE declarado: {nire_e}')
        if descricao_e:
            linhas.append(f'  Descrição: {descricao_e}')
    linhas.append('')

    linhas.append('=== STATUS DOS DOCUMENTOS ===')
    linhas.append(f'Minuta: {"PRESENTE" if texto_minuta else "AUSENTE"}')
    if ultima_ignorada:
        linhas.append('Última Alteração: IGNORADA PELO USUÁRIO')
    else:
        linhas.append(f'Última Alteração: {"PRESENTE" if texto_ultima else "AUSENTE"}')
    for i, (estab, edoc) in enumerate(zip(estabelecimentos, docs_por_estab)):
        tipo_e = estab.get('tipo', 'matriz')
        label_e = 'Matriz' if tipo_e == 'matriz' else f'Filial {i}'
        for tipo_doc in TIPOS_DOC:
            ignorado = edoc['ignorados'].get(tipo_doc, False)
            texto_doc = edoc['docs'].get(tipo_doc)
            label_doc = LABELS_DOC[tipo_doc]
            if ignorado:
                linhas.append(f'{label_e} — {label_doc}: IGNORADO PELO USUÁRIO')
            elif texto_doc:
                linhas.append(f'{label_e} — {label_doc}: PRESENTE')
            else:
                linhas.append(f'{label_e} — {label_doc}: AUSENTE')
    linhas.append('')

    linhas.append('=== SCHEMA JSON ESPERADO ===')
    linhas.append('''\
Retorne um objeto JSON com esta estrutura exata (sem texto antes ou depois, sem markdown):
{
  "cabecalho": {
    "razao_social": "string|null",
    "cnpj": "string|null",
    "nire": "string|null",
    "numero_alteracao": "number|null",
    "numero_alteracao_confianca": "Alta|Media|Baixa|Nao localizado",
    "tipo_processo": "string",
    "data_assinatura": "string|null",
    "local_assinatura": "string|null",
    "foro": "string|null"
  },
  "conferencia_geral": [
    {"campo": "string", "status": "status_valido", "observacao": "string",
     "entra_pendencias": false, "confianca": "string"}
  ],
  "por_estabelecimento": [
    {
      "indice": 0, "tipo": "matriz|filial", "titulo": "string",
      "itens": [
        {"documento": "string", "campo": "string", "status": "status_valido",
         "observacao": "string", "entra_pendencias": false, "entra_alertas": false}
      ]
    }
  ],
  "comparativo_consolidacao": [
    {"clausula": "string", "tema_provavel": "string", "status": "status_valido",
     "conteudo_ultima": "string|null", "conteudo_minuta": "string|null", "observacao": "string"}
  ],
  "quadro_societario": [
    {"nome": "string", "cpf_cnpj": "string",
     "condicao_anterior": "string|null", "condicao_minuta": "string|null",
     "movimento": "Remanescente|Ingressante|Retirante|Nao identificado",
     "confianca": "string", "duvida_extracao": false, "observacao": "string|null"}
  ],
  "comparativo_cnpj_nire": [
    {"campo": "string", "status": "status_valido", "observacao": "string", "entra_pendencias": false}
  ],
  "comparativo_enderecos": [
    {"estabelecimento": "string", "campo": "string",
     "valor_esperado": "string|null", "valor_encontrado": "string|null",
     "status": "status_valido", "observacao": "string",
     "entra_pendencias": false, "entra_alertas": false}
  ],
  "comparativo_atividades": [
    {"estabelecimento": "string", "campo": "string",
     "valor_esperado": "string|null", "valor_encontrado": "string|null",
     "status": "status_valido", "observacao": "string", "entra_pendencias": false}
  ],
  "revisao_textual": [
    {"campo": "string", "status": "status_valido", "observacao": "string", "entra_alertas": false}
  ],
  "ceps": [
    {"cep_encontrado": "string", "problema": "string", "ocorrencias": 1,
     "formato_esperado": "string", "status": "status_valido"}
  ],
  "pendencias_objetivas": ["string"],
  "alertas_manuais": ["string"],
  "falhas_extracao": ["string"],
  "documentos_ignorados_lista": ["string"],
  "conclusao": {
    "pendencias_count": 0, "alertas_count": 0,
    "falhas_count": 0, "ignorados_count": 0, "texto": "string"
  },
  "meta": {
    "ia_chamada": true, "modelo": "string",
    "data_hora": "string", "documentos_truncados": false
  }
}''')
    linhas.append('')

    linhas.append('=== TEXTOS DOS DOCUMENTOS ===')
    if texto_minuta:
        linhas.append('--- MINUTA DO CONTRATO SOCIAL ---')
        linhas.append(texto_minuta)
        linhas.append('')
    else:
        linhas.append('--- MINUTA DO CONTRATO SOCIAL: AUSENTE ---')
        linhas.append('')

    if ultima_ignorada:
        linhas.append('--- ÚLTIMA ALTERAÇÃO CONTRATUAL: IGNORADA PELO USUÁRIO ---')
        linhas.append('')
    elif texto_ultima:
        linhas.append('--- ÚLTIMA ALTERAÇÃO CONTRATUAL ---')
        linhas.append(texto_ultima)
        linhas.append('')
    else:
        linhas.append('--- ÚLTIMA ALTERAÇÃO CONTRATUAL: AUSENTE ---')
        linhas.append('')

    for i, (estab, edoc) in enumerate(zip(estabelecimentos, docs_por_estab)):
        tipo_e = estab.get('tipo', 'matriz')
        cnpj_e = estab.get('cnpj', '') or ''
        label_e = 'MATRIZ' if tipo_e == 'matriz' else f'FILIAL {i}'
        if cnpj_e:
            label_e += f' (CNPJ: {cnpj_e})'
        for tipo_doc in TIPOS_DOC:
            ignorado = edoc['ignorados'].get(tipo_doc, False)
            texto_doc = edoc['docs'].get(tipo_doc)
            label_doc = LABELS_DOC[tipo_doc].upper()
            if ignorado:
                linhas.append(f'--- {label_e} — {label_doc}: IGNORADO PELO USUÁRIO ---')
                linhas.append('')
            elif texto_doc:
                linhas.append(f'--- {label_e} — {label_doc} ---')
                linhas.append(texto_doc)
                linhas.append('')
            else:
                linhas.append(f'--- {label_e} — {label_doc}: AUSENTE ---')
                linhas.append('')

    return '\n'.join(linhas)


def _chamar_ia_conferencia(prompt_usuario: str, log: dict) -> dict:
    """Chama a API Anthropic e retorna o JSON da conferência."""
    import anthropic

    api_key = os.environ.get('ANTHROPIC_API_KEY', '')
    if not api_key:
        raise RuntimeError('ANTHROPIC_API_KEY não configurada no ambiente.')

    cliente = anthropic.Anthropic(api_key=api_key)

    t_inicio = time.time()
    resposta = cliente.messages.create(
        model=MODELO_IA,
        max_tokens=MAX_TOKENS_IA,
        temperature=TEMPERATURA_IA,
        system=SYSTEM_PROMPT_CONFERENCIA,
        messages=[{'role': 'user', 'content': prompt_usuario}],
    )
    log['ia_tempo_s'] = round(time.time() - t_inicio, 1)
    log['ia_input_tokens'] = resposta.usage.input_tokens
    log['ia_output_tokens'] = resposta.usage.output_tokens

    texto_resposta = resposta.content[0].text.strip()

    # Extrair JSON se vier com blocos markdown (proteção extra)
    if '```' in texto_resposta:
        m = re.search(r'```(?:json)?\s*([\s\S]+?)\s*```', texto_resposta)
        if m:
            texto_resposta = m.group(1).strip()

    dados = json.loads(texto_resposta)
    return dados


def _neutralizar_termo(texto: str) -> str:
    """Remove termos jurídicos proibidos de um texto de observação."""
    if not texto:
        return texto
    t_lower = texto.lower()
    for termo in TERMOS_JURIDICOS_PROIBIDOS:
        if termo in t_lower:
            return 'Atenção para conferência manual. Verificar manualmente.'
    return texto


def _validar_json_ia(dados: dict, log: dict) -> dict:
    """Valida campos obrigatórios e sanitiza termos proibidos no JSON da IA."""
    if not isinstance(dados, dict):
        raise ValueError('Resposta da IA não é um objeto JSON válido.')

    # Garantir chaves obrigatórias
    chaves_lista = [
        'conferencia_geral', 'por_estabelecimento', 'comparativo_consolidacao',
        'quadro_societario', 'comparativo_cnpj_nire', 'comparativo_enderecos',
        'comparativo_atividades', 'revisao_textual', 'ceps',
        'pendencias_objetivas', 'alertas_manuais', 'falhas_extracao',
        'documentos_ignorados_lista',
    ]
    for chave in chaves_lista:
        if chave not in dados or not isinstance(dados[chave], list):
            dados[chave] = []
    for chave in ('cabecalho', 'conclusao', 'meta'):
        if chave not in dados or not isinstance(dados[chave], dict):
            dados[chave] = {}

    violacoes = 0

    # Sanitizar itens simples
    for secao_key in ['conferencia_geral', 'comparativo_cnpj_nire',
                      'comparativo_enderecos', 'comparativo_atividades', 'revisao_textual']:
        for item in dados.get(secao_key, []):
            if not isinstance(item, dict):
                continue
            if item.get('status') not in STATUS_VALIDOS_IA:
                item['status'] = 'Atenção para conferência manual'
                violacoes += 1
            obs_orig = item.get('observacao', '')
            item['observacao'] = _neutralizar_termo(obs_orig)
            if item['observacao'] != obs_orig:
                violacoes += 1

    # Sanitizar por estabelecimento
    for estab in dados.get('por_estabelecimento', []):
        if not isinstance(estab, dict):
            continue
        for item in estab.get('itens', []):
            if not isinstance(item, dict):
                continue
            if item.get('status') not in STATUS_VALIDOS_IA:
                item['status'] = 'Atenção para conferência manual'
                violacoes += 1
            obs_orig = item.get('observacao', '')
            item['observacao'] = _neutralizar_termo(obs_orig)
            if item['observacao'] != obs_orig:
                violacoes += 1

    # Sanitizar comparativo de consolidação
    for item in dados.get('comparativo_consolidacao', []):
        if not isinstance(item, dict):
            continue
        if item.get('status') not in STATUS_VALIDOS_IA:
            item['status'] = 'Atenção para conferência manual'
            violacoes += 1
        obs_orig = item.get('observacao', '')
        item['observacao'] = _neutralizar_termo(obs_orig)
        if item['observacao'] != obs_orig:
            violacoes += 1

    log['ia_violacoes_sanitizadas'] = violacoes

    # Garantir meta
    meta = dados['meta']
    meta['ia_chamada'] = True
    meta['modelo'] = MODELO_IA
    meta.setdefault('data_hora', datetime.now().strftime('%d/%m/%Y %H:%M:%S'))
    meta.setdefault('documentos_truncados', False)

    return dados


def _json_para_html(
    dados: dict,
    nome_analista: str,
    data_hoje: str,
    tipo_processo: str,
    numero_alteracao,
    razao_social: str,
    cnpj_sociedade: str,
    estabelecimentos: list,
    ultima_ignorada: bool,
) -> str:
    """Converte o JSON da IA em HTML do relatório (21 seções)."""
    secoes = []

    cab = dados.get('cabecalho') or {}
    meta = dados.get('meta') or {}

    num_alt = cab.get('numero_alteracao') or numero_alteracao
    conf_num = cab.get('numero_alteracao_confianca', 'Média')
    tipo_label = TIPO_PROCESSO_LABEL.get(
        cab.get('tipo_processo') or tipo_processo,
        tipo_processo or 'Não identificado',
    )

    if num_alt:
        num_str = (
            f'Possível {num_alt}ª Alteração Contratual — confirmar manualmente'
            if conf_num == 'Baixa'
            else f'{num_alt}ª Alteração Contratual'
        )
    else:
        num_str = 'Número não identificado — conferir manualmente'

    rs_exib = cab.get('razao_social') or razao_social or '—'
    cnpj_exib = cab.get('cnpj') or cnpj_sociedade or '—'

    # ── SEÇÃO 1/2: RESUMO DO PROCESSO ─────────────────────────────────────────
    cab_itens = [
        _info(f'<strong>Analista:</strong> {_esc(nome_analista)}'),
        _info(f'<strong>Data da Conferência:</strong> {_esc(data_hoje)}'),
        _info(f'<strong>Tipo de Processo:</strong> {_esc(tipo_label)}'),
        _info(f'<strong>Numeração:</strong> {_esc(num_str)}'),
        _info(f'<strong>Razão Social:</strong> {_esc(rs_exib)}'),
        _info(f'<strong>CNPJ da Sociedade:</strong> {_esc(cnpj_exib)}'),
    ]
    if cab.get('nire'):
        cab_itens.append(_info(f'<strong>NIRE:</strong> {_esc(cab["nire"])}'))
    if cab.get('data_assinatura'):
        cab_itens.append(_info(f'<strong>Data de Assinatura:</strong> {_esc(cab["data_assinatura"])}'))
    if cab.get('local_assinatura'):
        cab_itens.append(_info(f'<strong>Local de Assinatura:</strong> {_esc(cab["local_assinatura"])}'))
    if cab.get('foro'):
        cab_itens.append(_info(f'<strong>Foro:</strong> {_esc(cab["foro"])}'))
    cab_itens.append(_info(
        f'<strong>Conferência realizada por IA:</strong> {_esc(meta.get("modelo", MODELO_IA))} '
        f'— {_esc(meta.get("data_hora", data_hoje))}'
    ))
    if meta.get('documentos_truncados'):
        cab_itens.append(_atencao(
            '<strong>Atenção para conferência manual:</strong> Um ou mais documentos foram '
            'lidos parcialmente. A conferência pode estar incompleta.'
        ))
    secoes.append(_secao('Resumo do Processo', cab_itens))

    # ── SEÇÃO 3: TIPO DE PROCESSO E EVENTOS ────────────────────────────────────
    tipo_itens = [
        _info(f'<strong>Tipo:</strong> {_esc(tipo_label)}'),
        _info(f'<strong>Instrumento:</strong> {_esc(num_str)}'),
    ]
    secoes.append(_secao('Tipo de Processo e Eventos Identificados', tipo_itens))

    # ── SEÇÕES 4/5/6: DOCUMENTOS ──────────────────────────────────────────────
    doc_recebidos = []
    doc_ignorados_sec = []
    doc_ausentes = []

    for item in dados.get('conferencia_geral', []):
        campo = item.get('campo', '').lower()
        st = item.get('status', '')
        if not any(k in campo for k in ['minuta', 'alteração', 'viabilidade',
                                        'fcpj', 'fcn', 'uso do solo', 'predial']):
            continue
        if st == 'Ignorado pelo usuário':
            doc_ignorados_sec.append(_item_html(item))
        elif st in ('Ausente', 'Não localizado'):
            doc_ausentes.append(_item_html(item))
        else:
            doc_recebidos.append(_item_html(item))

    # Fallback: doc_recebidos e ausentes da lista da IA
    for txt in dados.get('documentos_ignorados_lista', []):
        if not any(txt in d for d in doc_ignorados_sec):
            doc_ignorados_sec.append(_atencao(f'Ignorado pelo usuário: {_esc(txt)}'))

    if not doc_recebidos:
        doc_recebidos.append(_info('Documentos recebidos registrados na conferência geral.'))
    if not doc_ignorados_sec:
        doc_ignorados_sec.append(_info('Nenhum documento marcado como ignorado pelo usuário.'))
    if not doc_ausentes:
        doc_ausentes.append(_info('Nenhum documento ausente identificado.'))

    secoes.append(_secao('Documentos Anexados', doc_recebidos))
    secoes.append(_secao('Documentos Ignorados pelo Usuário — Status', doc_ignorados_sec))
    secoes.append(_secao('Documentos Ausentes', doc_ausentes))

    # ── SEÇÃO 7: CONFERÊNCIA GERAL DA MINUTA E ÚLTIMA ALTERAÇÃO ────────────────
    conf_geral = dados.get('conferencia_geral', [])
    geral_itens = [
        _item_html(i) for i in conf_geral
        if not any(k in i.get('campo', '').lower()
                   for k in ['minuta', 'alteração', 'viabilidade',
                              'fcpj', 'fcn', 'uso do solo', 'predial'])
    ]
    if not geral_itens:
        geral_itens.append(_info(
            'Conferência geral não realizada — minuta não anexada ou sem documentos suficientes.'
        ))
    secoes.append(_secao('Conferência da Minuta e Última Alteração', geral_itens))

    # ── SEÇÃO 8: COMPARATIVO DA CONSOLIDAÇÃO CONTRATUAL ────────────────────────
    consol = dados.get('comparativo_consolidacao', [])
    if consol:
        consol_itens = []
        for item in consol:
            status = item.get('status', 'Não localizado')
            clausula = item.get('clausula', '')
            tema = item.get('tema_provavel', '')
            obs = item.get('observacao', '')
            cont_ult = item.get('conteudo_ultima') or ''
            cont_min = item.get('conteudo_minuta') or ''

            linha = f'<strong>{_esc(clausula)}</strong>'
            if tema:
                linha += f' — <em>{_esc(tema)}</em>'
            if obs:
                linha += f': {_esc(obs)}'

            if status in ('Conforme', 'Alterada conforme evento', 'Nova cláusula'):
                consol_itens.append(_ok(f'<strong>{_esc(status)}:</strong> {linha}'))
            elif status in ('Divergente', 'Ausente'):
                detalhes = ''
                if cont_ult:
                    detalhes += (
                        f'<br><span style="color:#666;font-size:12px;">'
                        f'Última alteração: {_esc(cont_ult[:300])}</span>'
                    )
                if cont_min:
                    detalhes += (
                        f'<br><span style="color:#666;font-size:12px;">'
                        f'Minuta atual: {_esc(cont_min[:300])}</span>'
                    )
                consol_itens.append(_erro(f'<strong>{_esc(status)}:</strong> {linha}{detalhes}'))
            else:
                consol_itens.append(_atencao(f'<strong>{_esc(status)}:</strong> {linha}'))
        secoes.append(_secao(
            'Comparativo da Consolidação Contratual — Última Alteração × Minuta Atual',
            consol_itens,
        ))
    else:
        secoes.append(_secao(
            'Comparativo da Consolidação Contratual — Última Alteração × Minuta Atual',
            [_info('Comparativo não realizado — última alteração não anexada ou não aplicável.')],
        ))

    # ── SEÇÕES 9/10: POR ESTABELECIMENTO ──────────────────────────────────────
    por_estab = dados.get('por_estabelecimento', [])
    matrizes = [e for e in por_estab if isinstance(e, dict) and e.get('tipo') == 'matriz']
    filiais = [e for e in por_estab if isinstance(e, dict) and e.get('tipo') == 'filial']

    if matrizes:
        for ed in matrizes:
            titulo_e = ed.get('titulo', 'Matriz')
            itens_e = ed.get('itens', [])
            html_itens = [_item_html(i) for i in itens_e if isinstance(i, dict)]
            if not html_itens:
                html_itens = [_info('Nenhum item de conferência para este estabelecimento.')]
            secoes.append(_secao(f'Conferência da Matriz — {titulo_e}', html_itens))
    elif any(e.get('tipo') == 'matriz' for e in estabelecimentos):
        secoes.append(_secao('Conferência da Matriz', [
            _info('Matriz presente mas sem itens de conferência retornados pela IA.')
        ]))

    if filiais:
        for ed in filiais:
            titulo_e = ed.get('titulo', 'Filial')
            itens_e = ed.get('itens', [])
            html_itens = [_item_html(i) for i in itens_e if isinstance(i, dict)]
            if not html_itens:
                html_itens = [_info('Nenhum item de conferência para este estabelecimento.')]
            secoes.append(_secao(f'Conferência de Filial — {titulo_e}', html_itens))
    elif any(e.get('tipo') == 'filial' for e in estabelecimentos):
        secoes.append(_secao('Conferência das Filiais', [
            _info('Filiais presentes mas sem itens de conferência retornados pela IA.')
        ]))

    # ── SEÇÃO 11: COMPARATIVO CNPJ/NIRE ────────────────────────────────────────
    cnpj_nire = dados.get('comparativo_cnpj_nire', [])
    secoes.append(_secao('Comparativo de CNPJ/NIRE', [
        _item_html(i) for i in cnpj_nire if isinstance(i, dict)
    ] or [_info('Nenhum comparativo de CNPJ/NIRE retornado.')]))

    # ── SEÇÃO 12: COMPARATIVO DE ENDEREÇOS ─────────────────────────────────────
    enderecos = dados.get('comparativo_enderecos', [])
    secoes.append(_secao('Comparativo de Endereços', [
        _item_html(i) for i in enderecos if isinstance(i, dict)
    ] or [_info('Nenhum comparativo de endereço retornado.')]))

    # ── SEÇÃO 13: COMPARATIVO DE ATIVIDADES/CNAEs ──────────────────────────────
    atividades = dados.get('comparativo_atividades', [])
    secoes.append(_secao('Comparativo de Atividades Econômicas / CNAEs', [
        _item_html(i) for i in atividades if isinstance(i, dict)
    ] or [_info('Nenhum comparativo de atividades retornado.')]))

    # ── SEÇÃO 14: QUADRO SOCIETÁRIO ────────────────────────────────────────────
    qs = dados.get('quadro_societario', [])
    if qs:
        qs_itens = []
        linhas_qs = []
        for s in qs:
            if not isinstance(s, dict):
                continue
            nome = s.get('nome', '—')
            cpf = s.get('cpf_cnpj', '—')
            mov = s.get('movimento', '—')
            conf = s.get('confianca', '')
            duvida = s.get('duvida_extracao', False)

            conf_badge = (
                f' <span style="font-size:10px;color:#92400e;">[{_esc(conf)}]</span>'
                if conf and conf != 'Alta' else ''
            )
            duvida_badge = (
                ' <span style="font-size:10px;color:#dc2626;">[dúvida na extração]</span>'
                if duvida else ''
            )

            if mov == 'Remanescente':
                status_str = '<span class="status-ok">Remanescente</span>'
            elif mov == 'Ingressante':
                status_str = '<span class="status-ok">Ingressante</span>'
            elif mov == 'Retirante':
                status_str = '<span class="status-na">Retirante</span>'
            else:
                status_str = f'<span style="color:#92400e;">{_esc(mov)}</span>'

            linhas_qs.append(
                f'<tr>'
                f'<td>{_esc(nome)}{conf_badge}{duvida_badge}</td>'
                f'<td style="font-family:monospace">{_esc(cpf)}</td>'
                f'<td>{_esc(s.get("condicao_anterior") or "—")}</td>'
                f'<td>{_esc(s.get("condicao_minuta") or "—")}</td>'
                f'<td>{status_str}</td>'
                f'</tr>'
            )

        if linhas_qs:
            qs_itens.append(
                '<table class="tabela-cruzamento">'
                '<thead><tr>'
                '<th>Sócio</th><th>CPF/CNPJ</th>'
                '<th>Condição anterior</th><th>Condição na minuta</th><th>Status</th>'
                '</tr></thead>'
                '<tbody>' + ''.join(linhas_qs) + '</tbody></table>'
            )

        remanescentes = [s for s in qs if isinstance(s, dict) and s.get('movimento') == 'Remanescente']
        ingressantes = [s for s in qs if isinstance(s, dict) and s.get('movimento') == 'Ingressante']
        retirantes = [s for s in qs if isinstance(s, dict) and s.get('movimento') == 'Retirante']
        nao_id = [s for s in qs if isinstance(s, dict) and s.get('movimento') == 'Não identificado']

        if remanescentes:
            qs_itens.append(_ok(f'{len(remanescentes)} sócio(s) remanescente(s).'))
        if ingressantes:
            qs_itens.append(_info(
                f'Ingressante(s): {", ".join(_esc(s.get("nome","—")) for s in ingressantes)}'
            ))
        if retirantes:
            qs_itens.append(_info(
                f'Retirante(s): {", ".join(_esc(s.get("nome","—")) for s in retirantes)}'
            ))
        for s in nao_id:
            qs_itens.append(_atencao(
                f'<strong>Possível falha de extração:</strong> '
                f'{_esc(s.get("cpf_cnpj", "CPF/CNPJ não informado"))} — '
                f'{_esc(s.get("observacao") or "nome do sócio não identificado com segurança. Conferir manualmente.")}'
            ))

        secoes.append(_secao('Comparativo — Quadro Societário', qs_itens))
    else:
        secoes.append(_secao('Comparativo — Quadro Societário', [
            _info('Quadro societário não identificado automaticamente nos documentos.')
        ]))

    # ── SEÇÃO 15: REVISÃO TEXTUAL ───────────────────────────────────────────────
    rev_itens = [_item_html(i) for i in dados.get('revisao_textual', []) if isinstance(i, dict)]

    ceps = dados.get('ceps', [])
    if ceps:
        linhas_cep = (
            '<table style="border-collapse:collapse;width:100%;font-size:13px;margin-top:6px;">'
            '<thead><tr style="background:#fef3c7;">'
            '<th style="padding:4px 8px;text-align:left;border:1px solid #d97706;">CEP encontrado</th>'
            '<th style="padding:4px 8px;text-align:left;border:1px solid #d97706;">Problema</th>'
            '<th style="padding:4px 8px;text-align:center;border:1px solid #d97706;">Ocorrências</th>'
            '<th style="padding:4px 8px;text-align:left;border:1px solid #d97706;">Formato esperado</th>'
            '</tr></thead><tbody>'
        )
        for ci in ceps:
            if not isinstance(ci, dict):
                continue
            linhas_cep += (
                f'<tr>'
                f'<td style="padding:4px 8px;border:1px solid #fcd34d;">{_esc(ci.get("cep_encontrado",""))}</td>'
                f'<td style="padding:4px 8px;border:1px solid #fcd34d;">{_esc(ci.get("problema",""))}</td>'
                f'<td style="padding:4px 8px;border:1px solid #fcd34d;text-align:center;">{ci.get("ocorrencias",1)}</td>'
                f'<td style="padding:4px 8px;border:1px solid #fcd34d;">{_esc(ci.get("formato_esperado",""))}</td>'
                f'</tr>'
            )
        linhas_cep += '</tbody></table>'
        rev_itens.append(_atencao(f'{len(ceps)} CEP(s) com formatação irregular:{linhas_cep}'))

    if not rev_itens:
        rev_itens.append(_ok('Nenhuma inconsistência de formatação ou preenchimento identificada.'))
    secoes.append(_secao('Revisão Textual — Formatação e Preenchimento', rev_itens))

    # ── SEÇÃO 16: PENDÊNCIAS OBJETIVAS ─────────────────────────────────────────
    pendencias = dados.get('pendencias_objetivas', [])
    if pendencias:
        secoes.append(_secao(
            'Pendências Objetivas — Verificar Antes do Protocolo',
            [_erro(f'{i + 1}. {_esc(p)}') for i, p in enumerate(pendencias)],
        ))

    # ── SEÇÃO 17: ALERTAS PARA CONFERÊNCIA MANUAL ──────────────────────────────
    alertas = dados.get('alertas_manuais', [])
    if alertas:
        secoes.append(_secao(
            'Alertas para Conferência Manual',
            [_atencao(f'{i + 1}. {_esc(a)}') for i, a in enumerate(alertas)],
        ))

    # ── SEÇÃO 18: POSSÍVEIS FALHAS DE EXTRAÇÃO ──────────────────────────────────
    falhas = dados.get('falhas_extracao', [])
    if falhas:
        secoes.append(_secao(
            'Possíveis Falhas de Extração',
            [_info(f'{i + 1}. {_esc(f)}') for i, f in enumerate(falhas)],
        ))

    # ── SEÇÃO 19: DOCUMENTOS IGNORADOS ─────────────────────────────────────────
    ignorados_lista = dados.get('documentos_ignorados_lista', [])
    if ignorados_lista and not (len(ignorados_lista) == 1 and
                                 'nenhum' in ignorados_lista[0].lower()):
        secoes.append(_secao(
            'Documentos Ignorados pelo Usuário',
            [_atencao(f'{i + 1}. {_esc(ig)}') for i, ig in enumerate(ignorados_lista)],
        ))

    # ── SEÇÃO 20: CONCLUSÃO OPERACIONAL ────────────────────────────────────────
    conclusao = dados.get('conclusao') or {}
    n_pend = conclusao.get('pendencias_count', len(pendencias))
    n_alert = conclusao.get('alertas_count', len(alertas))
    n_falha = conclusao.get('falhas_count', len(falhas))
    n_ign = conclusao.get('ignorados_count', len(ignorados_lista))
    texto_concl = conclusao.get('texto', '')

    if not texto_concl:
        partes = []
        if n_pend:
            partes.append(f'<strong>{n_pend}</strong> pendência(s) objetiva(s)')
        if n_alert:
            partes.append(f'<strong>{n_alert}</strong> alerta(s) para conferência manual')
        if n_falha:
            partes.append(f'<strong>{n_falha}</strong> possível(is) falha(s) de extração')
        if n_ign:
            partes.append(f'<strong>{n_ign}</strong> documento(s) ignorado(s) pelo usuário')
        texto_concl = (
            'Foram identificados: ' + ', '.join(partes) + '.' if partes
            else 'Nenhuma divergência objetiva identificada. Proceda com a revisão antes do protocolo.'
        )
    else:
        texto_concl = _esc(texto_concl)

    concl_item = (
        _erro(texto_concl) if n_pend
        else _atencao(texto_concl) if (n_alert or n_falha)
        else _ok(texto_concl)
    )
    secoes.append(_secao('Conclusão Operacional', [concl_item]))

    # ── SEÇÃO 21: AVISO FINAL OBRIGATÓRIO ──────────────────────────────────────
    secoes.append(f'<div class="secao">{AVISO_JURIDICO_FINAL}</div>')

    return '\n'.join(secoes)


def _gerar_relatorio_com_ia(
    texto_minuta,
    texto_ultima,
    ultima_ignorada,
    docs_por_estab,
    estabelecimentos,
    tipo_processo,
    numero_alteracao,
    razao_social,
    cnpj_sociedade,
    nome_analista,
    data_hoje,
    log,
) -> str:
    """Gera o relatório HTML usando IA — chama Anthropic e converte JSON para HTML."""
    log['ia_chamada'] = True
    log['modelo_ia'] = MODELO_IA

    prompt_usuario = _montar_prompt_usuario_ia(
        texto_minuta=texto_minuta,
        texto_ultima=texto_ultima,
        ultima_ignorada=ultima_ignorada,
        docs_por_estab=docs_por_estab,
        estabelecimentos=estabelecimentos,
        tipo_processo=tipo_processo,
        numero_alteracao=numero_alteracao,
        razao_social=razao_social,
        cnpj_sociedade=cnpj_sociedade,
    )

    dados_raw = _chamar_ia_conferencia(prompt_usuario, log)
    dados = _validar_json_ia(dados_raw, log)

    html = _json_para_html(
        dados=dados,
        nome_analista=nome_analista,
        data_hoje=data_hoje,
        tipo_processo=tipo_processo,
        numero_alteracao=numero_alteracao,
        razao_social=razao_social,
        cnpj_sociedade=cnpj_sociedade,
        estabelecimentos=estabelecimentos,
        ultima_ignorada=ultima_ignorada,
    )
    return html


def _banner_fallback(motivo: str) -> str:
    """Banner de aviso quando a conferência com IA falha e usa fallback determinístico."""
    return (
        '<div class="secao" style="border-left:4px solid #f59e0b;background:#fffbeb;'
        'padding:16px;margin-bottom:16px;">'
        '<strong>Aviso — Conferência com IA indisponível</strong><br>'
        'A conferência com IA não pôde ser realizada neste momento. '
        'O relatório abaixo foi gerado pelo método determinístico (heurísticas/regex).<br>'
        f'<span style="font-size:12px;color:#92400e;">'
        f'Motivo técnico: {_esc(motivo[:300])}</span>'
        '</div>'
    )


# ─── Validações determinísticas auxiliares ────────────────────────────────────

def _parse_brl(s):
    """Converte string monetária BR para float. Ex: '1.000,00' → 1000.0"""
    try:
        return float(re.sub(r'[^\d,]', '', s).replace(',', '.'))
    except (ValueError, AttributeError, TypeError):
        return None


def _formatar_brl(v):
    """Formata float como string monetária BR. Ex: 1000.0 → 'R$ 1.000,00'"""
    s = f'{v:,.2f}'
    return 'R$ ' + s.replace(',', 'X').replace('.', ',').replace('X', '.')


def _validar_capital_social(texto_minuta):
    """Regra #10 — Valida aritmeticamente o capital social da minuta.
    Extrai o capital total declarado e a soma das quotas individuais por sócio.
    Retorna lista de strings para pendencias (se diverge) ou lista vazia.
    Comportamento conservador: na dúvida não acusa erro conclusivo."""
    if not texto_minuta:
        return []

    # Capital social total declarado — múltiplos padrões de fraseologia
    _PADROES_TOTAL = [
        (r'capital\s+social\s+(?:total\s+)?'
         r'(?:é\s+de|de|no\s+valor\s+de|importa\s+em|importando\s+em|'
         r'no\s+valor\s+total\s+de|corresponde\s+a|no\s+montante\s+de|'
         r'importa\s+na\s+import[aâ]ncia\s+de)\s+'
         r'R\$\s*([\d\.]+,\d{2})'),
        r'cujo\s+capital\s+social\s+(?:[éê]\s+de|de|no\s+valor\s+de)\s+R\$\s*([\d\.]+,\d{2})',
    ]
    m_total = None
    for _pat in _PADROES_TOTAL:
        m_total = re.search(_pat, texto_minuta, re.IGNORECASE)
        if m_total:
            break
    if not m_total:
        return []
    capital_total = _parse_brl(m_total.group(1))
    if not capital_total or capital_total <= 0:
        return []

    # Valores individuais de quotas — padrões estendidos
    pat_quota = re.compile(
        r'(?:no\s+valor\s+total\s+de|totalizando|equivalente\s+a|'
        r'no\s+montante\s+de|correspondente\s+a|importando\s+em|'
        r'integralizando|integraliza[çc][aã]o\s+de|'
        r'quotas?\s+(?:no\s+valor\s+(?:total\s+)?de|totalizando)|'
        r'no\s+valor\s+de)\s+'
        r'R\$\s*([\d\.]+,\d{2})',
        re.IGNORECASE,
    )
    valores = []
    for m in pat_quota.finditer(texto_minuta):
        v = _parse_brl(m.group(1))
        if v and 0 < v < capital_total:
            valores.append(v)

    if len(valores) < 2:
        return []  # Quotas individuais insuficientes para somar — comportamento conservador

    soma = round(sum(valores), 2)
    if abs(soma - capital_total) > 0.02:
        return [
            f'Capital social declarado: {_formatar_brl(capital_total)} — '
            f'soma das quotas identificadas ({len(valores)} sócio(s)): {_formatar_brl(soma)} — '
            f'diferença: {_formatar_brl(abs(soma - capital_total))}. '
            'Verificar se os valores de quotas estão corretos e conferem com o capital total.'
        ]
    return []


def _verificar_desimpedimento(texto_minuta):
    """Regra #13 — Verifica presença de cláusula de desimpedimento quando há cláusula
    de administração. Retorna lista de alertas (vazia se tudo OK)."""
    if not texto_minuta:
        return []

    tem_admin = bool(re.search(
        r'\b(?:administra[çc][aã]o|administrador|gerente\s+geral)\b',
        texto_minuta, re.IGNORECASE,
    ))
    if not tem_admin:
        return []

    tem_desimped = bool(re.search(
        r'\b(?:desimpedid[oa]|desimpedimento|livre\s+de\s+impedimento'
        r'|n[aã]o\s+(?:est[aá]|é|se\s+encontra)\s+impedid)',
        texto_minuta, re.IGNORECASE,
    ))
    if not tem_desimped:
        return [
            'Cláusula de administração identificada na minuta, mas declaração de '
            'desimpedimento não localizada. Verificar se a cláusula de desimpedimento '
            'está presente e adequada.'
        ]
    return []


def _extrair_capital_bruto(texto):
    """Extrai o valor de capital social de qualquer documento (DBE, FCPJ, FCN).
    Retorna (float|None, str) — valor e nota sobre a extração.
    Conservador: só retorna quando há correspondência razoavelmente inequívoca."""
    if not texto:
        return None, 'não localizado'
    padroes = [
        # Contratos e minutas
        (r'capital\s+social\s+(?:total\s+)?'
         r'(?:é\s+de|de|no\s+valor\s+de|importa\s+em|importando\s+em|'
         r'no\s+valor\s+total\s+de|corresponde\s+a|no\s+montante\s+de|'
         r'importa\s+na\s+import[aâ]ncia\s+de)\s+'
         r'R\$\s*([\d\.]+,\d{2})'),
        r'cujo\s+capital\s+social\s+(?:[éê]\s+de|de|no\s+valor\s+de)\s+R\$\s*([\d\.]+,\d{2})',
        # DBE/FCPJ costumam apresentar: "Capital Social: R$ X.XXX,XX" ou "Capital: R$ X"
        r'capital\s+social[:\s]+R\$\s*([\d\.]+,\d{2})',
        r'\bcapital[:\s]+R\$\s*([\d\.]+,\d{2})',
    ]
    for pat in padroes:
        m = re.search(pat, texto, re.IGNORECASE)
        if m:
            v = _parse_brl(m.group(1))
            if v and v > 0:
                return v, 'localizado'
    return None, 'não localizado'


def _verificar_retirante_administrador(socios_retirantes, texto_ultima, texto_minuta):
    """Regra #12 — Verifica se sócios retirantes eram administradores no ato anterior.

    socios_retirantes : lista de dicts com 'cpf', 'nome', 'confianca'
                        (somente os sócios cujo CPF está em retirantes)
    texto_ultima      : texto da última alteração consolidada (pode ser None)
    texto_minuta      : texto da minuta nova

    Retorna (qs_itens: list[str], pendencias: list[str]).
    Comportamento conservador: sem ato anterior → atenção manual; baixa confiança → alerta,
    não pendência conclusiva.
    """
    itens, pendencias = [], []

    if not socios_retirantes or not texto_minuta:
        return itens, pendencias

    # ── Extrair trecho de administração do ato ANTERIOR ──────────────────────
    # Usa findall para cabeçalhos de cláusula E busca DOTALL para capturar
    # CPF/nomes em parágrafos multilinhas da mesma cláusula.
    admin_ultima = ''
    if texto_ultima:
        trechos_ult = re.findall(
            r'(?i)(?:CL[AÁ]USULA|Cl[áa]usula)\s+\w+[^\n]{0,600}'
            r'(?:administra[çc][aã]o|administrador)[^\n]{0,600}',
            texto_ultima,
        )
        admin_ultima = ' '.join(trechos_ult)
        # Sempre complementar com contexto amplo (DOTALL) — captura CPF/nome
        # em parágrafos que vêm após o cabeçalho da cláusula
        m_adm = re.search(
            r'.{0,400}(?:administra[çc][aã]o|administrador).{0,400}',
            texto_ultima, re.IGNORECASE | re.DOTALL,
        )
        if m_adm:
            admin_ultima = (admin_ultima + ' ' + m_adm.group(0)).strip()

    # ── Extrair trecho de administração da MINUTA NOVA ────────────────────────
    admin_minuta = ''
    trechos_min = re.findall(
        r'(?i)(?:CL[AÁ]USULA|Cl[áa]usula)\s+\w+[^\n]{0,600}'
        r'(?:administra[çc][aã]o|administrador)[^\n]{0,600}',
        texto_minuta,
    )
    admin_minuta = ' '.join(trechos_min)
    # Sempre complementar com contexto amplo (DOTALL)
    m_adm2 = re.search(
        r'.{0,300}(?:administra[çc][aã]o|administrador).{0,300}',
        texto_minuta, re.IGNORECASE | re.DOTALL,
    )
    if m_adm2:
        admin_minuta = (admin_minuta + ' ' + m_adm2.group(0)).strip()

    # ── Padrão de destituição global na minuta nova ───────────────────────────
    tem_destituicao_global = bool(re.search(
        r'\b(?:destitui[çc][aã]o|destitui[dr]|exonerado|exonera[çc][aã]o|'
        r'retirada\s+da\s+administra[çc][aã]o|cessando\s+(?:a\s+)?administra[çc][aã]o|'
        r'n[aã]o\s+mais\s+administra|dispensado\s+da\s+administra[çc][aã]o)\b',
        texto_minuta, re.IGNORECASE,
    ))

    # ── Verificar cada sócio retirante ───────────────────────────────────────
    for s in socios_retirantes:
        cpf = s['cpf']       # 11 dígitos sem formatação
        nome = s['nome']
        confianca = s['confianca']

        # Sem ato anterior disponível → alerta manual
        if texto_ultima is None:
            if confianca in ('Alta', 'Média'):
                itens.append(_atencao(
                    f'Sócio retirante <strong>{_esc(nome)}</strong>: '
                    'ato anterior não disponível — verificar manualmente se havia '
                    'vínculo de administração que exija cláusula de destituição.'
                ))
            continue

        # ── Verificar se era administrador no ato ANTERIOR ───────────────────
        era_admin_ultima = False
        metodo = ''

        if admin_ultima:
            # Tentativa 1 — CPF no trecho de administração (mais confiável)
            cpf_re = rf'{cpf[:3]}[.\s]?{cpf[3:6]}[.\s]?{cpf[6:9]}[-/\s]?{cpf[9:]}'
            if re.search(cpf_re, admin_ultima):
                era_admin_ultima = True
                metodo = 'CPF na cláusula de administração do ato anterior'

            if not era_admin_ultima and confianca in ('Alta', 'Média'):
                # Tentativa 2 — nome normalizado: exige ≥ 2 tokens consecutivos
                nome_norm = _normalizar(nome)
                tokens = [t for t in nome_norm.split() if len(t) >= 3]
                admin_norm = _normalizar(admin_ultima)
                for i in range(len(tokens) - 1):
                    par = tokens[i] + r'\s+' + tokens[i + 1]
                    if re.search(par, admin_norm):
                        era_admin_ultima = True
                        metodo = (
                            f'nome na cláusula de administração do ato anterior '
                            f'(confiança: {confianca})'
                        )
                        break

        if not era_admin_ultima:
            # Não foi possível confirmar vínculo de administração — sem alerta
            continue

        # ── Era admin: verificar destituição e administração atualizada ───────
        # Refinamento: buscar destituição próxima ao nome na minuta
        destituido = tem_destituicao_global
        if not destituido:
            nome_norm = _normalizar(nome)
            tokens_nome = [t for t in nome_norm.split() if len(t) >= 3]
            if tokens_nome:
                jan = re.search(
                    r'.{0,300}' + re.escape(tokens_nome[0]) + r'.{0,300}',
                    _normalizar(texto_minuta),
                )
                if jan:
                    destituido = bool(re.search(
                        r'destitui|exonera|retirada\s+da\s+admin|cessando',
                        jan.group(0), re.IGNORECASE,
                    ))

        tem_admin_nova = bool(admin_minuta)

        if not destituido and not tem_admin_nova:
            itens.append(_atencao(
                f'Sócio retirante <strong>{_esc(nome)}</strong> '
                f'identificado como administrador no ato anterior ({metodo}). '
                'Não foram localizados na minuta: termo de destituição da administração '
                'e cláusula de administração atualizada. Verificar manualmente.'
            ))
            pendencias.append(
                f'Sócio retirante "{nome}": era administrador no ato anterior — '
                'verificar destituição e cláusula de administração atualizada na minuta.'
            )
        elif not destituido:
            itens.append(_atencao(
                f'Sócio retirante <strong>{_esc(nome)}</strong> '
                f'identificado como administrador no ato anterior ({metodo}). '
                'Cláusula de administração atualizada localizada na minuta, mas '
                'termo de destituição não identificado explicitamente. '
                'Verificar se há cláusula de destituição ou retirada da administração.'
            ))
            pendencias.append(
                f'Sócio retirante "{nome}": era administrador — '
                'verificar destituição explícita na minuta.'
            )
        else:
            itens.append(_info(
                f'Sócio retirante <strong>{_esc(nome)}</strong>: '
                f'era administrador no ato anterior ({metodo}) — '
                'término de administração identificado na minuta.'
            ))

    return itens, pendencias


# ─── Gerador do relatório comparativo (sem IA) ────────────────────────────────

def _gerar_relatorio_html_comparativo(
    texto_minuta,
    texto_ultima,
    ultima_ignorada,
    docs_por_estab,
    estabelecimentos,
    tipo_processo,
    numero_alteracao,
    razao_social,
    cnpj_sociedade,
    nome_analista,
    data_hoje,
    log,
):
    """
    Gera o relatório comparativo em HTML puro — SEM chamadas à IA/Claude.

    docs_por_estab: lista paralela a estabelecimentos, cada item:
        {'docs': {'viabilidade': texto|None, ...}, 'ignorados': {'viabilidade': bool, ...}}
    """
    log['ia_chamada'] = False
    log['conferencia_sem_ia'] = True
    log['USE_AI_FOR_MAIN_CONFERENCE'] = USE_AI_FOR_MAIN_CONFERENCE

    print(
        f'[conferencia] RELATÓRIO COMPARATIVO SEM IA | '
        f'tipo={tipo_processo} | estabs={len(estabelecimentos)} | '
        f'USE_AI={USE_AI_FOR_MAIN_CONFERENCE}'
    )

    secoes = []
    pendencias = []
    alertas_manuais = []
    falhas_extracao = []

    tipo_label = TIPO_PROCESSO_LABEL.get(tipo_processo, tipo_processo or 'Não identificado')

    # Numeração: usar função unificada aplicada à minuta para consistência total
    num_minuta_unif, conf_num_minuta = (
        _extrair_numero_alteracao_unificado(texto_minuta) if texto_minuta
        else (numero_alteracao, 'Média' if numero_alteracao else 'Não localizado')
    )
    # Se o frontend enviou número (da análise prévia) mas a função unificada não achou,
    # respeitar o do frontend quando a confiança for pelo menos Média
    if num_minuta_unif is None and numero_alteracao:
        num_minuta_unif = numero_alteracao
        conf_num_minuta = 'Média'

    if num_minuta_unif:
        if conf_num_minuta == 'Baixa':
            num_str = f'Possível {num_minuta_unif}ª Alteração Contratual — confirmar manualmente'
        else:
            num_str = f'{num_minuta_unif}ª Alteração Contratual'
    else:
        num_str = 'Número não identificado — conferir manualmente'

    # Atualizar número para uso interno consistente
    numero_alteracao = num_minuta_unif

    # ── 1. CABEÇALHO DO PROCESSO ──────────────────────────────────────────────
    cab_itens = [
        _info(f'<strong>Analista:</strong> {_esc(nome_analista)}'),
        _info(f'<strong>Data da Conferência:</strong> {_esc(data_hoje)}'),
        _info(f'<strong>Tipo de Processo:</strong> {_esc(tipo_label)}'),
        _info(f'<strong>Numeração:</strong> {_esc(num_str)}'),
    ]
    if razao_social:
        cab_itens.append(_info(f'<strong>Razão Social (minuta):</strong> {_esc(razao_social)}'))
    if cnpj_sociedade:
        cab_itens.append(_info(f'<strong>CNPJ da Sociedade:</strong> {_esc(cnpj_sociedade)}'))
    secoes.append(_secao('Resumo do Processo', cab_itens))

    # ── 2. STATUS DOS DOCUMENTOS ──────────────────────────────────────────────
    doc_status = []
    if texto_minuta:
        doc_status.append(_ok(
            f'{_tag_doc("Minuta do Contrato Social")} — Recebida '
            f'({len(texto_minuta)} caracteres extraídos)'
        ))
    else:
        doc_status.append(_erro(
            f'{_tag_doc("Minuta do Contrato Social")} — AUSENTE — FALHA CRÍTICA'
        ))
        pendencias.append('Minuta do contrato social não foi anexada.')

    if ultima_ignorada:
        doc_status.append(_atencao(
            f'{_tag_doc("Última Alteração Contratual")} — Ignorada pelo usuário — '
            'comparativo de quadro societário e sequência não realizado'
        ))
    elif texto_ultima:
        doc_status.append(_ok(
            f'{_tag_doc("Última Alteração Contratual")} — Recebida '
            f'({len(texto_ultima)} caracteres extraídos)'
        ))
    else:
        if tipo_processo == 'abertura':
            doc_status.append(_info(
                f'{_tag_doc("Última Alteração Contratual")} — Não aplicável para processo de abertura'
            ))
        else:
            doc_status.append(_atencao(
                f'{_tag_doc("Última Alteração Contratual")} — Não anexada — '
                'comparativo de quadro societário e sequência não realizado'
            ))

    for i, (estab, edoc) in enumerate(zip(estabelecimentos, docs_por_estab)):
        tipo_e = estab.get('tipo', 'matriz')
        cnpj_e = estab.get('cnpj') or ''
        label_e = 'Matriz' if tipo_e == 'matriz' else f'Filial {i}'
        if cnpj_e:
            label_e += f' ({cnpj_e})'
        for tipo_doc in TIPOS_DOC:
            label_doc = LABELS_DOC[tipo_doc]
            ignorado = edoc['ignorados'].get(tipo_doc, False)
            texto_doc = edoc['docs'].get(tipo_doc)
            if ignorado:
                doc_status.append(_atencao(
                    f'{_tag_doc(label_doc)} — {_esc(label_e)} — Ignorado pelo usuário'
                ))
            elif texto_doc:
                doc_status.append(_ok(
                    f'{_tag_doc(label_doc)} — {_esc(label_e)} — Recebido'
                ))
            else:
                doc_status.append(_info(
                    f'{_tag_doc(label_doc)} — {_esc(label_e)} — Ausente / não anexado'
                ))

    secoes.append(_secao('Documentos — Status', doc_status))

    # ── 3. CONFERÊNCIA GERAL DA MINUTA ────────────────────────────────────────
    geral_itens = []

    if texto_minuta and texto_ultima and not ultima_ignorada:
        # Sequência numérica
        num_ultima_doc = _extrair_numero_alteracao(texto_ultima)
        num_minuta_doc = _extrair_numero_alteracao(texto_minuta)

        if num_ultima_doc is not None and num_minuta_doc is not None:
            esperado = num_ultima_doc + 1
            if num_minuta_doc == esperado:
                geral_itens.append(_ok(
                    f'Sequência numérica correta: última alteração é a {num_ultima_doc}ª, '
                    f'a minuta é a {num_minuta_doc}ª (N+1).'
                ))
            else:
                geral_itens.append(_erro(
                    f'Sequência numérica inconsistente: a última alteração é a {num_ultima_doc}ª, '
                    f'porém a minuta indica ser a {num_minuta_doc}ª (esperado: {esperado}ª).'
                ))
                pendencias.append(
                    f'Verificar numeração da alteração: esperada {esperado}ª, '
                    f'encontrada {num_minuta_doc}ª na minuta.'
                )
        elif num_ultima_doc is not None:
            geral_itens.append(_atencao(
                f'Número da última alteração identificado: {num_ultima_doc}ª. '
                'Número da minuta não identificado — verificar manualmente.'
            ))
        elif num_minuta_doc is not None:
            geral_itens.append(_atencao(
                f'Número da minuta identificado: {num_minuta_doc}ª. '
                'Número da última alteração não identificado — verificar manualmente.'
            ))
        else:
            geral_itens.append(_atencao(
                'Números de alteração não identificados automaticamente nos documentos — '
                'verificar sequência manualmente.'
            ))

        # CNPJ consistente entre os documentos
        cnpjs_minuta = _extrair_cnpjs(texto_minuta)
        cnpjs_ultima = _extrair_cnpjs(texto_ultima)
        if cnpjs_minuta and cnpjs_ultima:
            if _cnpjs_iguais(cnpjs_minuta[0], cnpjs_ultima[0]):
                geral_itens.append(_ok(
                    f'CNPJ da sociedade consistente entre os documentos: {_esc(cnpjs_minuta[0])}'
                ))
            else:
                geral_itens.append(_erro(
                    f'CNPJ divergente: minuta indica {_esc(cnpjs_minuta[0])}, '
                    f'última alteração indica {_esc(cnpjs_ultima[0])}.'
                ))
                pendencias.append('CNPJ divergente entre minuta e última alteração contratual.')

    elif texto_minuta and tipo_processo != 'abertura' and not ultima_ignorada:
        geral_itens.append(_atencao(
            'Última alteração não enviada — verificação de sequência numérica e CNPJ não realizada.'
        ))

    # Local de assinatura vs. foro
    if texto_minuta:
        foro_m = re.search(
            r'foro\s+(?:da\s+|de\s+|do\s+)?([A-ZÁÉÍÓÚÂÊÔÃÕa-záéíóúâêôãõ\s]{3,40})',
            texto_minuta, re.IGNORECASE,
        )
        cidade_data_m = re.search(
            r'([A-ZÁÉÍÓÚÂÊÔÃÕa-záéíóúâêôãõ]{4,40}),\s*\d{1,2}\s+de\s+[a-záéíóúâêôãõ]+\s+de\s+\d{4}',
            texto_minuta,
        )
        if foro_m and cidade_data_m:
            foro_n = _normalizar(foro_m.group(1))
            cidade_n = _normalizar(cidade_data_m.group(1))
            if foro_n[:8] in cidade_n or cidade_n[:8] in foro_n:
                geral_itens.append(_ok(
                    f'Local de assinatura e foro compatíveis: '
                    f'"{_esc(foro_m.group(1).strip())}"'
                ))
            else:
                geral_itens.append(_atencao(
                    f'Verificar local de assinatura e foro: foro indica '
                    f'"{_esc(foro_m.group(1).strip())}", cidade da data de assinatura indica '
                    f'"{_esc(cidade_data_m.group(1).strip())}".'
                ))

    if not geral_itens:
        geral_itens.append(_info(
            'Conferência geral não realizada — minuta não anexada ou sem última alteração.'
        ))

    secoes.append(_secao('Conferência da Minuta e Última Alteração', geral_itens))

    # ── 4. POR ESTABELECIMENTO ────────────────────────────────────────────────
    for i, (estab, edoc) in enumerate(zip(estabelecimentos, docs_por_estab)):
        tipo_e = estab.get('tipo', 'matriz')
        cnpj_e = estab.get('cnpj') or ''
        nire_e = estab.get('nire') or ''
        descricao_e = estab.get('descricao', '')

        if tipo_e == 'matriz':
            titulo_e = 'Matriz' + (f' — CNPJ: {cnpj_e}' if cnpj_e else '')
        else:
            # Evitar duplicação: se a descrição já começa com "Filial N", usar só ela
            if descricao_e and descricao_e.strip():
                desc_strip = descricao_e.strip()
                if re.match(r'^[Ff]ilial\s+\d+', desc_strip):
                    titulo_e = desc_strip
                else:
                    titulo_e = f'Filial {i} — {desc_strip}'
            else:
                titulo_e = f'Filial {i}'
            if cnpj_e:
                titulo_e += f' (CNPJ: {cnpj_e})'

        estab_itens = []
        docs_e = edoc['docs']
        ignorados_e = edoc['ignorados']

        # Alerta viabilidade vencida
        if docs_e.get('viabilidade'):
            data_viab = _extrair_data_viabilidade(docs_e['viabilidade'])
            if data_viab:
                hoje = date.today()
                dias = (hoje - data_viab).days
                if dias > 20:
                    estab_itens.append(_erro(
                        f'ATENÇÃO — A Consulta Prévia/Viabilidade foi protocolada em '
                        f'{data_viab.strftime("%d/%m/%Y")} ({dias} dias antes desta conferência). '
                        'Verifique se o prazo de validade ainda está vigente antes do protocolo.'
                    ))
                    pendencias.append(
                        f'Viabilidade ({titulo_e}): prazo de {dias} dias — verificar validade.'
                    )

        # CNPJ do bloco na minuta
        if cnpj_e and texto_minuta:
            cnpjs_minuta_all = _extrair_cnpjs(texto_minuta)
            if any(_cnpjs_iguais(cnpj_e, c) for c in cnpjs_minuta_all):
                estab_itens.append(_ok(
                    f'CNPJ {_esc(cnpj_e)} encontrado na minuta.'
                ))
            else:
                estab_itens.append(_atencao(
                    f'CNPJ {_esc(cnpj_e)} (indicado para este bloco) '
                    'não localizado na minuta — verificar manualmente.'
                ))

        # Viabilidade, FCPJ, FCN
        for tipo_doc in ['viabilidade', 'fcpj', 'dbe', 'fcn']:
            texto_doc = docs_e.get(tipo_doc)
            ignorado = ignorados_e.get(tipo_doc, False)
            label_doc = LABELS_DOC[tipo_doc]

            if ignorado:
                estab_itens.append(_atencao(f'{_tag_doc(label_doc)} — Ignorado pelo usuário.'))
                continue

            if not texto_doc:
                estab_itens.append(_info(f'{_tag_doc(label_doc)} — Não anexado.'))
                continue

            # CNPJ no documento — classificar cada CNPJ encontrado
            cnpjs_doc = _extrair_cnpjs(texto_doc)
            # Lista de todos os CNPJs dos estabelecimentos do processo
            cnpjs_todos_estabs = [
                e.get('cnpj', '') for e in estabelecimentos if e.get('cnpj')
            ]
            if not cnpjs_doc:
                estab_itens.append(_atencao(
                    f'{_tag_doc(label_doc)} — <strong>Não localizado</strong>: '
                    'nenhum CNPJ encontrado no documento.'
                ))
            else:
                cnpj_esperado_encontrado = False
                for cnpj_d in cnpjs_doc:
                    if cnpj_e and _cnpjs_iguais(cnpj_e, cnpj_d):
                        estab_itens.append(_ok(
                            f'{_tag_doc(label_doc)} — <strong>Conforme</strong>: '
                            f'CNPJ esperado {_esc(cnpj_d)} localizado.'
                        ))
                        cnpj_esperado_encontrado = True
                    elif any(_cnpjs_iguais(cnpj_d, c) for c in cnpjs_todos_estabs if c != cnpj_e):
                        estab_itens.append(_atencao(
                            f'{_tag_doc(label_doc)} — <strong>Atenção</strong>: '
                            f'CNPJ {_esc(cnpj_d)} pertence a outro estabelecimento do processo. '
                            'Verificar se o documento correto foi anexado.'
                        ))
                        alertas_manuais.append(
                            f'{label_doc} ({titulo_e}): CNPJ {cnpj_d} pertence a outro '
                            'estabelecimento do processo.'
                        )
                    else:
                        estab_itens.append(_atencao(
                            f'{_tag_doc(label_doc)} — <strong>CNPJ adicional</strong>: '
                            f'{_esc(cnpj_d)} localizado sem vínculo automático com este estabelecimento.'
                        ))
                        alertas_manuais.append(
                            f'{label_doc} ({titulo_e}): CNPJ adicional {cnpj_d} sem vínculo '
                            'automático identificado.'
                        )
                if cnpj_e and not cnpj_esperado_encontrado:
                    estab_itens.append(_erro(
                        f'{_tag_doc(label_doc)} — <strong>Divergente</strong>: '
                        f'CNPJ esperado {_esc(cnpj_e)} não encontrado no documento.'
                    ))
                    pendencias.append(
                        f'{label_doc} ({titulo_e}): CNPJ esperado {cnpj_e} não localizado.'
                    )

            # NIRE
            if nire_e:
                nires_doc = _extrair_nires(texto_doc)
                nire_e_norm = re.sub(r'[^\d]', '', nire_e)
                if any(n == nire_e_norm for n in nires_doc):
                    estab_itens.append(_ok(
                        f'{_tag_doc(label_doc)} — NIRE confere: {_esc(nire_e)}'
                    ))
                elif nires_doc:
                    estab_itens.append(_atencao(
                        f'{_tag_doc(label_doc)} — NIRE divergente: documento contém '
                        f'{_esc(nires_doc[0])}, esperado {_esc(nire_e)}.'
                    ))
                    pendencias.append(f'{label_doc} ({titulo_e}): NIRE divergente.')
                else:
                    estab_itens.append(_info(
                        f'{_tag_doc(label_doc)} — NIRE não localizado no documento.'
                    ))

            # CNAEs (viabilidade e FCN)
            if tipo_doc in ['viabilidade', 'fcn'] and texto_minuta:
                cnaes_doc = _extrair_cnaes(texto_doc)
                cnaes_minuta = _extrair_cnaes(texto_minuta)
                if cnaes_doc and cnaes_minuta:
                    ausentes_minuta = [c for c in cnaes_doc if c not in cnaes_minuta]
                    ausentes_doc = [c for c in cnaes_minuta if c not in cnaes_doc]
                    if not ausentes_minuta and not ausentes_doc:
                        estab_itens.append(_ok(
                            f'{_tag_doc(label_doc)} — CNAEs conferem com a minuta: '
                            f'{", ".join(cnaes_doc)}'
                        ))
                    else:
                        partes = []
                        if ausentes_minuta:
                            partes.append(
                                f'CNAEs no documento mas não na minuta: {", ".join(ausentes_minuta)}'
                            )
                        if ausentes_doc:
                            partes.append(
                                f'CNAEs na minuta mas não no documento: {", ".join(ausentes_doc)}'
                            )
                        estab_itens.append(_atencao(
                            f'{_tag_doc(label_doc)} — Divergência de CNAEs: '
                            + '; '.join(partes)
                        ))
                elif cnaes_doc:
                    estab_itens.append(_info(
                        f'{_tag_doc(label_doc)} — CNAEs encontrados: {", ".join(cnaes_doc)}'
                    ))

        # Uso do Solo
        if ignorados_e.get('uso_do_solo'):
            estab_itens.append(_atencao(f'{_tag_doc("Uso do Solo")} — Ignorado pelo usuário.'))
        elif docs_e.get('uso_do_solo'):
            if texto_minuta:
                end_minuta = _extrair_endereco_contexto(texto_minuta, cnpj_e)
                end_doc = docs_e['uso_do_solo']
                if end_minuta and end_doc:
                    tokens = [t for t in _normalizar(end_minuta).split() if len(t) > 3]
                    matches = sum(1 for t in tokens if t in _normalizar(end_doc))
                    pct = matches / len(tokens) if tokens else 0
                    if pct >= 0.5:
                        estab_itens.append(_ok(
                            f'{_tag_doc("Uso do Solo")} — Endereço compatível com a minuta.'
                        ))
                    else:
                        estab_itens.append(_atencao(
                            f'{_tag_doc("Uso do Solo")} — Endereço pode não corresponder '
                            'ao indicado na minuta. Verificar manualmente.'
                        ))
                else:
                    estab_itens.append(_info(
                        f'{_tag_doc("Uso do Solo")} — Recebido. '
                        'Verificar endereço e atividade manualmente.'
                    ))
            else:
                estab_itens.append(_info(f'{_tag_doc("Uso do Solo")} — Recebido.'))
        else:
            estab_itens.append(_info(f'{_tag_doc("Uso do Solo")} — Não anexado.'))

        # Número Predial
        if ignorados_e.get('numero_predial'):
            estab_itens.append(_atencao(
                f'{_tag_doc("Número Predial")} — Ignorado pelo usuário.'
            ))
        elif docs_e.get('numero_predial'):
            if texto_minuta:
                num_pred_min = re.search(
                    r'(?:n[º°]|número|nº)\s*(\d+[A-Za-z]?)',
                    texto_minuta, re.IGNORECASE,
                )
                num_pred_doc = re.search(
                    r'(?:n[º°]|número|nº)\s*(\d+[A-Za-z]?)',
                    docs_e['numero_predial'], re.IGNORECASE,
                )
                if num_pred_min and num_pred_doc:
                    if num_pred_min.group(1).upper() == num_pred_doc.group(1).upper():
                        estab_itens.append(_ok(
                            f'{_tag_doc("Número Predial")} — Número confere: '
                            f'{_esc(num_pred_min.group(1))}'
                        ))
                    else:
                        estab_itens.append(_erro(
                            f'{_tag_doc("Número Predial")} — Número divergente: '
                            f'minuta indica {_esc(num_pred_min.group(1))}, '
                            f'documento indica {_esc(num_pred_doc.group(1))}.'
                        ))
                        pendencias.append(
                            f'Número predial divergente para {titulo_e}.'
                        )
                else:
                    estab_itens.append(_info(
                        f'{_tag_doc("Número Predial")} — Recebido. '
                        'Verificar número e logradouro manualmente.'
                    ))
            else:
                estab_itens.append(_info(f'{_tag_doc("Número Predial")} — Recebido.'))
        else:
            estab_itens.append(_info(f'{_tag_doc("Número Predial")} — Não anexado.'))

        # Filial sem vínculo identificado
        if tipo_e == 'filial' and not cnpj_e and not nire_e:
            estab_itens.append(_atencao(
                'Não foi possível vincular automaticamente esta filial a uma cláusula '
                'específica da minuta. Conferir manualmente.'
            ))

        secoes.append(_secao(f'Por Estabelecimento — {titulo_e}', estab_itens))

    # ── 5. COMPARATIVO QUADRO SOCIETÁRIO ──────────────────────────────────────
    if texto_minuta and texto_ultima and not ultima_ignorada:
        qs_itens = []
        socios_ultima_raw = _extrair_socios_v2(texto_ultima)
        socios_minuta_raw = _extrair_socios_v2(texto_minuta)

        # Separar válidos de falhas de extração
        socios_ultima = [s for s in socios_ultima_raw if s['valido']]
        socios_minuta = [s for s in socios_minuta_raw if s['valido']]
        falhas_ult = [s for s in socios_ultima_raw if not s['valido']]
        falhas_min = [s for s in socios_minuta_raw if not s['valido']]

        # Registrar falhas de extração
        for s in falhas_ult + falhas_min:
            cpf_fmt = f'{s["cpf"][:3]}.{s["cpf"][3:6]}.{s["cpf"][6:9]}-{s["cpf"][9:]}'
            msg = (
                f'CPF {_esc(cpf_fmt)} localizado, mas nome do sócio não identificado '
                f'com segurança (confiança: {s["confianca"]}). Conferir manualmente.'
            )
            qs_itens.append(_atencao(f'<strong>Possível falha de extração:</strong> {msg}'))
            falhas_extracao.append(f'Quadro Societário: {msg}')

        if socios_ultima or socios_minuta:
            cpfs_ultima = {s['cpf'] for s in socios_ultima}
            cpfs_minuta = {s['cpf'] for s in socios_minuta}
            remanescentes = cpfs_ultima & cpfs_minuta
            ingressantes = cpfs_minuta - cpfs_ultima
            retirantes = cpfs_ultima - cpfs_minuta

            linhas = []
            for s in socios_ultima:
                conf_badge = (
                    '' if s['confianca'] == 'Alta'
                    else f' <span style="font-size:10px;color:#92400e;">[{s["confianca"]}]</span>'
                )
                if s['cpf'] in remanescentes:
                    status_str = '<span class="status-ok">Remanescente</span>'
                else:
                    status_str = '<span class="status-na">Retirante</span>'
                linhas.append(
                    f'<tr><td>{_esc(s["nome"])}{conf_badge}</td>'
                    f'<td style="font-family:monospace">'
                    f'{s["cpf"][:3]}.{s["cpf"][3:6]}.{s["cpf"][6:9]}-{s["cpf"][9:]}'
                    f'</td>'
                    f'<td>Sócio anterior</td><td>{status_str}</td></tr>'
                )
            for s in socios_minuta:
                if s['cpf'] in ingressantes:
                    conf_badge = (
                        '' if s['confianca'] == 'Alta'
                        else f' <span style="font-size:10px;color:#92400e;">[{s["confianca"]}]</span>'
                    )
                    linhas.append(
                        f'<tr><td>{_esc(s["nome"])}{conf_badge}</td>'
                        f'<td style="font-family:monospace">'
                        f'{s["cpf"][:3]}.{s["cpf"][3:6]}.{s["cpf"][6:9]}-{s["cpf"][9:]}'
                        f'</td>'
                        f'<td>—</td>'
                        f'<td><span class="status-ok">Ingressante</span></td></tr>'
                    )

            if linhas:
                tabela = (
                    '<table class="tabela-cruzamento">'
                    '<thead><tr>'
                    '<th>Sócio</th><th>CPF</th>'
                    '<th>Situação anterior</th><th>Status</th>'
                    '</tr></thead>'
                    '<tbody>' + ''.join(linhas) + '</tbody></table>'
                )
                qs_itens.append(tabela)

            if remanescentes:
                qs_itens.append(_ok(
                    f'{len(remanescentes)} sócio(s) remanescente(s) identificado(s).'
                ))
            if ingressantes:
                nomes_ing = [s['nome'] for s in socios_minuta if s['cpf'] in ingressantes]
                qs_itens.append(_info(
                    f'Ingressante(s): {", ".join(_esc(n) for n in nomes_ing)}'
                ))
            if retirantes:
                nomes_ret = [s['nome'] for s in socios_ultima if s['cpf'] in retirantes]
                qs_itens.append(_info(
                    f'Retirante(s): {", ".join(_esc(n) for n in nomes_ret)}'
                ))
                # Verificar se retirante ainda aparece na consolidação da minuta
                # (apenas para nomes com confiança Alta ou Média)
                for s in socios_ultima:
                    if s['cpf'] not in retirantes:
                        continue
                    if s['confianca'] not in ('Alta', 'Média'):
                        continue
                    nome = s['nome']
                    if _normalizar(nome[:12]) in _normalizar(texto_minuta):
                        qs_itens.append(_atencao(
                            f'Sócio retirante "{_esc(nome)}" pode ainda constar na minuta — '
                            'verificar se foi removido da consolidação contratual.'
                        ))
                        alertas_manuais.append(
                            f'Verificar se sócio retirante "{nome}" '
                            'foi removido da consolidação.'
                        )

                # Regra #12 — Retirante que era administrador (via ato anterior)
                _r12_itens, _r12_pends = _verificar_retirante_administrador(
                    [s for s in socios_ultima if s['cpf'] in retirantes],
                    texto_ultima,
                    texto_minuta,
                )
                qs_itens.extend(_r12_itens)
                pendencias.extend(_r12_pends)

        elif not falhas_ult and not falhas_min:
            qs_itens.append(_info(
                'Sócios com CPF não identificados automaticamente nos documentos. '
                'Verificar quadro societário manualmente.'
            ))

        secoes.append(_secao('Comparativo — Quadro Societário', qs_itens))

    # ── 5b. CONFERÊNCIA DE CLÁUSULAS (capital + desimpedimento) ───────────────
    if texto_minuta:
        clausulas_itens = []

        # Regra #10 — Validação aritmética do capital social
        erros_capital = _validar_capital_social(texto_minuta)
        if erros_capital:
            for msg in erros_capital:
                clausulas_itens.append(_erro(msg))
                pendencias.append(msg)
        else:
            # Verificar se conseguiu extrair algum capital (para informar)
            if re.search(
                r'capital\s+social\s+(?:total\s+)?(?:é\s+de|de|no\s+valor\s+de)\s+R\$',
                texto_minuta, re.IGNORECASE,
            ):
                clausulas_itens.append(_ok(
                    'Validação aritmética do capital social: '
                    'valores conferem ou não foi possível extrair quotas individuais para somar.'
                ))

        # Regra #10 — Comparação do capital da minuta com documentos externos (DBE, FCPJ, FCN)
        # Extrai o capital declarado na minuta para usar como referência
        _cap_minuta = None
        for _p_cap in [
            (r'capital\s+social\s+(?:total\s+)?'
             r'(?:é\s+de|de|no\s+valor\s+de|importa\s+em|importando\s+em|'
             r'no\s+valor\s+total\s+de|corresponde\s+a|no\s+montante\s+de)\s+'
             r'R\$\s*([\d\.]+,\d{2})'),
            r'cujo\s+capital\s+social\s+(?:[éê]\s+de|de)\s+R\$\s*([\d\.]+,\d{2})',
            r'capital\s+social[:\s]+R\$\s*([\d\.]+,\d{2})',
        ]:
            _m_cap = re.search(_p_cap, texto_minuta, re.IGNORECASE)
            if _m_cap:
                _cap_minuta = _parse_brl(_m_cap.group(1))
                if _cap_minuta and _cap_minuta > 0:
                    break
                _cap_minuta = None

        if _cap_minuta:
            _TIPOS_CAP_EXT = ('dbe', 'fcpj', 'fcn')
            for _edoc in docs_por_estab:
                for _tipo_cap in _TIPOS_CAP_EXT:
                    if _edoc['ignorados'].get(_tipo_cap, False):
                        continue
                    _txt_ext = _edoc['docs'].get(_tipo_cap)
                    if not _txt_ext:
                        continue
                    _cap_ext, _ = _extrair_capital_bruto(_txt_ext)
                    if _cap_ext is None:
                        continue
                    _label_cap = LABELS_DOC.get(_tipo_cap, _tipo_cap.upper())
                    if abs(_cap_ext - _cap_minuta) > 0.02:
                        clausulas_itens.append(_atencao(
                            f'Capital social na minuta ({_formatar_brl(_cap_minuta)}) '
                            f'difere do valor encontrado no {_esc(_label_cap)} '
                            f'({_formatar_brl(_cap_ext)}). '
                            'Verificar se os documentos correspondem ao mesmo processo.'
                        ))
                        alertas_manuais.append(
                            f'Capital social divergente entre minuta e {_label_cap}: '
                            f'{_formatar_brl(_cap_minuta)} vs {_formatar_brl(_cap_ext)}.'
                        )
                    else:
                        clausulas_itens.append(_ok(
                            f'Capital social confere com {_esc(_label_cap)}: '
                            f'{_formatar_brl(_cap_minuta)}.'
                        ))

        # Regra #13 — Desimpedimento do administrador
        alertas_desimped = _verificar_desimpedimento(texto_minuta)
        if alertas_desimped:
            for msg in alertas_desimped:
                clausulas_itens.append(_atencao(msg))
                alertas_manuais.append(msg)
        else:
            if re.search(
                r'\b(?:administra[çc][aã]o|administrador)\b',
                texto_minuta, re.IGNORECASE,
            ):
                clausulas_itens.append(_ok(
                    'Declaração de desimpedimento localizada na minuta.'
                ))

        if clausulas_itens:
            secoes.append(_secao('Conferência de Cláusulas', clausulas_itens))

    # ── 6. REVISÃO TEXTUAL ────────────────────────────────────────────────────
    if texto_minuta:
        rev_itens = []

        # CEP em formato incorreto — agrupar por valor único com contagem
        ceps_raw = re.findall(r'\bCEP[:\s]*(\d[\d.\s\-\.]{5,11})\b', texto_minuta, re.IGNORECASE)
        ceps_invalidos: dict = {}  # cep_num -> {exemplo, problema, count}
        for cep in ceps_raw:
            cep_strip = cep.strip()
            cep_num = re.sub(r'[^\d]', '', cep_strip)
            if len(cep_num) != 8:
                continue
            cep_correto = f'{cep_num[:5]}-{cep_num[5:]}'
            if cep_strip == cep_correto:
                continue  # formato correto, ignorar
            if '.' in cep_strip:
                problema = 'ponto indevido'
            elif ' ' in cep_strip:
                problema = 'espaço indevido'
            elif '-' not in cep_strip:
                problema = 'sem hífen'
            else:
                problema = 'formato inesperado'
            if cep_num not in ceps_invalidos:
                ceps_invalidos[cep_num] = {'exemplo': cep_strip, 'problema': problema, 'count': 0}
            ceps_invalidos[cep_num]['count'] += 1

        if ceps_invalidos:
            linhas_cep = (
                '<table style="border-collapse:collapse;width:100%;font-size:13px;margin-top:6px;">'
                '<thead><tr style="background:#fef3c7;">'
                '<th style="padding:4px 8px;text-align:left;border:1px solid #d97706;">CEP encontrado</th>'
                '<th style="padding:4px 8px;text-align:left;border:1px solid #d97706;">Problema</th>'
                '<th style="padding:4px 8px;text-align:center;border:1px solid #d97706;">Ocorrências</th>'
                '<th style="padding:4px 8px;text-align:left;border:1px solid #d97706;">Formato esperado</th>'
                '</tr></thead><tbody>'
            )
            for cep_num, info in ceps_invalidos.items():
                cep_correto = f'{cep_num[:5]}-{cep_num[5:]}'
                linhas_cep += (
                    f'<tr><td style="padding:4px 8px;border:1px solid #fcd34d;">{_esc(info["exemplo"])}</td>'
                    f'<td style="padding:4px 8px;border:1px solid #fcd34d;">{_esc(info["problema"])}</td>'
                    f'<td style="padding:4px 8px;border:1px solid #fcd34d;text-align:center;">{info["count"]}</td>'
                    f'<td style="padding:4px 8px;border:1px solid #fcd34d;">{_esc(cep_correto)}</td></tr>'
                )
                alertas_manuais.append(
                    f'CEP com formatação irregular: {info["exemplo"]} '
                    f'({info["problema"]}, {info["count"]} ocorrência(s)) — '
                    f'formato esperado: {cep_correto}'
                )
            linhas_cep += '</tbody></table>'
            rev_itens.append(_atencao(
                f'{len(ceps_invalidos)} CEP(s) com formatação irregular na minuta:{linhas_cep}'
            ))

        # Campos em branco (sublinhados longos)
        brancos = re.findall(r'_{5,}', texto_minuta)
        if brancos:
            rev_itens.append(_atencao(
                f'Há {len(brancos)} campo(s) aparentemente em branco na minuta '
                '(sublinhados). Verificar se todos estão preenchidos.'
            ))

        # Verificar nome empresarial antigo na minuta
        if razao_social and texto_ultima:
            rs_old_m = re.search(
                r'(?:denominada|razão social)[:\s]+([A-ZÁÉÍÓÚÂÊÔÃÕÇ][^\n,;]{3,80})',
                texto_ultima, re.IGNORECASE,
            )
            if rs_old_m:
                rs_old = rs_old_m.group(1).strip().rstrip('.,;')
                rs_norm = _normalizar(razao_social)
                rs_old_norm = _normalizar(rs_old)
                if rs_old_norm != rs_norm and len(rs_old_norm) > 5:
                    if rs_old_norm[:12] in _normalizar(texto_minuta):
                        rev_itens.append(_atencao(
                            f'O nome empresarial anterior ("{_esc(rs_old[:60])}") '
                            'pode estar aparecendo na minuta. Verificar se foi atualizado '
                            'em todos os campos e na consolidação.'
                        ))
                        pendencias.append(
                            'Verificar se o nome empresarial anterior '
                            'foi substituído em toda a minuta.'
                        )

        # Cláusulas identificadas (informativo)
        clausulas = re.findall(
            r'(?:Cláusula|CLÁUSULA)\s+(\w+)', texto_minuta, re.IGNORECASE,
        )
        if len(clausulas) > 2:
            rev_itens.append(_info(
                f'{len(clausulas)} cláusula(s) identificada(s) na minuta. '
                'Verificar sequência e numeração.'
            ))

        if not rev_itens:
            rev_itens.append(_ok(
                'Nenhuma inconsistência de formatação ou preenchimento '
                'identificada automaticamente.'
            ))

        secoes.append(_secao('Revisão Textual — Formatação e Preenchimento', rev_itens))

    # ── 7. PENDÊNCIAS OBJETIVAS ───────────────────────────────────────────────
    if pendencias:
        pend_itens = [_erro(f'{i + 1}. {_esc(p)}') for i, p in enumerate(pendencias)]
        secoes.append(_secao(
            'Pendências Objetivas — Verificar Antes do Protocolo', pend_itens
        ))

    # ── 7b. ALERTAS PARA CONFERÊNCIA MANUAL ──────────────────────────────────
    if alertas_manuais:
        alerta_itens = [_atencao(f'{i + 1}. {_esc(a)}') for i, a in enumerate(alertas_manuais)]
        secoes.append(_secao('Alertas para Conferência Manual', alerta_itens))

    # ── 7c. POSSÍVEIS FALHAS DE EXTRAÇÃO ──────────────────────────────────────
    if falhas_extracao:
        falha_itens = [
            _info(f'{i + 1}. {_esc(f)}') for i, f in enumerate(falhas_extracao)
        ]
        secoes.append(_secao('Possíveis Falhas de Extração', falha_itens))

    # ── 8. CONCLUSÃO ──────────────────────────────────────────────────────────
    partes_concl = []
    if pendencias:
        partes_concl.append(f'<strong>{len(pendencias)}</strong> pendência(s) objetiva(s)')
    if alertas_manuais:
        partes_concl.append(f'<strong>{len(alertas_manuais)}</strong> alerta(s) para conferência manual')
    if falhas_extracao:
        partes_concl.append(f'<strong>{len(falhas_extracao)}</strong> possível(is) falha(s) de extração')

    if partes_concl:
        concl_item = _erro(
            'Foram identificados: ' + ', '.join(partes_concl) + '. '
            'Consulte as seções acima antes do protocolo.'
        )
    else:
        concl_item = _ok(
            'A conferência documental comparativa não identificou divergências objetivas '
            'nos documentos analisados. Proceda com a revisão jurídica antes do protocolo.'
        )

    secoes.append(_secao('Conclusão Operacional', [concl_item]))

    # ── 9. AVISO FINAL ────────────────────────────────────────────────────────
    secoes.append(f'<div class="secao">{AVISO_JURIDICO_FINAL}</div>')

    return '\n'.join(secoes)


# ─── Job: execução em background (sem IA) ─────────────────────────────────────

def _executar_comparativo_background(job_id, kwargs, log_info):
    """Gera o relatório comparativo em thread separada.
    Tenta IA primeiro; se falhar, usa conferência determinística como fallback."""
    t_inicio = time.time()
    try:
        with _jobs_lock:
            _jobs[job_id]['status'] = 'running'
            _jobs[job_id]['msg'] = (
                f'Analisando documentos com IA ({MODELO_IA})...'
                if USE_AI_FOR_MAIN_CONFERENCE
                else 'Gerando relatório comparativo...'
            )

        if USE_AI_FOR_MAIN_CONFERENCE:
            try:
                html = _gerar_relatorio_com_ia(**kwargs, log=log_info)
                modo = f'IA ({MODELO_IA})'
            except Exception as e_ia:
                log_info['ia_erro'] = str(e_ia)[:300]
                log_info['fallback_ativado'] = True
                print(f'[conferencia] job={job_id} IA falhou ({e_ia}) — ativando fallback')
                with _jobs_lock:
                    _jobs[job_id]['msg'] = 'IA indisponível — usando conferência determinística...'
                html_fallback = _gerar_relatorio_html_comparativo(**kwargs, log=log_info)
                html = _banner_fallback(str(e_ia)) + html_fallback
                modo = 'fallback (determinístico)'
        else:
            html = _gerar_relatorio_html_comparativo(**kwargs, log=log_info)
            modo = 'determinístico'

        tempo_total = round(time.time() - t_inicio, 1)
        log_info['tempo_total_s'] = tempo_total

        with _jobs_lock:
            _jobs[job_id]['status'] = 'done'
            _jobs[job_id]['html'] = html
            _jobs[job_id]['msg'] = f'Concluído em {tempo_total}s ({modo})'
            _jobs[job_id]['log'] = log_info

        print(
            f'[conferencia] job={job_id} CONCLUÍDO ({modo}) em {tempo_total}s | '
            f'ia_chamada={log_info.get("ia_chamada")}'
        )

    except Exception as e:
        tempo_total = round(time.time() - t_inicio, 1)
        log_info['erro'] = str(e)[:500]
        with _jobs_lock:
            _jobs[job_id]['status'] = 'error'
            _jobs[job_id]['erro'] = str(e)
            _jobs[job_id]['msg'] = f'Erro após {tempo_total}s'
            _jobs[job_id]['log'] = log_info
        print(f'[conferencia] job={job_id} ERRO: {e}')


# ─── Rotas ────────────────────────────────────────────────────────────────────

@conferencia_bp.route('/', methods=['GET'])
def index():
    if login_obrigatorio():
        return redirect(url_for('auth.login'))
    return render_template('conferencia/index.html')


@conferencia_bp.route('/analisar-minuta', methods=['POST'])
def analisar_minuta():
    """Lê a minuta usando APENAS heurísticas (sem IA).
    Retorna a classificação do processo para o cabeçalho do relatório.
    NÃO define blocos de estabelecimento — o usuário cria manualmente."""
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401

    t_inicio = time.time()
    log = {
        'ia_chamada': False,
        'modo': 'heuristicas_apenas',
        'tempo_s': 0,
    }

    arquivo_minuta = request.files.get('minuta')
    arquivo_ultima = request.files.get('ultima_alteracao')

    texto_minuta = _salvar_e_extrair(arquivo_minuta)
    texto_ultima = _salvar_e_extrair(arquivo_ultima)

    log['tempo_s'] = round(time.time() - t_inicio, 2)

    if not texto_minuta:
        return jsonify({'erro': 'Minuta não enviada ou não legível.', '_log': log}), 400

    heur = _extrair_heuristicas(texto_minuta)

    resultado = {
        'sucesso': True,
        'modo_fallback': False,
        'ia_chamada': False,
        'tipo_processo': heur['tipo_inferido'] or 'alteracao_contratual',
        'numero_alteracao': heur['numero_alteracao'],
        'cnpj_sociedade': heur['cnpjs'][0] if heur['cnpjs'] else None,
        'razao_social': heur['razao_social'],
        'avisos': [],
        # estabelecimentos_alterados vazio — blocos são criados manualmente pelo usuário
        'estabelecimentos_alterados': [],
        '_log': log,
    }

    print(
        f'[conferencia] analisar-minuta SEM IA | '
        f'tipo={resultado["tipo_processo"]} | '
        f'num={resultado["numero_alteracao"]} | '
        f'tempo={log["tempo_s"]}s'
    )

    return jsonify(resultado)


@conferencia_bp.route('/analisar', methods=['POST'])
def analisar():
    """Conferência comparativa completa — SEM IA.
    Recebe documentos organizados por estabelecimento (criados manualmente pelo usuário).
    Retorna {job_id} imediatamente. Frontend faz polling em /status/<job_id>."""
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401

    _cleanup_old_jobs()

    try:
        analise_minuta = json.loads(request.form.get('estabelecimentos_json', '{}'))
    except (json.JSONDecodeError, ValueError):
        analise_minuta = {}

    estabelecimentos = analise_minuta.get('estabelecimentos_alterados') or []
    tipo_processo = analise_minuta.get('tipo_processo', 'alteracao_contratual')
    numero_alteracao = analise_minuta.get('numero_alteracao')
    razao_social = analise_minuta.get('razao_social', '')
    cnpj_sociedade = analise_minuta.get('cnpj_sociedade', '')

    log_info = {
        'razao_social': razao_social,
        'cnpj': cnpj_sociedade,
        'tipo_processo': tipo_processo,
        'estabelecimentos': len(estabelecimentos),
        'docs_recebidos': [],
        'docs_ignorados': [],
        'docs_ausentes': [],
        'ia_chamada': False,
        'conferencia_sem_ia': True,
        'USE_AI_FOR_MAIN_CONFERENCE': USE_AI_FOR_MAIN_CONFERENCE,
    }

    t_extracao = time.time()

    texto_minuta = _salvar_e_extrair(request.files.get('minuta'))
    texto_ultima = _salvar_e_extrair(request.files.get('ultima_alteracao'))
    ultima_ignorada = request.form.get('ignorar_ultima') == '1'

    if texto_minuta:
        log_info['docs_recebidos'].append('minuta')
    else:
        log_info['docs_ausentes'].append('minuta')

    if ultima_ignorada:
        log_info['docs_ignorados'].append('ultima_alteracao')
    elif texto_ultima:
        log_info['docs_recebidos'].append('ultima_alteracao')
    else:
        log_info['docs_ausentes'].append('ultima_alteracao')

    # Documentos por estabelecimento
    docs_por_estab = []

    for i, estab in enumerate(estabelecimentos):
        tipo_e = estab.get('tipo', 'matriz')
        docs_e = {}
        ignorados_e = {}

        for tipo_doc in TIPOS_DOC:
            campo_arq = f'doc_{i}_{tipo_doc}'
            campo_ignorar = f'ignorar_{i}_{tipo_doc}'
            arq = request.files.get(campo_arq)
            ignorado = request.form.get(campo_ignorar) == '1'
            ignorados_e[tipo_doc] = ignorado

            if ignorado:
                docs_e[tipo_doc] = None
                log_info['docs_ignorados'].append(f'{tipo_e}_{i}_{tipo_doc}')
            elif arq:
                texto_doc = _salvar_e_extrair(arq)
                docs_e[tipo_doc] = texto_doc if texto_doc else None
                if texto_doc:
                    log_info['docs_recebidos'].append(f'{tipo_e}_{i}_{tipo_doc}')
                else:
                    log_info['docs_ausentes'].append(f'{tipo_e}_{i}_{tipo_doc}')
            else:
                docs_e[tipo_doc] = None
                log_info['docs_ausentes'].append(f'{tipo_e}_{i}_{tipo_doc}')

        docs_por_estab.append({'docs': docs_e, 'ignorados': ignorados_e})

    log_info['tempo_extracao_s'] = round(time.time() - t_extracao, 2)

    from database import get_user_by_id
    user = get_user_by_id(session.get('user_id'))
    nome_analista = user['nome'] if user else 'Sigma Contabilidade — Setor Societário'

    kwargs = dict(
        texto_minuta=texto_minuta or None,
        texto_ultima=texto_ultima or None,
        ultima_ignorada=ultima_ignorada,
        docs_por_estab=docs_por_estab,
        estabelecimentos=estabelecimentos,
        tipo_processo=tipo_processo,
        numero_alteracao=numero_alteracao,
        razao_social=razao_social,
        cnpj_sociedade=cnpj_sociedade,
        nome_analista=nome_analista,
        data_hoje=date.today().strftime('%d/%m/%Y'),
    )

    job_id = str(uuid.uuid4())[:12]
    with _jobs_lock:
        _jobs[job_id] = {
            'status': 'pending',
            'html': '',
            'msg': (
                f'Iniciando conferência com IA ({MODELO_IA})...'
                if USE_AI_FOR_MAIN_CONFERENCE
                else 'Iniciando conferência comparativa (sem IA)...'
            ),
            'erro': None,
            'log': log_info,
            'criado': datetime.now(),
        }

    t = threading.Thread(
        target=_executar_comparativo_background,
        args=(job_id, kwargs, log_info),
        daemon=True,
    )
    t.start()

    print(
        f'[conferencia] job={job_id} iniciado SEM IA | '
        f'empresa={razao_social} | estabs={len(estabelecimentos)} | '
        f'docs={len(log_info["docs_recebidos"])}'
    )
    return jsonify({'job_id': job_id})


@conferencia_bp.route('/status/<job_id>', methods=['GET'])
def status_job(job_id):
    """Retorna o status atual de um job de conferência (polling do frontend).

    Resposta:
      status:  pending | running | done | error
      html:    HTML gerado (parcial ou completo)
      msg:     mensagem de progresso legível
      erro:    string com o erro, se status=error
      log:     informações técnicas do job
    """
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401

    with _jobs_lock:
        job = _jobs.get(job_id)

    if not job:
        return jsonify({'erro': 'Job não encontrado ou expirado'}), 404

    return jsonify({
        'status': job['status'],
        'html': job['html'],
        'msg': job['msg'],
        'erro': job.get('erro'),
        'log': job.get('log', {}),
    })
