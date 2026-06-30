"""
blueprints/procuracoes.py — Módulo Elaboração de Procurações
Preenchimento determinístico de templates DOCX — sem uso de IA.
"""
import io
import os
import re
import uuid
import time
import logging
from datetime import date
from copy import deepcopy
from flask import Blueprint, render_template, request, jsonify, session, redirect, url_for, send_file, abort

from blueprints.auth import login_obrigatorio

procuracoes_bp = Blueprint('procuracoes', __name__, url_prefix='/procuracoes')

MODELOS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'modelos_procuracao')

MODELOS_DISPONÍVEIS = [
    {
        'nome': 'Procuração Para Serviços Externos - Terceirizado',
        'arquivo': 'procuracao_para_servicos_externos_terceirizado.docx',
        'descricao': 'Procuração para prestadores de serviços externos / terceirizados',
        'tipo': 'servicos_externos',
    },
    {
        'nome': 'Procuração Poderes Específicos - Sigma',
        'arquivo': 'procuracao_poderes_especificos_sigma.docx',
        'descricao': 'Procuração com poderes específicos — modelo padrão Sigma',
        'tipo': 'poderes_especificos',
    },
    {
        'nome': 'Procuração de Serviços Contábeis',
        'arquivo': 'procuracao_sigma.docx',
        'descricao': 'Procuração geral para serviços contábeis',
        'tipo': 'procuracao_geral',
    },
    {
        'nome': 'Substabelecimento Com Reserva - Sigma',
        'arquivo': 'substabelecimento_com_reserva_sigma.docx',
        'descricao': 'Substabelecimento com reserva de poderes — modelo padrão Sigma',
        'tipo': 'subst_com_reserva',
    },
    {
        'nome': 'Substabelecimento Sem Reserva - Sigma',
        'arquivo': 'substabelecimento_sem_reserva_sigma.docx',
        'descricao': 'Substabelecimento sem reserva de poderes — modelo padrão Sigma',
        'tipo': 'subst_sem_reserva',
    },
]

BRUNO_CONTADOR = {
    'nome': 'Bruno Augusto De Leles Carvalho',
    'doc': '009.216.091-32',
    'qualificacao': (
        'Bruno Augusto De Leles Carvalho, brasileiro, Gerente Contábil, '
        'portador da RG n° 4.296.269 DGPC/GO e CPF nº 009.216.091-32, '
        'com endereço profissional situado à Avenida T-2, nº 471, '
        'Focus Business Center, Sala 507, Setor Bueno, Goiânia/GO, CEP 74210-005'
    ),
}

_PROC_CACHE: dict = {}

MESES_PT = [
    'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
    'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro',
]

logger = logging.getLogger(__name__)


# ─── Formatação automática ────────────────────────────────────────────────────

def _fmt_cpf(v: str) -> str:
    digits = re.sub(r'\D', '', v)
    if len(digits) == 11:
        return f'{digits[:3]}.{digits[3:6]}.{digits[6:9]}-{digits[9:]}'
    return v

def _fmt_cnpj(v: str) -> str:
    digits = re.sub(r'\D', '', v)
    if len(digits) == 14:
        return f'{digits[:2]}.{digits[2:5]}.{digits[5:8]}/{digits[8:12]}-{digits[12:]}'
    return v

def _fmt_cep(v: str) -> str:
    digits = re.sub(r'\D', '', v)
    if len(digits) == 8:
        return f'{digits[:5]}-{digits[5:]}'
    return v

def _title(v: str) -> str:
    """Primeira letra de cada palavra maiúscula (Title Case)."""
    return v.strip().title() if v else v

def _upper(v: str) -> str:
    return v.strip().upper() if v else v

def _lower(v: str) -> str:
    return v.strip().lower() if v else v

def _data_extenso() -> str:
    hoje = date.today()
    return f'{hoje.day} de {MESES_PT[hoje.month - 1]} de {hoje.year}'

def _extrair_municipio_uf(endereco: str) -> str:
    """Extrai 'Cidade/UF' do endereço completo.
    Exemplo: 'Rua X, 123, Goiânia/GO, cep 74000-000' → 'Goiânia/GO'
    """
    if not endereco:
        return ''
    m = re.search(r'([\w\s\-]+)/([A-Z]{2})(?:\b|\s|,|$)', endereco, re.UNICODE)
    if m:
        return f'{m.group(1).strip()}/{m.group(2)}'
    return ''

def _fmt_rg(v: str) -> str:
    """Uppercase no órgão emissor do RG. Ex: '12346 ssp/go' → '12346 SSP/GO'"""
    if not v:
        return v
    v = v.strip()
    parts = v.rsplit(' ', 1)
    if len(parts) == 2 and re.search(r'[A-Za-z]', parts[1]):
        return f'{parts[0]} {parts[1].upper()}'
    return v

def _fmt_endereco(v: str) -> str:
    """Title case + normaliza espaços + corrige CEP sem formatação + UF em maiúsculo."""
    if not v:
        return v
    v = re.sub(r'  +', ' ', v.strip())
    v = _title(v)
    v = re.sub(r'\bCep\b', 'CEP', v)
    v = re.sub(r'\bCEP\s+(\d{5})-?(\d{3})\b', r'CEP \1-\2', v)
    # Restaura sigla de UF (2 letras após '/') para maiúsculo: Goiânia/Go → Goiânia/GO
    v = re.sub(r'/([A-Za-z]{2})\b', lambda m: '/' + m.group(1).upper(), v)
    return v

def _fmt_vigencia(v: str) -> str:
    """Formata DDMMYYYY ou YYYY-MM-DD para DD/MM/YYYY. Texto livre passa sem alteração."""
    if not v or not v.strip():
        return 'prazo indeterminado'
    v = v.strip()
    if re.match(r'^\d{8}$', v):
        return f'{v[:2]}/{v[2:4]}/{v[4:]}'
    m = re.match(r'^(\d{4})-(\d{2})-(\d{2})$', v)
    if m:
        return f'{m.group(3)}/{m.group(2)}/{m.group(1)}'
    return v


# ─── Preenchimento de template DOCX ──────────────────────────────────────────

def _substituir_em_paragrafo(paragrafo, valores: dict):
    """Substitui placeholders {{CAMPO}} nos runs do parágrafo, preservando formatação."""
    for run in paragrafo.runs:
        for chave, valor in valores.items():
            ph = '{{' + chave + '}}'
            if ph in run.text:
                run.text = run.text.replace(ph, valor)


def _preencher_template(arquivo: str, valores: dict) -> io.BytesIO:
    """
    Abre o template DOCX, substitui placeholders com os valores formatados
    e retorna BytesIO do documento resultante.
    O documento final é idêntico ao template em fonte, espaçamento e formatação.
    """
    from docx import Document
    caminho = os.path.join(MODELOS_DIR, arquivo)
    doc = Document(caminho)

    for paragrafo in doc.paragraphs:
        _substituir_em_paragrafo(paragrafo, valores)

    # Tabelas (se houver)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _substituir_em_paragrafo(paragrafo, valores)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _valores_outorgante(dados: dict) -> dict:
    """Extrai e formata os campos do outorgante comuns a vários modelos."""
    endereco = dados.get('endereco', '')
    return {
        'RAZAO_SOCIAL':           _upper(dados.get('razao_social', '')),
        'CNPJ':                   _fmt_cnpj(dados.get('cnpj', '')),
        'ENDERECO':               _fmt_endereco(endereco),
        'NOME_REPRESENTANTE':     _upper(dados.get('nome_representante', '')),
        'NACIONALIDADE':          _lower(dados.get('nacionalidade', '')),
        'ESTADO_CIVIL':           _lower(dados.get('estado_civil', '')),
        'PROFISSAO':              _lower(dados.get('profissao', '')),
        'RG_REPRESENTANTE':       _fmt_rg(dados.get('rg_representante', '')),
        'DOMICILIO_REPRESENTANTE':_fmt_endereco(dados.get('domicilio_representante', '')),
        'CPF_REPRESENTANTE':      _fmt_cpf(dados.get('cpf_representante', '')),
        'DATA_EXTENSO':           _data_extenso(),
        'MUNICIPIO_UF':           _extrair_municipio_uf(endereco),
    }


def _gerar_procuracao(dados: dict) -> io.BytesIO:
    tipo = dados.get('tipo_modelo', '')

    if tipo == 'procuracao_geral':
        valores = _valores_outorgante(dados)
        return _preencher_template('procuracao_sigma.docx', valores)

    elif tipo == 'poderes_especificos':
        valores = _valores_outorgante(dados)
        valores['PODERES']  = dados.get('poderes', '').strip()
        valores['VIGENCIA'] = _fmt_vigencia(dados.get('vigencia', ''))
        return _preencher_template('procuracao_poderes_especificos_sigma.docx', valores)

    elif tipo == 'subst_com_reserva':
        valores = {
            'NOME_SUBSTABELECIDO': _title(dados.get('nome_substabelecido', '')),
            'CPF_SUBSTABELECIDO':  _fmt_cpf(dados.get('cpf_substabelecido', '')),
            'RG_SUBSTABELECIDO':   _fmt_rg(dados.get('rg_substabelecido', '')),
            'EMPRESA_CONCEDENTE':  _upper(dados.get('empresa_concedente', '')),
            'CNPJ_CONCEDENTE':     _fmt_cnpj(dados.get('cnpj_concedente', '')),
            'PODERES':             dados.get('poderes', '').strip(),
            'VIGENCIA':            _fmt_vigencia(dados.get('vigencia', '')),
            'DATA_EXTENSO':        _data_extenso(),
        }
        return _preencher_template('substabelecimento_com_reserva_sigma.docx', valores)

    elif tipo == 'subst_sem_reserva':
        valores = {
            'NOME_SUBSTABELECIDO': _title(dados.get('nome_substabelecido', '')),
            'CPF_SUBSTABELECIDO':  _fmt_cpf(dados.get('cpf_substabelecido', '')),
            'RG_SUBSTABELECIDO':   _fmt_rg(dados.get('rg_substabelecido', '')),
            'EMPRESA_CONCEDENTE':  _upper(dados.get('empresa_concedente', '')),
            'CNPJ_CONCEDENTE':     _fmt_cnpj(dados.get('cnpj_concedente', '')),
            'PODERES':             dados.get('poderes', '').strip(),
            'VIGENCIA':            _fmt_vigencia(dados.get('vigencia', '')),
            'DATA_EXTENSO':        _data_extenso(),
        }
        return _preencher_template('substabelecimento_sem_reserva_sigma.docx', valores)

    elif tipo == 'servicos_externos':
        valores = _valores_outorgante(dados)
        valores['NOME_OUTORGADO'] = _title(dados.get('nome_outorgado', ''))
        valores['CPF_OUTORGADO']  = _fmt_cpf(dados.get('cpf_outorgado', ''))
        valores['ORGAO']          = dados.get('orgao', '').strip()
        valores['VIGENCIA']       = _fmt_vigencia(dados.get('vigencia', ''))
        return _preencher_template('procuracao_para_servicos_externos_terceirizado.docx', valores)

    else:
        raise ValueError(f'Tipo de modelo desconhecido: {tipo}')


def _limpar_cache():
    agora = time.time()
    for token in list(_PROC_CACHE.keys()):
        if agora - _PROC_CACHE[token]['ts'] > 3600:
            del _PROC_CACHE[token]


# ─── Rotas ────────────────────────────────────────────────────────────────────

@procuracoes_bp.route('/')
def index():
    if login_obrigatorio():
        return redirect(url_for('auth.login'))
    return render_template('procuracoes/index.html',
                           modelos=MODELOS_DISPONÍVEIS,
                           bruno=BRUNO_CONTADOR)


@procuracoes_bp.route('/modelo/<arquivo>')
def baixar_modelo(arquivo):
    if login_obrigatorio():
        return redirect(url_for('auth.login'))
    nomes_validos = {m['arquivo'] for m in MODELOS_DISPONÍVEIS}
    if arquivo not in nomes_validos:
        abort(404)
    # Serve o arquivo _ORIG (template limpo sem placeholders) para download manual
    orig = arquivo.replace('.docx', '_ORIG.docx')
    caminho_orig = os.path.join(MODELOS_DIR, orig)
    caminho = caminho_orig if os.path.exists(caminho_orig) else os.path.join(MODELOS_DIR, arquivo)
    if not os.path.exists(caminho):
        abort(404)
    nome_download = arquivo.replace('_', ' ').replace('.docx', '').title() + '.docx'
    return send_file(caminho, as_attachment=True, download_name=nome_download,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@procuracoes_bp.route('/gerar', methods=['POST'])
def gerar():
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401

    dados = request.get_json()
    if not dados:
        return jsonify({'erro': 'Dados inválidos'}), 400

    tipo = dados.get('tipo_modelo', '').strip()
    if not tipo:
        return jsonify({'erro': 'Selecione um modelo de procuração'}), 400

    # Validação mínima por tipo
    obrigatorios = {
        'procuracao_geral':   ['razao_social', 'cnpj', 'nome_representante', 'cpf_representante'],
        'poderes_especificos':['razao_social', 'cnpj', 'nome_representante', 'cpf_representante', 'poderes'],
        'subst_com_reserva':  ['nome_substabelecido', 'cpf_substabelecido', 'empresa_concedente', 'cnpj_concedente'],
        'subst_sem_reserva':  ['nome_substabelecido', 'cpf_substabelecido', 'empresa_concedente', 'cnpj_concedente'],
        'servicos_externos':  ['razao_social', 'cnpj', 'nome_representante', 'cpf_representante', 'nome_outorgado', 'cpf_outorgado', 'orgao'],
    }
    labels = {
        'razao_social': 'Razão Social', 'cnpj': 'CNPJ', 'nome_representante': 'Nome do Representante',
        'cpf_representante': 'CPF do Representante', 'poderes': 'Poderes outorgados',
        'nome_substabelecido': 'Nome do Substabelecido', 'cpf_substabelecido': 'CPF do Substabelecido',
        'empresa_concedente': 'Empresa que originou a procuração', 'cnpj_concedente': 'CNPJ da empresa concedente',
        'nome_outorgado':  'Nome do Outorgado',
        'cpf_outorgado':   'CPF do Outorgado',
        'orgao':           'Órgão / escopo dos poderes',
    }
    for campo in obrigatorios.get(tipo, []):
        if not dados.get(campo, '').strip():
            return jsonify({'erro': f'Campo obrigatório: {labels.get(campo, campo)}'}), 400

    try:
        buf = _gerar_procuracao(dados)
    except ValueError as e:
        return jsonify({'erro': str(e)}), 400
    except Exception as e:
        logger.exception("Erro ao gerar procuração")
        return jsonify({'erro': 'Erro interno ao gerar o documento.'}), 500

    _limpar_cache()
    token = str(uuid.uuid4())
    tipo_labels = {
        'procuracao_geral':    'ServicosContabeis',
        'poderes_especificos': 'PoderesEspecificos',
        'subst_com_reserva':   'SubstComReserva',
        'subst_sem_reserva':   'SubstSemReserva',
        'servicos_externos':   'ServicosExternos',
    }
    label_tipo = tipo_labels.get(tipo, tipo)
    empresa = dados.get('razao_social', dados.get('empresa_concedente', 'documento'))
    nome_base = f"Procuracao_{label_tipo}_{empresa.replace(' ', '_')[:25]}"
    _PROC_CACHE[token] = {
        'docx': buf.read(),
        'nome': nome_base,
        'ts':   time.time(),
    }

    return jsonify({'token': token, 'nome': nome_base})


@procuracoes_bp.route('/download/<token>')
def download(token):
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401
    entrada = _PROC_CACHE.get(token)
    if not entrada:
        return jsonify({'erro': 'Documento expirado. Gere novamente.'}), 404
    return send_file(
        io.BytesIO(entrada['docx']),
        as_attachment=True,
        download_name=f'{entrada["nome"]}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
