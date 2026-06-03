"""
blueprints/declaracoes.py — Módulo Elaboração de Declarações e Requerimentos
Gera declarações e requerimentos via Claude API.
"""
import io
import os
import re
import uuid
import time
from datetime import date
from flask import Blueprint, render_template, request, jsonify, session, redirect, url_for, send_file
import anthropic

from blueprints.auth import login_obrigatorio

declaracoes_bp = Blueprint('declaracoes', __name__, url_prefix='/declaracoes')

_DECL_CACHE: dict = {}

MESES_PT = [
    'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
    'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro'
]

TIPOS_DOCUMENTO = [
    # Especial — formulário dedicado
    'Requerimento de Uso do Solo - Goiânia',
    # Declarações empresariais
    'Declaração de Enquadramento como Microempresa (ME)',
    'Declaração de Enquadramento como Empresa de Pequeno Porte (EPP)',
    'Declaração de Enquadramento como MEI',
    'Declaração de Ausência de Débitos',
    'Declaração de Inexistência de Impedimento Legal',
    'Declaração de Capacidade Técnica',
    'Declaração de Quadro Societário',
    'Declaração de Endereço da Empresa',
    'Declaração de Atividade',
    'Declaração de Faturamento',
    'Declaração de Não Contratação de Menores',
    'Declaração de Ausência de Vínculo Empregatício',
    # Requerimentos
    'Requerimento de Certidão',
    'Requerimento de Cancelamento de Inscrição',
    'Requerimento de Baixa de CNPJ',
    'Requerimento de Alteração Cadastral',
    'Requerimento de Parcelamento de Débitos',
    'Requerimento de Vista de Processo',
    'Requerimento de Certidão Negativa de Débitos',
    'Requerimento ao Corpo de Bombeiros',
    'Requerimento à Vigilância Sanitária',
    'Requerimento à Prefeitura',
    'Requerimento à Receita Federal',
    'Requerimento ao INSS',
    # Outros
    'Carta de Apresentação',
    'Carta de Anuência',
    'Termo de Responsabilidade',
    'Termo de Ciência e Concordância',
    'Declaração de Responsabilidade Técnica',
    'Outro (descrever)',
]


def _data_extenso():
    hoje = date.today()
    return f'Goiânia, {hoje.day} de {MESES_PT[hoje.month - 1]} de {hoje.year}.'


def _gerar_documento_claude(dados: dict) -> str:
    """Chama Claude para gerar declaração/requerimento."""
    api_key = os.getenv('ANTHROPIC_API_KEY')
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY não configurada.")

    client = anthropic.Anthropic(api_key=api_key)

    tipo = dados.get('tipo', '')
    empresa_nome = dados.get('empresa_nome', '')
    empresa_cnpj = dados.get('empresa_cnpj', '')
    empresa_qualificacao = dados.get('empresa_qualificacao', '')
    representante_nome = dados.get('representante_nome', '')
    representante_qualificacao = dados.get('representante_qualificacao', '')
    destinatario = dados.get('destinatario', '')
    objeto = dados.get('objeto', '')
    observacoes = dados.get('observacoes', '')

    e_requerimento = 'requerimento' in tipo.lower() or 'carta' in tipo.lower() or 'termo' in tipo.lower()

    prompt = f"""Você é um especialista em direito empresarial e societário brasileiro.
Elabore {"um requerimento" if e_requerimento else "uma declaração"} profissional e juridicamente correto(a) com base nos dados abaixo.

TIPO DE DOCUMENTO: {tipo}

EMPRESA/DECLARANTE:
Nome: {empresa_nome}
CNPJ: {empresa_cnpj}
{empresa_qualificacao}

REPRESENTANTE LEGAL:
{representante_nome}
{representante_qualificacao}

{f"DESTINATÁRIO: {destinatario}" if destinatario else ""}
{f"OBJETO/FINALIDADE: {objeto}" if objeto else ""}
{f"INFORMAÇÕES ADICIONAIS: {observacoes}" if observacoes else ""}

DATA: {_data_extenso()}

INSTRUÇÕES:
- Escreva o documento completo em português brasileiro formal e jurídico
- Use marcadores <<<EMPRESA>>> onde o nome da empresa deve aparecer em destaque
- Use marcadores <<<REPRESENTANTE>>> onde o nome do representante deve aparecer em destaque
- Estrutura: cabeçalho do tipo de documento, qualificação do declarante/requerente, conteúdo principal, encerramento, local e data, linha de assinatura
- Tom solene, formal, preciso — sem redundâncias
- NÃO inclua explicações ou comentários — apenas o texto do documento
"""

    mensagem = client.messages.create(
        model='claude-sonnet-4-6',
        max_tokens=3000,
        messages=[{'role': 'user', 'content': prompt}]
    )
    return mensagem.content[0].text


def _gerar_docx_declaracao(texto: str, titulo: str) -> io.BytesIO:
    """Gera Word (.docx) da declaração/requerimento."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    def set_font(run, bold=False, size=12):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.insert(0, rFonts)

    linhas = texto.split('\n')
    for linha in linhas:
        linha_stripped = linha.strip()
        if not linha_stripped:
            doc.add_paragraph()
            continue

        is_titulo = (linha_stripped.isupper() and len(linha_stripped) < 80) or \
                    any(linha_stripped.startswith(p) for p in
                        ['DECLARAÇÃO', 'REQUERIMENTO', 'CARTA', 'TERMO'])

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_titulo else WD_ALIGN_PARAGRAPH.JUSTIFY

        partes = re.split(r'(<<<.+?>>>)', linha_stripped)
        for parte in partes:
            if parte.startswith('<<<') and parte.endswith('>>>'):
                conteudo = parte[3:-3]
                run = p.add_run(conteudo)
                set_font(run, bold=True, size=14 if is_titulo else 12)
                run.font.color.rgb = RGBColor(0xA7, 0x2C, 0x31)
            else:
                run = p.add_run(parte)
                set_font(run, bold=is_titulo, size=14 if is_titulo else 12)

    # Rodapé
    footer = doc.sections[0].footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run('Sigma Contabilidade — Além da Contabilidade | www.gsigma.com.br')
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(0xA7, 0x2C, 0x31)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _limpar_cache():
    agora = time.time()
    for token in list(_DECL_CACHE.keys()):
        if agora - _DECL_CACHE[token]['ts'] > 3600:
            del _DECL_CACHE[token]


# ─── Rotas ────────────────────────────────────────────────────────────────────

@declaracoes_bp.route('/')
def index():
    if login_obrigatorio():
        return redirect(url_for('auth.login'))
    return render_template('declaracoes/index.html', tipos=TIPOS_DOCUMENTO)


@declaracoes_bp.route('/gerar', methods=['POST'])
def gerar():
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401

    dados = request.get_json()
    if not dados:
        return jsonify({'erro': 'Dados inválidos'}), 400

    for campo in ['tipo', 'empresa_nome', 'representante_nome', 'objeto']:
        if not dados.get(campo, '').strip():
            return jsonify({'erro': f'Campo obrigatório: {campo}'}), 400

    try:
        texto = _gerar_documento_claude(dados)
    except ValueError as e:
        return jsonify({'erro': str(e)}), 400
    except Exception as e:
        return jsonify({'erro': f'Erro ao gerar com IA: {str(e)}'}), 500

    try:
        buf = _gerar_docx_declaracao(texto, dados.get('tipo', 'Documento'))
    except Exception as e:
        return jsonify({'erro': f'Erro ao gerar documento: {str(e)}'}), 500

    _limpar_cache()
    token = str(uuid.uuid4())
    tipo_slug = dados.get('tipo', 'documento').split(' ')[0].replace('/', '')
    empresa_slug = dados.get('empresa_nome', '').replace(' ', '_')[:20]
    nome_arquivo = f"{tipo_slug}_{empresa_slug}.docx"
    _DECL_CACHE[token] = {
        'docx': buf.read(),
        'nome': nome_arquivo,
        'ts': time.time(),
    }

    return jsonify({'token': token, 'nome': nome_arquivo})


@declaracoes_bp.route('/uso-do-solo', methods=['POST'])
def uso_do_solo():
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401

    dados = request.get_json()
    if not dados:
        return jsonify({'erro': 'Dados inválidos'}), 400

    for campo in ['empresa_nome', 'inscricao_imobiliaria', 'endereco', 'area_m2']:
        if not dados.get(campo, '').strip():
            return jsonify({'erro': f'Campo obrigatório: {campo}'}), 400
    if not dados.get('cnaes_lista'):
        return jsonify({'erro': 'Informe pelo menos um CNAE.'}), 400

    try:
        buf_docx = _gerar_docx_uso_do_solo(dados)
        buf_pdf  = _gerar_pdf_uso_do_solo(dados)
    except Exception as e:
        return jsonify({'erro': f'Erro ao gerar documento: {str(e)}'}), 500

    _limpar_cache()
    token = str(uuid.uuid4())
    empresa_slug = dados.get('empresa_nome', '').replace(' ', '_')[:20]
    nome_docx = f"Uso_do_Solo_{empresa_slug}.docx"
    nome_pdf  = f"Uso_do_Solo_{empresa_slug}.pdf"
    _DECL_CACHE[token] = {
        'docx': buf_docx.read(),
        'pdf':  buf_pdf.read(),
        'nome': nome_docx,
        'nome_pdf': nome_pdf,
        'ts': time.time(),
    }
    return jsonify({'token': token, 'nome': nome_docx, 'nome_pdf': nome_pdf})



def _gerar_docx_uso_do_solo(dados: dict) -> io.BytesIO:
    """Gera DOCX usando o template oficial da Prefeitura de Goiânia (COD. 646)."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    template_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'static', 'templates', 'template_uso_solo.docx'
    )

    empresa_nome = dados.get('empresa_nome', '').upper()
    iptu         = dados.get('inscricao_imobiliaria', '')
    endereco     = dados.get('endereco', '')
    area_m2      = dados.get('area_m2', '')
    cnaes_lista  = dados.get('cnaes_lista', [])
    observacoes  = dados.get('observacoes', '').strip()

    with open(template_path, 'rb') as f:
        buf_in = io.BytesIO(f.read())
    doc = Document(buf_in)

    def _add_value(cell, value, size=9):
        """Adiciona run com valor após label+linebreak existentes no template."""
        para = cell.paragraphs[0]
        run = para.add_run(value)
        run.font.name = 'Arial'
        run.font.size = Pt(size)
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rPr.insert(0, rFonts)

    def _set_text(cell, value, size=9, center=False):
        """Limpa parágrafo e define texto na célula."""
        para = cell.paragraphs[0]
        para.clear()
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if value:
            run = para.add_run(value)
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rPr.insert(0, rFonts)

    # TABLE 2 — Seção 1: dados do requerente/imóvel
    t2 = doc.tables[2]
    _add_value(t2.rows[1].cells[0], empresa_nome)
    _add_value(t2.rows[2].cells[0], endereco)
    _add_value(t2.rows[3].cells[0], iptu)
    _add_value(t2.rows[4].cells[0], area_m2)

    # TABLE 3 — CNAEs (até 24 CNAEs em 12 pares)
    t3 = doc.tables[3]
    for i, cnae in enumerate(cnaes_lista[:24]):
        row_idx = 2 + (i // 2)
        col_code, col_esc = (1, 2) if i % 2 == 0 else (4, 5)
        _set_text(t3.rows[row_idx].cells[col_code], cnae.get('codigo', ''))
        esc_val = 'X' if cnae.get('escritorio', '').lower() == 'sim' else ''
        _set_text(t3.rows[row_idx].cells[col_esc], esc_val, center=True)

    # TABLE 4 — Observações
    t4 = doc.tables[4]
    _set_text(t4.rows[0].cells[1], observacoes, size=8)

    # Substituir data no documento — buscar em todos os parágrafos
    hoje = date.today()
    data_formatada = f'{hoje.day} de {MESES_PT[hoje.month - 1].capitalize()} de {hoje.year}'
    for para in doc.paragraphs:
        full = ''.join(r.text for r in para.runs)
        if 'GOIÂNIA,' in full and ('de' in full or '___' in full):
            for i, run in enumerate(para.runs):
                if 'GOIÂNIA,' in run.text:
                    run.text = f'GOIÂNIA, {data_formatada}.'
                elif any(x in run.text for x in ['de ', '___', '20_']):
                    run.text = ''

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _gerar_pdf_uso_do_solo(dados: dict) -> io.BytesIO:
    """Gera PDF do Requerimento de Uso do Solo com reportlab."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Image
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    empresa_nome = dados.get('empresa_nome', '').upper()
    iptu         = dados.get('inscricao_imobiliaria', '')
    endereco     = dados.get('endereco', '')
    area_m2      = dados.get('area_m2', '')
    cnaes_lista  = dados.get('cnaes_lista', [])
    observacoes  = dados.get('observacoes', '').strip()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
        topMargin=0.2*cm, bottomMargin=0.2*cm,
        leftMargin=2*cm, rightMargin=2*cm)

    W = 17 * cm
    gray = colors.HexColor('#D9D9D9')
    blk = colors.black

    def sty(size, bold=False, align=TA_LEFT, leading=None):
        fn = 'Helvetica-Bold' if bold else 'Helvetica'
        return ParagraphStyle(f's{size}{"b" if bold else ""}a{align}',
            fontName=fn, fontSize=size, alignment=align,
            leading=leading or (size * 1.3))

    story = []

    # ── CABEÇALHO: logo + instituição ──
    logo_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'static', 'templates', 'logo_prefeitura.png'
    )
    try:
        logo = Image(logo_path, width=7*cm, height=1.8*cm) if os.path.exists(logo_path) else ''
    except Exception:
        logo = ''

    inst_para = Paragraph(
        'Secretaria Municipal de Efici\u00eancia \u2013 <b>SEFIC</b><br/>'
        'Ger\u00eancia de Informa\u00e7\u00e3o do Uso do Solo e N\u00famero Predial \u2013 <b>GERINF</b>',
        ParagraphStyle('inst', fontName='Helvetica', fontSize=9, leading=12))

    t_hdr = Table([[logo, inst_para]], colWidths=[8*cm, 9*cm])
    t_hdr.setStyle(TableStyle([
        ('BOX', (0,0),(-1,-1), 0.5, blk),
        ('INNERGRID', (0,0),(-1,-1), 0.5, blk),
        ('VALIGN', (0,0),(-1,-1), 'MIDDLE'),
        ('LEFTPADDING', (0,0),(-1,-1), 4),
        ('RIGHTPADDING', (0,0),(-1,-1), 4),
        ('TOPPADDING', (0,0),(-1,-1), 4),
        ('BOTTOMPADDING', (0,0),(-1,-1), 4),
    ]))
    story.append(t_hdr)

    # ── TÍTULO ──
    t_title = Table([
        [Paragraph('REQUERIMENTO \u2013 USO DO SOLO ATIVIDADE ECON\u00d4MICA',
                   sty(12, True, TA_CENTER)),
         Paragraph('COD. 646', sty(12, True, TA_CENTER))]
    ], colWidths=[14*cm, 3*cm])
    t_title.setStyle(TableStyle([
        ('BOX', (0,0),(-1,-1), 0.5, blk),
        ('INNERGRID', (0,0),(-1,-1), 0.5, blk),
        ('VALIGN', (0,0),(-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0),(-1,-1), 5),
        ('BOTTOMPADDING', (0,0),(-1,-1), 5),
    ]))
    story.append(t_title)

    # ── SEÇÃO 1 ──
    def field_cell(label, value):
        return Paragraph(
            f'<b><font size="8">{label}</font></b><br/><font size="9">{value}</font>',
            ParagraphStyle('fc', fontName='Helvetica', fontSize=9, leading=12))

    t_s1 = Table([
        [Paragraph('1- REQUERENTE / DADOS DO IM\u00d3VEL / \u00c1REA OCUPADA',
                   sty(9, True))],
        [field_cell('NOME OU RAZ\u00c3O SOCIAL:', empresa_nome)],
        [field_cell('ENDERE\u00c7O DO ESTABELECIMENTO: (RUA/AV, QUADRA, LOTE, BAIRRO, CEP)',
                    endereco)],
        [field_cell('INSCRI\u00c7\u00c3O IMOBILI\u00c1RIA: (IPTU)', iptu)],
        [field_cell('\u00c1REA OCUPADA PELO ESTABELECIMENTO - M\u00b2:', area_m2)],
    ], colWidths=[W])
    t_s1.setStyle(TableStyle([
        ('BOX', (0,0),(-1,-1), 0.5, blk),
        ('INNERGRID', (0,0),(-1,-1), 0.5, blk),
        ('BACKGROUND', (0,0),(0,0), gray),
        ('TOPPADDING', (0,0),(-1,-1), 4),
        ('BOTTOMPADDING', (0,0),(-1,-1), 4),
        ('LEFTPADDING', (0,0),(-1,-1), 4),
        ('RIGHTPADDING', (0,0),(-1,-1), 4),
        ('MINROWHEIGHT', (0,1),(-1,-1), 0.8*cm),
    ]))
    story.append(t_s1)

    # ── TABELA CNAE ──
    sf = W / (14 * cm)
    cnae_cols = [w * cm * sf for w in [1, 4, 2, 1, 4, 2]]
    ch = sty(8, True, TA_CENTER)
    cl = sty(8)
    cv = sty(9, False, TA_CENTER)

    cnae_rows = [
        [Paragraph('2- ATIVIDADES \u2013 CNAE / DESCRI\u00c7\u00c3O:', ch),
         '', '', '', '', ''],
        ['', Paragraph('CNAE', ch), Paragraph('ESCRIT\u00d3RIO*', ch),
         '', Paragraph('CNAE', ch), Paragraph('ESCRIT\u00d3RIO*', ch)],
    ]
    for row_i in range(12):
        li, ri = row_i * 2, row_i * 2 + 1
        lc = cnaes_lista[li] if li < len(cnaes_lista) else {}
        rc = cnaes_lista[ri] if ri < len(cnaes_lista) else {}
        cnae_rows.append([
            Paragraph('N\u00ba CNAE:', cl),
            Paragraph(lc.get('codigo', '') if lc else '', cv),
            Paragraph('X' if (lc and lc.get('escritorio', '').lower() == 'sim') else '', cv),
            Paragraph('N\u00ba CNAE:', cl),
            Paragraph(rc.get('codigo', '') if rc else '', cv),
            Paragraph('X' if (rc and rc.get('escritorio', '').lower() == 'sim') else '', cv),
        ])

    t_cnae = Table(cnae_rows, colWidths=cnae_cols)
    t_cnae.setStyle(TableStyle([
        ('BOX', (0,0),(-1,-1), 0.5, blk),
        ('INNERGRID', (0,0),(-1,-1), 0.5, blk),
        ('BACKGROUND', (0,0),(5,0), gray),
        ('BACKGROUND', (0,1),(5,1), gray),
        ('SPAN', (0,0),(5,0)),
        ('TOPPADDING', (0,0),(-1,-1), 3),
        ('BOTTOMPADDING', (0,0),(-1,-1), 3),
        ('LEFTPADDING', (0,0),(-1,-1), 3),
        ('RIGHTPADDING', (0,0),(-1,-1), 3),
        ('VALIGN', (0,0),(-1,-1), 'MIDDLE'),
        ('MINROWHEIGHT', (0,2),(-1,-1), 0.55*cm),
    ]))
    story.append(t_cnae)

    # ── RODAPÉ (observações + Art.299 + assinatura) ──
    art299 = (
        'Art. 299 - Omitir, em documento p\u00fablico ou particular, declara\u00e7\u00e3o que dele devia '
        'constar, ou nele inserir ou fazer inserir declara\u00e7\u00e3o falsa ou diversa da que devia '
        'ser escrita, com o fim de prejudicar direito, criar obriga\u00e7\u00e3o ou alterar a verdade '
        'sobre fato juridicamente relevante: Pena - reclus\u00e3o, de um a cinco anos, '
        'e multa, se o documento \u00e9 p\u00fablico, e reclus\u00e3o de um a tr\u00eas anos, e multa, de '
        'quinhentos mil r\u00e9is a cinco contos de r\u00e9is, se o documento \u00e9 particular.'
    )

    obs_sty = ParagraphStyle('obs', fontName='Helvetica', fontSize=8, leading=10)
    art_sty = ParagraphStyle('art', fontName='Helvetica', fontSize=7, leading=9)
    sig_sty = ParagraphStyle('sig', fontName='Helvetica', fontSize=8,
                              alignment=TA_CENTER, leading=12)

    # Data formatada: "GOIÂNIA, 8 de Abril de 2026."
    hoje = date.today()
    data_str = f'GOI\u00c2NIA, {hoje.day} de {MESES_PT[hoje.month - 1].capitalize()} de {hoje.year}.'

    t_foot = Table([
        [Paragraph('<b>3- OBSERVA\u00c7\u00d5ES:</b>', obs_sty),
         Paragraph(observacoes, obs_sty)],
        ['', ''],
        ['', ''],
        ['', ''],
        [Paragraph('TELEFONE: (62) 3095-6240', obs_sty),
         Paragraph('E-MAIL: societario@gsigma.com.br', obs_sty)],
        [Paragraph(art299, art_sty), ''],
        [Paragraph(
            f'{data_str}'
            '<br/><br/><br/>________________________________________________________________________<br/>ASSINATURA DO REQUERENTE',
            sig_sty), ''],
    ], colWidths=[8.5*cm, 8.5*cm])
    t_foot.setStyle(TableStyle([
        ('BOX', (0,0),(-1,-1), 0.5, blk),
        ('INNERGRID', (0,0),(-1,-1), 0.5, blk),
        ('BACKGROUND', (0,0),(0,0), gray),
        ('SPAN', (0,5),(1,5)),
        ('SPAN', (0,6),(1,6)),
        ('TOPPADDING', (0,0),(-1,-1), 4),
        ('BOTTOMPADDING', (0,0),(-1,-1), 4),
        ('LEFTPADDING', (0,0),(-1,-1), 4),
        ('RIGHTPADDING', (0,0),(-1,-1), 4),
        ('VALIGN', (0,0),(-1,-1), 'TOP'),
        ('ALIGN', (0,6),(1,6), 'CENTER'),
        ('VALIGN', (0,6),(1,6), 'MIDDLE'),
        ('MINROWHEIGHT', (0,1),(1,3), 0.4*cm),
    ]))
    story.append(t_foot)

    def footer_page(canvas, _doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 7)
        canvas.drawCentredString(
            LETTER[0] / 2, 0.4 * cm,
            'Av. do Cerrado, 999 \u2013 Park Lozandes, Pa\u00e7o Municipal, '
            '2\u00ba andar, Bloco C, Goi\u00e2nia \u2013 GO. CEP: 74884-900'
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=footer_page, onLaterPages=footer_page)
    buf.seek(0)
    return buf



@declaracoes_bp.route('/download/<token>')
def download(token):
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401

    entrada = _DECL_CACHE.get(token)
    if not entrada:
        return jsonify({'erro': 'Documento expirado. Gere novamente.'}), 404

    return send_file(
        io.BytesIO(entrada['docx']),
        as_attachment=True,
        download_name=entrada['nome'],
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@declaracoes_bp.route('/download-pdf/<token>')
def download_pdf(token):
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401

    entrada = _DECL_CACHE.get(token)
    if not entrada or 'pdf' not in entrada:
        return jsonify({'erro': 'Documento expirado. Gere novamente.'}), 404

    return send_file(
        io.BytesIO(entrada['pdf']),
        as_attachment=True,
        download_name=entrada.get('nome_pdf', 'Uso_do_Solo.pdf'),
        mimetype='application/pdf'
    )
