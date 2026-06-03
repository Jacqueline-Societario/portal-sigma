import os
import io
import re
import json
import time
import uuid
import tempfile
from datetime import date, datetime
from flask import Flask, request, jsonify, render_template, send_file, session, redirect, url_for
from werkzeug.utils import secure_filename
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
import openai
from dotenv import load_dotenv

# ─── ReportLab (PDF) ─────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm as RL_CM
from reportlab.lib.colors import HexColor as RL_Color
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 Table, TableStyle, PageBreak, KeepTogether,
                                 Image as RL_Image)
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib import colors as rl_colors
from reportlab.pdfbase.pdfmetrics import stringWidth as rl_stringWidth
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

load_dotenv()

# ─── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(name)s: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
)
logger = logging.getLogger('portal.contrato')


# ─── Exceção customizada para falha na API da Anthropic ──────────────────────
class APIIndisponivel(Exception):
    """Erro de serviço da API Anthropic (créditos, autenticação, rate-limit, etc.)"""
    pass


app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
_secret_key = os.getenv("SECRET_KEY")
if not _secret_key:
    raise RuntimeError("SECRET_KEY nao configurada no ambiente. Defina no .env antes de iniciar o portal.")
app.secret_key = _secret_key
SENHA_ACESSO = 'gestao3095'


def login_obrigatorio():
    return not session.get('autenticado')

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extrair_texto_pdf(caminho):
    texto = []
    with pdfplumber.open(caminho) as pdf:
        for pagina in pdf.pages:
            t = pagina.extract_text()
            if t:
                texto.append(t)
    return '\n'.join(texto)


def extrair_texto_docx(caminho):
    doc = Document(caminho)
    return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())


def extrair_texto(caminho, extensao):
    if extensao == 'pdf':
        return extrair_texto_pdf(caminho)
    elif extensao == 'docx':
        return extrair_texto_docx(caminho)
    return ''


def gerar_com_claude(texto_contrato, alteracoes, tem_consolidacao=True):
    api_key = os.getenv('OPENAI_API_KEY')
    if not api_key:
        raise ValueError("OPENAI_API_KEY não configurada. Adicione a chave no arquivo .env")

    client = openai.OpenAI(api_key=api_key, timeout=600)

    # Filtrar "Consolidação" da lista de alterações efetivas
    alteracoes_efetivas = [a for a in alteracoes if a['tipo'] != 'Consolidação']

    alteracoes_formatadas = '\n'.join(
        f"- {a['tipo']}: {a['descricao']}" for a in alteracoes_efetivas
    )

    # Ementa pré-construída — exatamente os nomes dos eventos selecionados
    tipos_ementa = [a['tipo'] for a in alteracoes_efetivas]
    if len(tipos_ementa) == 1:
        ementa_pre = tipos_ementa[0] + '.'
    elif len(tipos_ementa) > 1:
        ementa_pre = ';\n'.join(tipos_ementa[:-1]) + ';\n' + tipos_ementa[-1] + '.'
    else:
        ementa_pre = 'Alteração do contrato social.'

    # Data de hoje em português
    meses = ['janeiro','fevereiro','março','abril','maio','junho',
             'julho','agosto','setembro','outubro','novembro','dezembro']
    hoje = date.today()
    data_hoje = f"{hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"

    # Bloco condicional: cláusula de consolidação + PARTE 3
    if tem_consolidacao:
        clausula_consolidacao = """
CLÁUSULA [ÚLTIMA] – Após as alterações acima, o contrato social passa a viger com as cláusulas e condições abaixo consolidadas, sendo que, aquelas não contempladas neste instrumento, tornam-se extintas.

[ATENÇÃO: A seção ===PARTE 3: CONSOLIDAÇÃO=== só deve aparecer DEPOIS desta última cláusula.]

===PARTE 3: CONSOLIDAÇÃO===

CONSOLIDAÇÃO DO CONTRATO SOCIAL DA SOCIEDADE EMPRESÁRIA LIMITADA
[NOME DA EMPRESA PÓS-ALTERAÇÃO em maiúsculas]
CNPJ nº [XXX]
N.I.R.E nº [XXX]

[Qualificação COMPLETA de cada sócio — mesma do início, já com dados atualizados pelas alterações. Envolva o nome de cada sócio em **asteriscos duplos**: **NOME COMPLETO**.]

[Único sócio / Únicos sócios] da empresa **[NOME DA EMPRESA]**, com sede na [endereço], com seu contrato social devidamente registrado e arquivado na Junta Comercial do Estado de Goiás sob o nº [NIRE], inscrita no CNPJ sob o nº [CNPJ], resolve(m) consolidar seu contrato social conforme as cláusulas e condições seguintes:

I – DA DENOMINAÇÃO SOCIAL, SEDE E FILIAIS
[Texto consolidado com nome empresarial e nome fantasia atualizados, endereço da sede, menção à possibilidade de filiais. Citar art. 997, II, Lei nº 10.406/2002.]
Parágrafo único: Ao presente instrumento aplicam-se, supletivamente, no que caberem, as disposições da Lei Sociedade por Ações (Lei 6.404/76) nos termos do parágrafo único do artigo 1053 do Código Civil (Lei 10.406/2002).

II – OBJETO SOCIAL
A sociedade limitada tem por objeto social a exploração do ramo de:
[Listar os CNAEs do objeto social — se alterado, usar o novo; se não alterado, manter o do contrato original.]

III – PRAZO DE DURAÇÃO
A sociedade limitada iniciou suas atividades em [data original] e seu prazo de duração é indeterminado.

IV – CAPITAL SOCIAL
O capital social é de R$ [valor] ([valor por extenso]), divididos em [nº] ([por extenso]) quotas de valor nominal de R$ 1,00 (um real) cada uma, subscritas e totalmente integralizadas em moeda corrente do País, ficando distribuídas da seguinte forma:
[Tabela: Sócio | Quotas | Valor (R$) | %]
[Linha por sócio]
[Total]
Parágrafo único: A responsabilidade de cada sócio é restrita ao valor de suas quotas, mas todos respondem solidariamente pela integralização do capital social. (art. 1.052, CC/2002).

V – ADMINISTRAÇÃO
A administração da sociedade é exercida [forma] pelo(s) sócio(s) **[NOME(S) DO(S) ADMINISTRADOR(ES)]**, com os poderes e atribuições de representação ativa e passiva na sociedade, judicial e extrajudicialmente, podendo praticar todos os atos compreendidos no objeto social, sempre de interesse da sociedade, autorizado o uso do nome empresarial, vedado, no entanto, fazê-lo em atividades estranhas ao interesse social ou assumir obrigações seja em favor de qualquer dos quotistas ou de terceiros, bem como onerar ou alienar bens imóveis da sociedade, sem autorização dos outros sócios. (arts. 997, VI; 1.013; 1.015; 1.064, CC/2002)

VI – DESIMPEDIMENTO
O(s) Administrador(es) declara(m), sob as penas da lei, de que não está(ão) impedido(s) de exercer a administração da empresa, por lei especial, ou em virtude de condenação criminal, ou por se encontrar sob os efeitos dela, a pena que vede, ainda que temporariamente, o acesso a cargos públicos; ou por crime falimentar, de prevaricação, peita ou suborno, concussão, peculato, ou contra a economia popular, contra o sistema financeiro nacional, contra normas de defesa da concorrência, contra as relações de consumo, fé pública, ou a propriedade. (art. 1.011, § 1º; CC/2002).

VII – EXERCÍCIO SOCIAL E BALANÇO PATRIMONIAL
O exercício social será coincidente com o ano-calendário, terminando em 31 de dezembro de cada ano, quando será procedido o levantamento do balanço patrimonial e efetuada a apuração de resultados, os quais serão atribuídos ao(s) sócio(s) proporcionalmente às suas quotas de capital. Podendo os lucros a critério do(s) mesmo(s), serem distribuídos ou ficarem em reserva na sociedade.
Parágrafo primeiro: A sociedade limitada poderá levantar demonstrações contábeis intermediárias, a qualquer tempo, para fins de cisão parcial ou total, fusão e incorporação, retirada do(s) sócio(s) ou ainda, para quaisquer atos julgados necessários pelo(s) sócio(s).
Parágrafo segundo: O(s) sócio(s) será(ão) obrigado(s) à reposição dos lucros e das quantias retiradas, a qualquer título ainda que autorizados pelo contrato, quando tais lucros ou quantia se distribuírem com prejuízo do capital.

VIII – RETIRADA OU FALECIMENTO DO SÓCIO
[Texto conforme número de sócios — usar o modelo correto do contrato original.]

IX – CONSELHO FISCAL
Fica estabelecida que a Sociedade Limitada não terá Conselho Fiscal.

X – DELIBERAÇÕES
Em suas deliberações os administradores adotarão preferencialmente a forma estabelecida no parágrafo 3º do artigo 1.072 do Código Civil (Lei 10.406/2002).

XI – CASOS OMISSOS
Os casos omissos no presente instrumento serão resolvidos com observância dos preceitos do Código Civil (Lei 10.406/2002) e seus dispositivos aplicáveis.

XII – FORO
Fica eleito o foro de Goiânia, Goiás para o exercício e o cumprimento dos direitos e obrigações resultantes do presente instrumento, com exclusão de qualquer outro, seja qual for ou vier a ser o futuro domicílio do(s) sócio(s).

Lavrado em 01 (uma) via, lido, conferido, compreendido, elaborado de conformidade e nos termos, condições e intenção propostas pelo(s) sócio(s) ora presente(s) e que o(s) mesmo(s) assine(m) e rubrique(m) este instrumento, assumindo integralmente as responsabilidades legais decorrentes do presente ato, obrigando-se fielmente por si, seus herdeiros e sucessores legais a cumpri-lo em todos os seus termos.

[MUNICÍPIO DA SEDE]/[UF], {data_hoje}.

**[NOME DO SÓCIO 1 EM MAIÚSCULAS]**
Sócio Administrador

**[NOME DO SÓCIO 2 EM MAIÚSCULAS — se houver]**
Sócio Administrador"""
        instrucao_partes = """Gere o instrumento em DOIS blocos separados por marcadores exatos:
===PARTE 1: INSTRUMENTO DE ALTERAÇÃO===
===PARTE 3: CONSOLIDAÇÃO==="""
        resolucao = "RESOLVE(M), proceder a uma Alteração no seu Contrato Social e Consolidar o mesmo, nos termos da Lei nº 10.406/2002, mediante as seguintes cláusulas:"
    else:
        clausula_consolidacao = ""
        instrucao_partes = """Gere o instrumento em UM bloco com o marcador:
===PARTE 1: INSTRUMENTO DE ALTERAÇÃO===
(NÃO inclua ===PARTE 3: CONSOLIDAÇÃO=== — o documento é apenas de alteração, sem consolidação)"""
        resolucao = "RESOLVE(M), proceder a uma Alteração no seu Contrato Social, nos termos da Lei nº 10.406/2002, mediante as seguintes cláusulas:"

    prompt = f"""Você é um especialista em direito empresarial brasileiro com vasta experiência na elaboração de instrumentos de alteração contratual societária. Siga rigorosamente o Código Civil (Lei 10.406/2002), as normas do DREI e o modelo padrão da Sigma Contábil e Gestão.

REGRA DE FORMATAÇÃO: Use **asteriscos duplos** SOMENTE para envolver nomes de sócios e nome da empresa no texto corrido. Para todo o restante do texto, use texto puro sem Markdown. Para tabelas de capital social, use o formato | col | col | com linha separadora |---|---|.

CONTRATO SOCIAL ANTERIOR:
{texto_contrato}

ALTERAÇÕES SOLICITADAS:
{alteracoes_formatadas}

{instrucao_partes}

ESTRUTURA OBRIGATÓRIA — PARTE 1 (Instrumento de Alteração):

[Nome da empresa ATUAL em maiúsculas — linha única]
CNPJ nº [XXX]
N.I.R.E nº [XXX]
[Nº ordinal em algarismo]ª ALTERAÇÃO CONTRATUAL

Ementa:
{ementa_pre}

[Qualificação COMPLETA de cada sócio: **NOME EM MAIÚSCULAS**, nacionalidade, estado civil, profissão, data de nascimento, nome dos pais, documento de identidade, CPF, endereço completo — extraídos do contrato anterior. Parágrafo único por sócio.]

[Único sócio / Únicos sócios] da empresa **[NOME DA EMPRESA]**, com sede [endereço], com seu contrato social devidamente registrado e arquivado na Junta Comercial do Estado de Goiás sob o nº [NIRE], inscrita no CNPJ sob o nº [CNPJ], {resolucao}

CLÁUSULA PRIMEIRA – [Redigir a cláusula relativa à 1ª alteração. Texto corrido, linguagem jurídica formal. Incluir parágrafos se necessário.]

CLÁUSULA SEGUNDA – [2ª alteração]
[Continuar numerando, uma cláusula por alteração.]
{clausula_consolidacao}

Lavrado em 01 (uma) via, lido, conferido, compreendido, elaborado de conformidade e nos termos, condições e intenção propostas pelo(s) sócio(s) ora presente(s) e que o(s) mesmo(s) assine(m) e rubrique(m) este instrumento, assumindo integralmente as responsabilidades legais decorrentes do presente ato, obrigando-se fielmente por si, seus herdeiros e sucessores legais a cumpri-lo em todos os seus termos.

[MUNICÍPIO DA SEDE]/[UF], {data_hoje}.

**[NOME DO SÓCIO 1 EM MAIÚSCULAS]**
Sócio Administrador

**[NOME DO SÓCIO 2 EM MAIÚSCULAS — se houver]**
Sócio Administrador

REGRAS ABSOLUTAS:
- NUNCA inventar dados: usar APENAS informações do contrato anterior
- Dados ausentes: indicar como [A PREENCHER]
- Linguagem jurídica formal brasileira, sem abreviações no corpo do texto
- Numeração ordinal das cláusulas: PRIMEIRA, SEGUNDA, TERCEIRA...
- Parágrafos: Parágrafo Único (se só um) ou Parágrafo Primeiro / Segundo (se mais de um)
- A EMENTA acima está FIXADA — use exatamente como fornecida, sem alterar
- Nomes de sócios e nome da empresa: envolva sempre em **asteriscos duplos** no corpo do texto
- ORDEM ADMISSÃO + SAÍDA (cessão de cotas): quando houver admissão de novo sócio E saída de sócio por cessão de cotas, a cláusula de ADMISSÃO do novo sócio vem PRIMEIRO; a cláusula de SAÍDA do sócio retirante vem IMEDIATAMENTE APÓS e deve citar o novo sócio já admitido como cessionário das cotas
- DESIMPEDIMENTO OBRIGATÓRIO: sempre que houver cláusula nomeando ou alterando ADMINISTRADOR da sociedade, a cláusula IMEDIATAMENTE SEGUINTE deve ser a de DESIMPEDIMENTO (declaração de que o administrador nomeado não está impedido de exercer a administração, conforme art. 1.011, § 1º, CC/2002)
- INTEGRALIZAÇÃO OBRIGATÓRIA: sempre que houver alteração na distribuição de cotas, no quadro societário ou no capital social, a cláusula que trata dessas alterações deve ser IMEDIATAMENTE SEGUIDA pela cláusula de RESPONSABILIDADE PELA INTEGRALIZAÇÃO DO CAPITAL SOCIAL (responsabilidade solidária, conforme art. 1.052 CC/2002)
- RERATIFICAÇÃO — ESPECIFICAR ERRO MATERIAL: quando a reratificação se destinar a corrigir erro material, a cláusula deve especificar expressamente "para correção de erro material" e indicar o instrumento de origem (ex: "para correção de erro material constante da [Xª] Alteração Contratual registrada em [data]")
- CLÁUSULA SEPARADA DE NOVA COMPOSIÇÃO DO CAPITAL: quando houver cessão de quotas entre sócios (entrada de novo sócio + saída de sócio), a nova composição do quadro societário NUNCA deve ficar embutida dentro da cláusula de saída/cessão; deve ser uma CLÁUSULA PRÓPRIA dedicada à "Nova Composição do Capital Social" com a tabela de distribuição e parágrafo único de integralização
- RENÚNCIA DO ADMINISTRADOR ANTERIOR: quando houver troca de administrador (novo administrador entra, outro sai), a cláusula de administração deve incluir Parágrafo Único com a renúncia expressa do administrador anterior ao cargo, declarando que não possui qualquer restrição ao referido ato
- LOCALIDADE DA ASSINATURA: substituir [MUNICÍPIO DA SEDE]/[UF] pelo município onde está a sede da empresa (extraído do endereço da sede no contrato anterior), NÃO usar Goiânia/GO por padrão; se a sede for em Trindade, usar Trindade/GO; se for em Goiânia, usar Goiânia/GO; etc.
- CONSOLIDAÇÃO — TEMPO VERBAL: nas cláusulas consolidadas, usar PRESENTE do indicativo para descrever o estado atual da sociedade; ERRADO: "a administração será exercida", "a sociedade terá por objeto", "a sede será"; CORRETO: "a administração é exercida", "a sociedade tem por objeto", "a sede é"; EXCEÇÃO: cláusulas que tratam de eventos futuros ou hipotéticos PODEM manter o futuro — exercício social ("o exercício social será coincidente"), falecimento do sócio ("a sociedade continuará"), conselho fiscal ("não terá Conselho Fiscal"), deliberações, casos omissos e foro"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except openai.APIStatusError as exc:
        # Log técnico completo — nunca vai ao frontend
        logger.error(
            "OpenAI API error | status=%s | request_id=%s | message=%s | ts=%s",
            exc.status_code,
            getattr(exc, 'request_id', 'n/a'),
            str(exc.message),
            datetime.now().isoformat(),
        )
        raise APIIndisponivel(
            "Não foi possível gerar o contrato no momento porque o serviço de IA "
            "está indisponível por limitação de créditos ou configuração da API. "
            "Entre em contato com o administrador do sistema."
        ) from exc
    except openai.APIConnectionError as exc:
        logger.error("OpenAI connection error | message=%s | ts=%s", str(exc), datetime.now().isoformat())
        raise APIIndisponivel(
            "Não foi possível gerar o contrato no momento porque o serviço de IA "
            "está indisponível por limitação de créditos ou configuração da API. "
            "Entre em contato com o administrador do sistema."
        ) from exc


LOGO_PATH = os.path.join(os.path.dirname(__file__), 'static', 'sigma_logo.png')
FONTE = 'Times New Roman'
COR_SIGMA = RGBColor(0xA7, 0x2C, 0x31)  # Vermelho bordô Sigma

ORDINAIS_PT = {
    1: 'primeira', 2: 'segunda', 3: 'terceira', 4: 'quarta', 5: 'quinta',
    6: 'sexta', 7: 'sétima', 8: 'oitava', 9: 'nona', 10: 'décima',
    11: 'décima primeira', 12: 'décima segunda', 13: 'décima terceira',
    14: 'décima quarta', 15: 'décima quinta', 16: 'décima sexta',
    17: 'décima sétima', 18: 'décima oitava', 19: 'décima nona', 20: 'vigésima'
}


# ─── PDF — Registro de fontes Times New Roman ────────────────────────────────
_FONT_WIN = '/mnt/c/Windows/Fonts'
_PDF_FONTE = 'TimesNewRoman'
try:
    pdfmetrics.registerFont(TTFont('TimesNewRoman',            f'{_FONT_WIN}/times.ttf'))
    pdfmetrics.registerFont(TTFont('TimesNewRoman-Bold',       f'{_FONT_WIN}/timesbd.ttf'))
    pdfmetrics.registerFont(TTFont('TimesNewRoman-Italic',     f'{_FONT_WIN}/timesi.ttf'))
    pdfmetrics.registerFont(TTFont('TimesNewRoman-BoldItalic', f'{_FONT_WIN}/timesbi.ttf'))
    pdfmetrics.registerFontFamily('TimesNewRoman',
        normal='TimesNewRoman', bold='TimesNewRoman-Bold',
        italic='TimesNewRoman-Italic', boldItalic='TimesNewRoman-BoldItalic')
except Exception:
    _PDF_FONTE = 'Times-Roman'  # fallback PDF padrão

_PDF_SIGMA_RED = RL_Color('#A72C31')

# Cache temporário de documentos gerados (token → {docx, pdf, nome, ts})
_DOCS_CACHE: dict = {}


def _limpar_cache_antigo():
    """Remove entradas com mais de 1 hora do cache."""
    agora = time.time()
    for token in list(_DOCS_CACHE.keys()):
        if agora - _DOCS_CACHE[token]['ts'] > 3600:
            del _DOCS_CACHE[token]


def _pdf_estilo(nome, bold=False, italic=False, size=12,
                align=TA_JUSTIFY, sb=0, sa=4, li=0, fi=0):
    """Cria ParagraphStyle para reportlab."""
    if bold and italic:
        fn = f'{_PDF_FONTE}-BoldItalic' if _PDF_FONTE == 'TimesNewRoman' else 'Times-BoldItalic'
    elif bold:
        fn = f'{_PDF_FONTE}-Bold' if _PDF_FONTE == 'TimesNewRoman' else 'Times-Bold'
    elif italic:
        fn = f'{_PDF_FONTE}-Italic' if _PDF_FONTE == 'TimesNewRoman' else 'Times-Italic'
    else:
        fn = _PDF_FONTE if _PDF_FONTE == 'TimesNewRoman' else 'Times-Roman'
    return ParagraphStyle(nome, fontName=fn, fontSize=size, alignment=align,
                          spaceBefore=sb, spaceAfter=sa,
                          leftIndent=li, firstLineIndent=fi,
                          leading=size * 1.35, wordWrap='LTR')


def _pdf_esc(texto):
    """Escapa HTML e converte marcadores <<<>>> para <b>."""
    t = texto.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    # Restaurar marcadores de negrito: &lt;&lt;&lt;NOME&gt;&gt;&gt; → <b>NOME</b>
    t = re.sub(r'&lt;&lt;&lt;(.+?)&gt;&gt;&gt;', r'<b>\1</b>', t)
    return t


def _flush_pending(story, pending, proximo):
    """Adiciona pending+proximo agrupados (KeepTogether) ou só proximo se não há pending."""
    if pending is not None:
        story.append(KeepTogether([pending, proximo]))
    else:
        story.append(proximo)


def _inserir_tabela_pdf(story, linhas_tabela, text_width):
    """Converte bloco Markdown em Table reportlab proporcional à largura do texto."""
    rows_raw = [l for l in linhas_tabela if not is_separador_tabela(l)]
    if not rows_raw:
        return
    rows = []
    for linha in rows_raw:
        cells = [c.strip() for c in linha.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    rows = [r + [''] * (num_cols - len(r)) for r in rows]

    # Proporções de coluna (mesmas do Word) — somam à largura total do texto
    if num_cols == 4:
        prop = [4234, 1512, 2117, 1209]
    elif num_cols == 3:
        prop = [4234, 2117, 2721]
    elif num_cols == 2:
        prop = [4914, 4158]
    else:
        prop = [1] * num_cols
    total_p = sum(prop)
    col_widths = [text_width * p / total_p for p in prop]

    fn_bold = f'{_PDF_FONTE}-Bold' if _PDF_FONTE == 'TimesNewRoman' else 'Times-Bold'
    fn_norm = _PDF_FONTE if _PDF_FONTE == 'TimesNewRoman' else 'Times-Roman'
    st_hdr  = ParagraphStyle('tbl_hdr',  fontName=fn_bold, fontSize=10, alignment=TA_CENTER, leading=13)
    st_cell = ParagraphStyle('tbl_cell', fontName=fn_norm, fontSize=10, alignment=TA_CENTER, leading=13)

    table_data = []
    for i, row in enumerate(rows):
        st = st_hdr if i == 0 else st_cell
        table_data.append([Paragraph(_pdf_esc(c), st) for c in row])

    t = Table(table_data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ('GRID',          (0, 0), (-1, -1), 0.5, rl_colors.black),
        ('BACKGROUND',    (0, 0), (-1,  0), rl_colors.HexColor('#F2F2F2')),
        ('ALIGN',         (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING',    (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING',   (0, 0), (-1, -1), 5),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 5),
    ]))
    story.append(t)
    story.append(Spacer(1, 6))


def _pdf_header_footer(canvas, doc):
    """Cabeçalho sem logo + rodapé: separador + 'Contrato elaborado por [logo]' + tagline."""
    canvas.saveState()
    page_w = A4[0]
    fn_norm = _PDF_FONTE if _PDF_FONTE == 'TimesNewRoman' else 'Times-Roman'
    fn_bold = f'{_PDF_FONTE}-Bold' if _PDF_FONTE == 'TimesNewRoman' else 'Times-Bold'
    right_x = page_w - 2.5 * RL_CM

    # Linha separadora
    footer_base = 3.2 * RL_CM
    canvas.setStrokeColor(_PDF_SIGMA_RED)
    canvas.setLineWidth(0.5)
    canvas.line(2.5 * RL_CM, footer_base, right_x, footer_base)

    canvas.setFillColor(_PDF_SIGMA_RED)

    # Linha 1: "Contrato elaborado por  " + logo na mesma linha (alinhados à direita)
    logo_w = 2.0 * RL_CM   # igual ao .docx (cx=720000 EMUs ≈ 2.0 cm)
    logo_h = 0.58 * RL_CM  # proporção original (cy=209334 EMUs ≈ 0.58 cm)
    logo_x = right_x - logo_w
    # Posicionar logo: base alinhada ao separador - 1 linha (texto + gap)
    logo_y = footer_base - logo_h - 8  # logo fica logo abaixo do separador
    line1_y = logo_y + 2               # baseline do texto alinhada à base da logo
    if os.path.exists(LOGO_PATH):
        canvas.drawImage(LOGO_PATH, logo_x, logo_y,
                         width=logo_w, height=logo_h,
                         preserveAspectRatio=True, anchor='sw', mask='auto')
    canvas.setFont(fn_norm, 7)
    canvas.drawRightString(logo_x, line1_y, 'Contrato elaborado por  ')

    # Linha 2: "Além da Contabilidade  |  gsigma.com.br" — bold, direita
    line2_y = logo_y - 12  # abaixo da logo, com 12pt de espaço
    footer_text = 'Além da Contabilidade  |  gsigma.com.br'
    canvas.setFont(fn_bold, 7)
    canvas.drawRightString(right_x, line2_y, footer_text)
    # Hyperlink clicável sobre "gsigma.com.br"
    url_text = 'gsigma.com.br'
    url_w = rl_stringWidth(url_text, fn_bold, 7)
    canvas.linkURL('https://gsigma.com.br',
                   (right_x - url_w, line2_y - 2, right_x, line2_y + 6), relative=0)

    canvas.restoreState()


def gerar_pdf(texto_gerado, nome_empresa, tem_consolidacao=True):
    """Gera PDF com reportlab usando o mesmo conteúdo do DOCX."""
    ST = {
        'corpo':       _pdf_estilo('corpo',       size=12, align=TA_JUSTIFY, sa=3),
        'center14':    _pdf_estilo('center14',     size=14, bold=True, align=TA_CENTER, sa=2),
        'num_alt':     _pdf_estilo('num_alt',      size=13, bold=True, italic=True, align=TA_RIGHT, sa=4),
        'ementa_item': _pdf_estilo('ementa_item',  size=12, bold=True, align=TA_JUSTIFY, li=18, fi=-18),
        'clausula':    _pdf_estilo('clausula',     size=12, align=TA_JUSTIFY, sa=3),
        'data':        _pdf_estilo('data',         size=12, align=TA_RIGHT, sb=4),
        'cargo':       _pdf_estilo('cargo',        size=12, italic=True, align=TA_CENTER, sa=2),
        'assinatura':  _pdf_estilo('assinatura',   size=12, bold=True, align=TA_CENTER, sb=16, sa=2),
        'romano':      _pdf_estilo('romano',       size=12, bold=True, align=TA_CENTER, sa=3),
        'titulo_cons': _pdf_estilo('titulo_cons',  size=15, bold=True, align=TA_CENTER, sb=18, sa=4),
    }

    buf = io.BytesIO()
    doc_rl = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5 * RL_CM, rightMargin=2.5 * RL_CM,
        topMargin=3.8 * RL_CM, bottomMargin=4.6 * RL_CM
    )
    text_w = doc_rl.width  # largura útil do texto para tabelas

    story = []

    # Separar as 3 partes (parte2 = assinaturas, mantém fluxo natural)
    partes_pdf = {'parte1': '', 'parte2': '', 'parte3': ''}
    if '===PARTE 1:' in texto_gerado:
        blocos = texto_gerado.split('===PARTE ')
        for bloco in blocos[1:]:
            if bloco.startswith('1:'):
                raw = bloco[bloco.find('\n') + 1:]
                fim = raw.find('===PARTE')
                partes_pdf['parte1'] = (raw[:fim] if fim >= 0 else raw).strip()
            elif bloco.startswith('2:'):
                raw = bloco[bloco.find('\n') + 1:]
                fim = raw.find('===PARTE')
                partes_pdf['parte2'] = (raw[:fim] if fim >= 0 else raw).strip()
            elif bloco.startswith('3:'):
                partes_pdf['parte3'] = bloco[bloco.find('\n') + 1:].strip()
    else:
        partes_pdf['parte1'] = texto_gerado

    def _flush_tabela(buf_tab):
        if buf_tab:
            temp = []
            _inserir_tabela_pdf(temp, buf_tab, text_w)
            if temp:
                story.append(KeepTogether(temp))
            buf_tab.clear()

    def _trim_spacers():
        """Remove Spacers extras no final de story antes de uma quebra."""
        while story and isinstance(story[-1], Spacer):
            story.pop()

    for chave in ['parte1', 'parte2', 'parte3']:
        conteudo = partes_pdf.get(chave, '').strip()
        if not conteudo:
            continue
        if chave == 'parte3' and not tem_consolidacao:
            continue

        conteudo = limpar_texto(conteudo)
        em_assinaturas = False
        apos_ementa = False
        _pending_clausula = None   # título de cláusula aguardando KeepTogether com próximo §
        buf_tabela: list = []

        for linha in conteudo.split('\n'):
            l = linha.strip()

            # — Buffer de tabela Markdown
            if is_linha_tabela(linha):
                buf_tabela.append(linha)
                continue
            else:
                _flush_tabela(buf_tabela)

            if not l:
                # Linha vazia: libera pending_clausula isolado (sem corpo logo abaixo)
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                    _pending_clausula = None
                story.append(Spacer(1, 3))
                continue

            tipo = classificar_linha(linha)
            esc  = _pdf_esc(l)

            # Detectar início de assinaturas
            if re.match(r'^Goi[aâ]nia/GO,', l, re.IGNORECASE):
                em_assinaturas = True

            # --- Ementa label ---
            if re.match(r'^Ementa\s*:', l, re.IGNORECASE):
                apos_ementa = True
                _flush_pending(story, _pending_clausula, Paragraph(f'<b>{esc}</b>', ST['corpo']))
                _pending_clausula = None
                continue

            # --- Itens de ementa ---
            _eh_clausula_pdf = re.match(
                r'^(CLÁUSULA|CLAUSULA|ARTIGO|ART\.|PARÁGRAFO|PARAGRAFO'
                r'|I{1,3}V?|VI{0,3}|IX|X{1,3}|CNPJ|N\.I\.R\.E|NIRE)',
                l, re.IGNORECASE)
            if apos_ementa and l and not _eh_clausula_pdf and len(l) < 120:
                _flush_pending(story, _pending_clausula, Paragraph(f'<b>\u2022  {esc}</b>', ST['ementa_item']))
                _pending_clausula = None
                if l.endswith('.'):
                    apos_ementa = False
                continue
            elif apos_ementa and (not l or _eh_clausula_pdf):
                apos_ementa = False

            # --- Assinatura: nome do sócio ---
            if em_assinaturas and detectar_assinatura(l):
                nome_limpo = l[3:-3].strip() if (l.startswith('<<<') and l.endswith('>>>')) else l
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                    _pending_clausula = None
                story.append(Paragraph(f'<b>{_pdf_esc(nome_limpo)}</b>', ST['assinatura']))
                continue

            # --- Renderizar tipo ---
            if tipo == 'titulo_consolidacao':
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                    _pending_clausula = None
                story.append(Paragraph(esc, ST['titulo_cons']))

            elif tipo == 'num_alteracao':
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                    _pending_clausula = None
                story.append(Paragraph(esc, ST['num_alt']))
                story.append(HRFlowable(width='100%', thickness=0.5,
                                        color=_PDF_SIGMA_RED, spaceAfter=4))

            elif tipo in ('cnpj_nire', 'maiusculas'):
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                    _pending_clausula = None
                story.append(Paragraph(esc, ST['center14']))

            elif tipo == 'data':
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                    _pending_clausula = None
                story.append(Paragraph(esc, ST['data']))

            elif tipo == 'cargo':
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                    _pending_clausula = None
                story.append(Paragraph(esc, ST['cargo']))

            elif tipo == 'romano':
                m = re.match(r'^((?:I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII)\s*[–\-]\s*)(.+)', l)
                txt = f'<b>{_pdf_esc(m.group(1))}</b>{_pdf_esc(m.group(2))}' if m else f'<b>{esc}</b>'
                # romanos na consolidação: keepWithNext também
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                _pending_clausula = Paragraph(txt, ST['romano'])

            elif tipo == 'clausula':
                lu = l.upper()
                if lu.startswith('PARÁGRAFO') or lu.startswith('PARAGRAFO'):
                    partes_c = l.split(':', 1)
                    txt = f'<b>{_pdf_esc(partes_c[0].rstrip())}: </b>'
                    if len(partes_c) > 1:
                        txt += _pdf_esc(partes_c[1].strip())
                elif '–' in l:
                    partes_c = l.split('–', 1)
                    txt = f'<b>{_pdf_esc(partes_c[0].rstrip())} – </b>{_pdf_esc(partes_c[1].strip())}'
                elif ' — ' in l:
                    partes_c = l.split(' — ', 1)
                    txt = f'<b>{_pdf_esc(partes_c[0].rstrip())} — </b>{_pdf_esc(partes_c[1].strip())}'
                elif ':' in l:
                    partes_c = l.split(':', 1)
                    txt = f'<b>{_pdf_esc(partes_c[0].rstrip())}: </b>'
                    if len(partes_c) > 1:
                        txt += _pdf_esc(partes_c[1].strip())
                elif ' - ' in l:
                    partes_c = l.split(' - ', 1)
                    txt = f'<b>{_pdf_esc(partes_c[0].rstrip())} - </b>{_pdf_esc(partes_c[1].strip())}'
                else:
                    txt = f'<b>{esc}</b>'
                # Guardar como pending: será agrupado com próximo parágrafo (keepWithNext)
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                _pending_clausula = Paragraph(txt, ST['clausula'])

            elif tipo == 'ementa_label':
                if _pending_clausula is not None:
                    story.append(_pending_clausula)
                    _pending_clausula = None
                story.append(Paragraph(f'<b>{esc}</b>', ST['corpo']))

            else:  # corpo
                corpo_para = Paragraph(esc, ST['corpo'])
                if _pending_clausula is not None:
                    # KeepTogether: título da cláusula não ficará sozinho no fim de página
                    story.append(KeepTogether([_pending_clausula, corpo_para]))
                    _pending_clausula = None
                else:
                    story.append(corpo_para)

        # Flush de itens pendentes no final da parte
        _flush_tabela(buf_tabela)
        if _pending_clausula is not None:
            story.append(_pending_clausula)
            _pending_clausula = None

        # Quebra de página entre partes (sem espaços extras antes)
        precisa_quebra = (
            (chave == 'parte1' and not partes_pdf.get('parte2', '').strip() and
             partes_pdf.get('parte3', '').strip() and tem_consolidacao)
            or
            (chave == 'parte2' and partes_pdf.get('parte3', '').strip() and tem_consolidacao)
        )
        if precisa_quebra:
            _trim_spacers()
            story.append(PageBreak())

    if not story:
        story.append(Paragraph('Documento gerado.', ST['corpo']))

    doc_rl.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
    buf.seek(0)
    return buf


def extrair_ordinal_filename(texto):
    """Extrai o número da alteração do texto e retorna nome de arquivo por extenso."""
    m = re.search(r'(\d+)[ªº°]\s*ALTERA[ÇC][AÃ]O\s+CONTRATUAL', texto, re.IGNORECASE)
    if m:
        n = int(m.group(1))
        palavra = ORDINAIS_PT.get(n, f'{n}ª')
        return f'{palavra} alteração contratual.docx'
    for palavra in ['primeira', 'segunda', 'terceira', 'quarta', 'quinta', 'sexta',
                    'sétima', 'oitava', 'nona', 'décima']:
        if re.search(rf'\b{palavra}\b\s*ALTERA[ÇC][AÃ]O', texto, re.IGNORECASE):
            return f'{palavra} alteração contratual.docx'
    return 'alteração contratual.docx'


def set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name = FONTE
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Idioma Português Brasil
    rPr = run._r.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'pt-BR')
    lang.set(qn('w:eastAsia'), 'pt-BR')
    rPr.append(lang)


def _add_para(container, text, align, size_pt, bold=False, italic=False, color=None,
              space_before=None, space_after=None):
    """Cria parágrafo com run único — helper interno."""
    p = container.add_paragraph()
    p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    set_font(run, size_pt, bold=bold, italic=italic, color=color)
    return p


def configurar_header(section):
    """Cabeçalho sem logo (logo foi movida para o rodapé inferior direito)."""
    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)


def _add_hyperlink_run(para, text, url):
    """Adiciona run com hyperlink clicável em um parágrafo."""
    part = para.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    hl.set(qn('w:history'), '1')

    run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), 'C00000')   # vermelho bordô — igual ao modelo
    rPr.append(color_el)
    fonts_el = OxmlElement('w:rFonts')
    fonts_el.set(qn('w:ascii'), FONTE)
    fonts_el.set(qn('w:hAnsi'), FONTE)
    rPr.append(fonts_el)
    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), '14')   # 7pt = 14 half-points
    rPr.append(sz_el)
    run_el.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    run_el.append(t)
    hl.append(run_el)
    para._p.append(hl)


def configurar_footer(section):
    """Rodapé Sigma: 3 parágrafos — separador + 'Contrato elaborado por + logo' + tagline."""
    footer = section.footer
    # Limpar parágrafo inicial
    fp0 = footer.paragraphs[0]
    fp0.clear()
    fp0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp0.paragraph_format.space_before = Pt(0)
    fp0.paragraph_format.space_after = Pt(0)
    r_sep = fp0.add_run('_' * 80)
    r_sep.font.name = FONTE
    r_sep.font.size = Pt(7)
    r_sep.font.color.rgb = COR_SIGMA

    # "Contrato elaborado por  " + LOGO na mesma linha, alinhado à direita
    fp1 = footer.add_paragraph()
    fp1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp1.paragraph_format.space_before = Pt(0)
    fp1.paragraph_format.space_after = Pt(0)
    r1 = fp1.add_run('Contrato elaborado por  ')
    r1.font.name = FONTE
    r1.font.size = Pt(7)
    r1.font.color.rgb = COR_SIGMA
    if os.path.exists(LOGO_PATH):
        run_logo = fp1.add_run()
        run_logo.add_picture(LOGO_PATH, width=Cm(2.0))

    # "Além da Contabilidade  |  gsigma.com.br" — bold, right, com hyperlink
    fp2 = footer.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp2.paragraph_format.space_before = Pt(0)
    fp2.paragraph_format.space_after = Pt(0)
    r2 = fp2.add_run('Além da Contabilidade  |  ')
    r2.font.name = FONTE
    r2.font.size = Pt(7)
    r2.font.bold = True
    r2.font.color.rgb = COR_SIGMA
    # Hyperlink clicável
    _add_hyperlink_run(fp2, 'gsigma.com.br', 'https://gsigma.com.br')


def classificar_linha(linha):
    """Retorna o tipo de formatação para cada linha de texto."""
    import re
    l = linha.strip()
    lu = l.upper()

    if not l:
        return 'vazio'

    # Consolidação (título grande centralizado)
    if 'CONSOLIDAÇÃO DO CONTRATO SOCIAL' in lu or 'CONTRATO SOCIAL CONSOLIDADO' in lu:
        return 'titulo_consolidacao'

    # Número de alteração ordinal (ex: "5ª ALTERAÇÃO CONTRATUAL")
    if re.match(r'^\d+[ªº°]\s*(ALTERA|Altera)', l):
        return 'num_alteracao'

    # Cabeçalho empresa (CNPJ, NIRE — linhas separadas)
    if lu.startswith('CNPJ') or lu.startswith('N.I.R.E') or lu.startswith('NIRE'):
        return 'cnpj_nire'

    # Data no estilo "Goiânia/GO, XX de XXXXXXXX de XXXX."
    if re.match(r'^Goi[aâ]nia/GO,', l, re.IGNORECASE):
        return 'data'

    # Assinatura: linha toda em maiúsculas + curta (nome do sócio)
    # e linha de cargo abaixo (Sócio Administrador / Sócio)
    if re.match(r'^Sócio\s*Administrador', l, re.IGNORECASE) or \
       re.match(r'^Sóci[ao]\s+[Rr]etirante', l, re.IGNORECASE) or \
       re.match(r'^Sóci[ao]\s+[Ii]ngressante', l, re.IGNORECASE) or \
       re.match(r'^Sóci[ao]\s+[Rr]emanescente', l, re.IGNORECASE) or \
       re.match(r'^Sóci[ao]$', l, re.IGNORECASE) or \
       re.match(r'^Administrador[ao]?$', l, re.IGNORECASE) or \
       l in ('Sócio Administrador', 'Sócio', 'Administrador',
             'Sócio Retirante', 'Sócio Ingressante', 'Sócio Remanescente',
             'Sócia Retirante', 'Sócia Ingressante', 'Sócia Remanescente'):
        return 'cargo'

    # Algarismos romanos — seções da consolidação (I –, II –, III –, IV –...)
    if re.match(r'^(I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII)\s*[–\-]', l):
        return 'romano'

    # Cláusulas / artigos / parágrafos
    if lu.startswith('CLÁUSULA') or lu.startswith('CLAUSULA') or \
       lu.startswith('ARTIGO ') or lu.startswith('ART. ') or \
       lu.startswith('PARÁGRAFO') or lu.startswith('PARAGRAFO'):
        return 'clausula'

    # Ementa (label) — "Ementa:"
    if re.match(r'^Ementa\s*:', l, re.IGNORECASE):
        return 'ementa_label'

    # Itens de ementa: a), b), c)... ou a. b. c.
    if re.match(r'^[a-z]\)', l) or re.match(r'^[a-z]\.\s', l):
        return 'ementa_item'

    # Linha toda em maiúsculas de tamanho médio (título de seção / nome de empresa)
    if l == l.upper() and len(l) > 8 and len(l) < 120 and not re.match(r'^\d', l):
        return 'maiusculas'

    return 'corpo'


def _add_corpo_runs(para, texto, size):
    """Adiciona runs ao parágrafo processando marcadores <<<>>> como negrito."""
    if '<<<' in texto and '>>>' in texto:
        partes = re.split(r'(<<<.+?>>>)', texto)
        for parte in partes:
            if parte.startswith('<<<') and parte.endswith('>>>'):
                run = para.add_run(parte[3:-3])
                set_font(run, size, bold=True)
            elif parte:
                run = para.add_run(parte)
                set_font(run, size, bold=False)
    else:
        run = para.add_run(texto)
        set_font(run, size, bold=False)


def adicionar_linha_doc(doc, linha):
    """Adiciona parágrafo ao documento com formatação exata do modelo Sigma 5ª Alt."""
    import re
    tipo = classificar_linha(linha)

    if tipo == 'vazio':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        return

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    if tipo == 'titulo_consolidacao':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(18)   # espaço de respiro antes do título de consolidação
        run = p.add_run(linha.strip())
        set_font(run, 16, bold=True)

    elif tipo == 'num_alteracao':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(linha.strip())
        set_font(run, 13, bold=True, italic=True)
        # Linha de borda abaixo — igual ao documento de referência
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom_el = OxmlElement('w:bottom')
        bottom_el.set(qn('w:val'), 'single')
        bottom_el.set(qn('w:sz'), '6')
        bottom_el.set(qn('w:space'), '1')
        bottom_el.set(qn('w:color'), 'A72C31')
        pBdr.append(bottom_el)
        pPr.append(pBdr)

    elif tipo == 'cnpj_nire':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(linha.strip())
        set_font(run, 14, bold=True)

    elif tipo == 'data':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(linha.strip())
        set_font(run, 12)

    elif tipo == 'cargo':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(linha.strip())
        set_font(run, 12, italic=True)

    elif tipo == 'romano':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        m = re.match(r'^((?:I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII)\s*[–\-]\s*)(.+)', linha.strip())
        if m:
            run1 = p.add_run(m.group(1))
            set_font(run1, 12, bold=True)
            run2 = p.add_run(m.group(2))
            set_font(run2, 12, bold=True)
        else:
            run = p.add_run(linha.strip())
            set_font(run, 12, bold=True)

    elif tipo == 'clausula':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.keep_with_next = True   # título nunca fica isolado no fim de página
        lu = linha.upper()
        if lu.startswith('PARÁGRAFO') or lu.startswith('PARAGRAFO'):
            partes = linha.split(':', 1)
            run1 = p.add_run(partes[0].rstrip() + ': ')
            set_font(run1, 12, bold=True)
            if len(partes) > 1:
                _add_corpo_runs(p, partes[1].strip(), 12)
        elif '–' in linha:
            partes = linha.split('–', 1)
            run1 = p.add_run(partes[0].rstrip() + ' – ')
            set_font(run1, 12, bold=True)
            _add_corpo_runs(p, partes[1].strip(), 12)
        elif ' — ' in linha:
            partes = linha.split(' — ', 1)
            run1 = p.add_run(partes[0].rstrip() + ' — ')
            set_font(run1, 12, bold=True)
            _add_corpo_runs(p, partes[1].strip(), 12)
        elif ':' in linha:
            partes = linha.split(':', 1)
            run1 = p.add_run(partes[0].rstrip() + ': ')
            set_font(run1, 12, bold=True)
            if len(partes) > 1:
                _add_corpo_runs(p, partes[1].strip(), 12)
        elif ' - ' in linha:
            partes = linha.split(' - ', 1)
            run1 = p.add_run(partes[0].rstrip() + ' - ')
            set_font(run1, 12, bold=True)
            _add_corpo_runs(p, partes[1].strip(), 12)
        else:
            run = p.add_run(linha.strip())
            set_font(run, 12, bold=True)

    elif tipo == 'ementa_label':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(linha.strip())
        set_font(run, 12, bold=True)

    elif tipo == 'ementa_item':
        # Bullet (•) com recuo — igual ao modelo PargrafodaLista
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
        run = p.add_run('\u2022  ' + linha.strip())
        set_font(run, 12, bold=True)

    elif tipo == 'maiusculas':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(linha.strip())
        set_font(run, 14, bold=True)

    else:  # corpo
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        l_strip = linha.strip()
        if '<<<' in l_strip and '>>>' in l_strip:
            # Marcadores de negrito: <<<NOME>>> → bold
            partes_bold = re.split(r'(<<<.+?>>>)', l_strip)
            for parte in partes_bold:
                if parte.startswith('<<<') and parte.endswith('>>>'):
                    run = p.add_run(parte[3:-3])
                    set_font(run, 12, bold=True)
                elif parte:
                    run = p.add_run(parte)
                    set_font(run, 12)
        else:
            run = p.add_run(l_strip)
            set_font(run, 12)

    return p


def limpar_texto(texto):
    """Converte marcadores Markdown: **texto** vira <<<texto>>> (negrito), resto remove."""
    texto = re.sub(r'\*\*(.+?)\*\*', r'<<<\1>>>', texto)  # bold → marcador interno
    texto = re.sub(r'__(.+?)__', r'<<<\1>>>', texto)
    texto = re.sub(r'\*(.+?)\*', r'\1', texto)             # itálico simples → remove
    texto = re.sub(r'^#{1,6}\s+', '', texto, flags=re.MULTILINE)
    return texto


def is_linha_tabela(linha):
    return linha.strip().startswith('|') and linha.strip().endswith('|') and len(linha.strip()) > 2


def is_separador_tabela(linha):
    l = linha.strip()
    return l.startswith('|') and bool(re.match(r'^[\|\-\s\:]+$', l))


def inserir_tabela_markdown(doc, linhas_tabela):
    """Converte bloco de linhas Markdown em tabela Word — visual idêntico ao PDF."""
    rows_raw = [l for l in linhas_tabela if not is_separador_tabela(l)]
    if not rows_raw:
        return
    rows = []
    for linha in rows_raw:
        cells = [c.strip() for c in linha.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    rows = [r + [''] * (num_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    # ── Larguras em twips (área útil: 21cm - 2×2.5cm = 16cm ≈ 9072 twips) ──
    LARGURAS_4COL = [4234, 1512, 2117, 1209]
    LARGURAS_3COL = [4234, 2117, 2721]
    LARGURAS_2COL = [4914, 4158]
    if num_cols == 4:
        col_twips = LARGURAS_4COL
    elif num_cols == 3:
        col_twips = LARGURAS_3COL
    elif num_cols == 2:
        col_twips = LARGURAS_2COL
    else:
        col_twips = [int(9072 / num_cols)] * num_cols
    total_twips = sum(col_twips)

    # ── tblPr: largura total + sem indentação + layout fixo ──────────────────
    tbl = table._tbl
    existing_tblPr = tbl.findall(qn('w:tblPr'))
    tblPr = existing_tblPr[0] if existing_tblPr else OxmlElement('w:tblPr')
    if not existing_tblPr:
        tbl.insert(0, tblPr)
    for el in tblPr.findall(qn('w:tblW')):
        tblPr.remove(el)
    for el in tblPr.findall(qn('w:tblInd')):
        tblPr.remove(el)
    for el in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(el)
    for el in tblPr.findall(qn('w:jc')):
        tblPr.remove(el)
    tblW_el = OxmlElement('w:tblW')
    tblW_el.set(qn('w:w'), str(total_twips))
    tblW_el.set(qn('w:type'), 'dxa')
    tblPr.append(tblW_el)
    tblInd_el = OxmlElement('w:tblInd')
    tblInd_el.set(qn('w:w'), '0')
    tblInd_el.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd_el)
    tblLayout_el = OxmlElement('w:tblLayout')
    tblLayout_el.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout_el)
    jc_el = OxmlElement('w:jc')
    jc_el.set(qn('w:val'), 'left')
    tblPr.append(jc_el)

    for i, row_data in enumerate(rows):
        is_header = (i == 0)
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # Largura da célula
            if j < len(col_twips):
                for existing_tcW in tcPr.findall(qn('w:tcW')):
                    tcPr.remove(existing_tcW)
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(col_twips[j]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
            # Fundo cinza no cabeçalho (igual ao PDF)
            if is_header:
                for existing_shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(existing_shd)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F2F2')
                tcPr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if is_header:
                run = p.add_run(cell_text)
                set_font(run, 11, bold=True)
            else:
                _add_corpo_runs(p, cell_text, 11)

    # ── cantSplit: impede que linhas da tabela sejam cortadas entre páginas ────
    for row in table.rows:
        tr = row._tr
        trPr_list = tr.findall(qn('w:trPr'))
        if trPr_list:
            trPr = trPr_list[0]
        else:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), '1')
        trPr.append(cantSplit)


def detectar_assinatura(linha):
    """Verifica se a linha é o nome de um sócio na seção de assinatura."""
    import re
    l = linha.strip()
    # Suporte a marcadores de negrito <<<NOME>>>
    if l.startswith('<<<') and l.endswith('>>>'):
        l = l[3:-3].strip()
    # Nome de sócio: tudo maiúsculas, pelo menos 2 palavras, sem pontuação estranha
    if l == l.upper() and 5 < len(l) < 80 and re.match(r'^[A-ZÁÉÍÓÚÃÕÂÊÎÔÛÀÇÜÑ ]+$', l):
        return True
    return False


def gerar_docx(texto_gerado, nome_empresa, tem_consolidacao=True):
    doc = Document()

    # Margens exatas do MINUTA 5ª ALTERACAO: top=2.0, bottom=4.0, left=2.5, right=2.5
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(4.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        configurar_header(section)
        configurar_footer(section)

    # Separar as partes (1: instrumento, 3: consolidação)
    partes = {'parte1': '', 'parte2': '', 'parte3': ''}

    if '===PARTE 1:' in texto_gerado:
        blocos = texto_gerado.split('===PARTE ')
        for bloco in blocos[1:]:
            if bloco.startswith('1:'):
                fim = bloco.find('===PARTE 2:') if '===PARTE 2:' in bloco else len(bloco)
                partes['parte1'] = bloco[bloco.find('\n')+1:fim].strip()
            elif bloco.startswith('2:'):
                fim = bloco.find('===PARTE 3:') if '===PARTE 3:' in bloco else len(bloco)
                partes['parte2'] = bloco[bloco.find('\n')+1:fim].strip()
            elif bloco.startswith('3:'):
                partes['parte3'] = bloco[bloco.find('\n')+1:].strip()
    else:
        partes['parte1'] = texto_gerado

    for chave in ['parte1', 'parte2', 'parte3']:
        conteudo = partes.get(chave, '').strip()
        if not conteudo:
            continue
        # Se não tem consolidação, pular parte3 mesmo que exista no texto
        if chave == 'parte3' and not tem_consolidacao:
            continue

        # Limpar formatação Markdown antes de processar
        conteudo = limpar_texto(conteudo)

        em_assinaturas = False
        apos_ementa = False
        buf_tabela = []   # buffer para acumular linhas de tabela Markdown

        linhas = conteudo.split('\n')
        idx = 0
        while idx < len(linhas):
            linha = linhas[idx]
            l = linha.strip()

            # Acumular tabela Markdown
            if is_linha_tabela(l):
                buf_tabela.append(l)
                idx += 1
                continue
            elif buf_tabela:
                # Fim da tabela — renderizar
                inserir_tabela_markdown(doc, buf_tabela)
                buf_tabela = []

            # Detectar início da seção de assinaturas
            if re.match(r'^Goi[aâ]nia/GO,', l, re.IGNORECASE):
                em_assinaturas = True

            # Detectar "Ementa:"
            if re.match(r'^Ementa\s*:', l, re.IGNORECASE):
                apos_ementa = True
                adicionar_linha_doc(doc, linha)
                idx += 1
                continue

            # Itens de ementa: qualquer linha curta enquanto apos_ementa=True
            # (antes de aparecer cláusula, romano, CNPJ, linha de qualificação longa)
            _eh_clausula = re.match(
                r'^(CLÁUSULA|CLAUSULA|ARTIGO|ART\.|PARÁGRAFO|PARAGRAFO'
                r'|I{1,3}V?|VI{0,3}|IX|X{1,3}|CNPJ|N\.I\.R\.E|NIRE)',
                l, re.IGNORECASE)
            if apos_ementa and l and not _eh_clausula and len(l) < 120:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run('\u2022  ' + l)
                set_font(run, 12, bold=True)
                # Último item (termina com ".") encerra bloco de ementa
                if l.endswith('.'):
                    apos_ementa = False
                idx += 1
                continue
            elif apos_ementa and (not l or _eh_clausula):
                apos_ementa = False

            # Assinaturas: nome em MAIÚSCULAS após a data
            if em_assinaturas and detectar_assinatura(l):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(0)
                # Remove marcadores <<<>>> se presentes
                nome_limpo = l[3:-3].strip() if (l.startswith('<<<') and l.endswith('>>>')) else l
                run = p.add_run(nome_limpo)
                set_font(run, 12, bold=True)
            else:
                adicionar_linha_doc(doc, linha)

            idx += 1

        # Renderizar tabela pendente ao final da parte
        if buf_tabela:
            inserir_tabela_markdown(doc, buf_tabela)
            buf_tabela = []

        # Quebra de página entre partes
        prox = 'parte2' if chave == 'parte1' else 'parte3'
        if chave != 'parte3' and partes.get(prox, '').strip():
            doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route('/login', methods=['GET', 'POST'])
def login():
    erro = None
    if request.method == 'POST':
        senha = request.form.get('senha', '')
        if senha == SENHA_ACESSO:
            session['autenticado'] = True
            return redirect(url_for('index'))
        erro = 'Senha incorreta. Tente novamente.'
    return render_template('login.html', erro=erro)


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


@app.route('/')
def index():
    if login_obrigatorio():
        return redirect(url_for('login'))
    resp = render_template('index.html')
    from flask import make_response
    r = make_response(resp)
    r.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
    return r


@app.route('/upload', methods=['POST'])
def upload():
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401
    if 'contrato' not in request.files:
        return jsonify({'erro': 'Nenhum arquivo enviado'}), 400

    arquivo = request.files['contrato']
    if arquivo.filename == '':
        return jsonify({'erro': 'Nenhum arquivo selecionado'}), 400

    if not allowed_file(arquivo.filename):
        return jsonify({'erro': 'Formato inválido. Use PDF ou DOCX'}), 400

    filename = secure_filename(arquivo.filename)
    extensao = filename.rsplit('.', 1)[1].lower()

    # Salvar temporariamente
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{extensao}') as tmp:
            tmp_path = tmp.name
            arquivo.save(tmp.name)
        texto = extrair_texto(tmp_path, extensao)
    except Exception as e:
        return jsonify({'erro': f'Erro ao ler o arquivo: {str(e)}'}), 400
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

    if not texto.strip():
        return jsonify({'erro': 'Não foi possível extrair texto do arquivo. Verifique se o PDF não é escaneado.'}), 400

    # Extrair nome da empresa (primeiras linhas do contrato)
    linhas = [l.strip() for l in texto.split('\n') if l.strip()]
    nome_empresa = linhas[0] if linhas else 'EMPRESA'

    return jsonify({
        'sucesso': True,
        'texto': texto,
        'nome_empresa': nome_empresa,
        'tamanho': len(texto)
    })


@app.route('/gerar', methods=['POST'])
def gerar():
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401
    dados = request.get_json()
    if not dados:
        return jsonify({'erro': 'Dados inválidos'}), 400

    texto_contrato = dados.get('texto_contrato', '')
    alteracoes = dados.get('alteracoes', [])
    nome_empresa = dados.get('nome_empresa', 'EMPRESA')

    if not texto_contrato:
        return jsonify({'erro': 'Contrato não encontrado'}), 400
    if not alteracoes:
        return jsonify({'erro': 'Nenhuma alteração informada'}), 400

    # Detectar se consolidação foi selecionada
    tem_consolidacao = any(a.get('tipo') == 'Consolidação' for a in alteracoes)

    try:
        texto_gerado = gerar_com_claude(texto_contrato, alteracoes, tem_consolidacao)
    except APIIndisponivel as e:
        return jsonify({'erro': str(e)}), 503
    except ValueError as e:
        return jsonify({'erro': str(e)}), 400
    except Exception as e:
        logger.exception("Erro inesperado em gerar_com_claude (app.py)")
        return jsonify({'erro': 'Erro interno ao processar o contrato. Tente novamente ou contate o administrador.'}), 500

    try:
        buffer = gerar_docx(texto_gerado, nome_empresa, tem_consolidacao)
    except Exception as e:
        return jsonify({'erro': f'Erro ao gerar documento Word: {str(e)}'}), 500

    # Nome do arquivo: número ordinal por extenso + "alteração contratual"
    nome_arquivo = extrair_ordinal_filename(texto_gerado)

    # Gerar PDF também
    pdf_erro = None
    try:
        buffer_pdf = gerar_pdf(texto_gerado, nome_empresa, tem_consolidacao)
        pdf_bytes = buffer_pdf.read()
    except Exception as e_pdf:
        import traceback
        pdf_bytes = None
        pdf_erro = traceback.format_exc()
        print(f'[AVISO] Falha ao gerar PDF: {e_pdf}\n{pdf_erro}', flush=True)

    # Guardar no cache com token único
    _limpar_cache_antigo()
    token = str(uuid.uuid4())
    _DOCS_CACHE[token] = {
        'docx': buffer.read(),
        'pdf':  pdf_bytes,
        'nome': nome_arquivo,
        'ts':   time.time(),
    }

    return jsonify({'token': token, 'nome': nome_arquivo, 'pdf_ok': pdf_bytes is not None})


@app.route('/download/<token>/<formato>')
def download(token, formato):
    if login_obrigatorio():
        return jsonify({'erro': 'Não autorizado'}), 401
    entrada = _DOCS_CACHE.get(token)
    if not entrada:
        return jsonify({'erro': 'Documento expirado. Gere novamente.'}), 404

    nome_base = entrada['nome'].replace('.docx', '')

    if formato == 'docx':
        return send_file(
            io.BytesIO(entrada['docx']),
            as_attachment=True,
            download_name=entrada['nome'],
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    elif formato == 'pdf':
        if not entrada.get('pdf'):
            return jsonify({'erro': 'PDF não disponível para este documento.'}), 500
        return send_file(
            io.BytesIO(entrada['pdf']),
            as_attachment=True,
            download_name=f'{nome_base}.pdf',
            mimetype='application/pdf'
        )
    else:
        return jsonify({'erro': 'Formato inválido'}), 400


if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    port = int(os.getenv('PORT', 5080))
    app.run(host='0.0.0.0', port=port, debug=False)
